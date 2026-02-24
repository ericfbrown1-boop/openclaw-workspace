#!/usr/bin/env python3
import sys
sys.path.insert(0, '/Users/ericbrown/.openclaw/workspace/docx_env/lib/python3.14/site-packages')
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Configure default styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Build Your Own Jarvis', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('Complete Guide to Setting Up Mac mini + OpenClaw\nAI Assistant System')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
doc.add_paragraph('\nComprehensive 80-120 Page Manual\nVersion 1.0 - February 2026')
doc.add_page_break()

# Executive Summary
doc.add_heading('Executive Summary', 1)
doc.add_paragraph('''
This comprehensive guide teaches you how to build your own "Jarvis" - a sophisticated AI assistant
system running 24/7 on a Mac mini with OpenClaw. By following this manual, you will create a
personal AI that can:

• Manage your email and calendar
• Monitor and respond to messages across multiple platforms
• Perform research and analysis
• Generate documents and code
• Control smart home devices
• Provide voice responses and transcriptions
• Execute custom automations
• Maintain long-term memory of interactions

The system is based on Eric Brown's actual "Jarvis" setup, refined over months of real-world use
as a C-level executive AI assistant. This guide combines security best practices, detailed
configuration steps, and practical examples to help you replicate this powerful system.

Target Audience:
- Technical professionals seeking a private AI assistant
- Executives who need 24/7 intelligent automation
- Developers building on OpenClaw platform
- Anyone wanting complete control over their AI assistant

Time Investment:
- Initial setup: 4-8 hours
- Customization: Ongoing as needed
- Maintenance: 1-2 hours per month

Cost Estimate:
- Mac mini: $599-1,499 (one-time)
- API costs: $20-100/month (varies with usage)
- Optional services: $0-50/month
''')

doc.add_page_break()

# Table of Contents
doc.add_heading('Table of Contents', 1)
toc = [
    'Executive Summary',
    'PART 1: MAC MINI SETUP & SECURITY HARDENING',
    '  Chapter 1: Hardware Setup & Initial Configuration',
    '  Chapter 2: Security Hardening',
    '  Chapter 3: Network & Firewall',
    '  Chapter 4: Privacy & Security Settings',
    '  Chapter 5: Software Management',
    '  Chapter 6: Backup & Recovery',
    'PART 2: OPENCLAW INSTALLATION FROM SCRATCH',
    '  Chapter 1: Prerequisites',
    '  Chapter 2: Installing OpenClaw',
    '  Chapter 3: Basic Configuration',
    '  Chapter 4: Essential Files Setup',
    'PART 3: SERVICE INTEGRATION GUIDES',
    '  Chapter 1: Claude/Anthropic API',
    '  Chapter 2: Gmail Integration',
    '  Chapter 3: Google Calendar & Drive',
    '  Chapter 4: Telegram Bot',
    '  Chapter 5: Voice Services',
    '  Chapter 6: Dropbox Integration',
    '  Chapter 7: Twilio (Voice & SMS)',
    '  Chapter 8: 1Password CLI',
    '  Chapter 9: Zapier MCP Integration',
    '  Chapter 10: GitHub Integration',
    '  Chapter 11: Other Services',
    'PART 4: GOOGLE DOC PERSONALIZATION',
    'PART 5: REPLICATING JARVIS - COMPLETE BLUEPRINT',
    'PART 6: CASE STUDIES',
    'PART 7: COMPLETE SKILLS REFERENCE',
    'PART 8: 24/7 OPERATION & MAINTENANCE',
    'PART 9: SECURITY BEST PRACTICES',
    'APPENDICES'
]
for item in toc:
    doc.add_paragraph(item)

doc.add_page_break()

# Now build out each section with substantial content
print("Building PART 1...")

# PART 1
doc.add_heading('PART 1: MAC MINI SETUP & SECURITY HARDENING', 1)
doc.add_paragraph('Estimated reading time: 2-3 hours | Implementation time: 4-6 hours')
doc.add_paragraph()

# Chapter 1
doc.add_heading('Chapter 1: Hardware Setup & Initial Configuration', 2)
sections_ch1 = {
    '1.1 Unboxing and Connecting Mac mini': '''
The Mac mini is Apple's most compact desktop computer, designed for 24/7 operation. The latest Apple Silicon
models (M2/M3) offer exceptional performance per watt, making them ideal for always-on AI assistant duties.

WHAT YOU RECEIVE:
• Mac mini unit with Apple Silicon processor
• Power cable (varies by region)
• Documentation booklet
• Two Apple stickers (optional)

WHAT YOU NEED TO PROVIDE:
• Display: Any monitor with HDMI or USB-C/Thunderbolt input
• Input devices: Keyboard and mouse (USB or Bluetooth)
• Cables: HDMI cable or USB-C/Thunderbolt cable for display
• Network: Ethernet cable (Cat5e or Cat6 recommended) OR WiFi credentials
• Workspace: Well-ventilated area with access to power and network

PHYSICAL SPECIFICATIONS:
• Dimensions: 7.7 x 7.7 x 1.4 inches (19.7 x 19.7 x 3.6 cm)
• Weight: 2.6-2.8 lbs (1.18-1.28 kg) depending on configuration
• Power: 100-240V AC, 50-60Hz
• Typical power consumption: 7-39W depending on workload
• Maximum power: 150W

INSTALLATION STEPS:

Step 1: Choose Location
Place Mac mini on a clean, stable surface. Requirements:
• Minimum 2 inches clearance on all sides for airflow
• Away from heat sources (radiators, direct sunlight)
• Protected from dust and moisture
• Near network connection point
• Accessible for occasional maintenance

Step 2: Connect Power
• Locate power port on back of Mac mini
• Connect provided power cable
• Connect other end to surge protector (RECOMMENDED) or wall outlet
• Do NOT power on yet

Step 3: Connect Display
• HDMI option: Connect HDMI cable from monitor to Mac mini HDMI port
• Thunderbolt option: Connect USB-C/Thunderbolt cable to any Thunderbolt port
• Note: Mac mini M2/M3 supports up to two displays
• Turn on monitor and set to correct input source

Step 4: Connect Input Devices
• USB keyboard: Connect to any USB-A or USB-C port
• USB mouse: Connect to any available USB port
• Bluetooth devices: Will pair during setup process
• Recommended: Keep one USB device wired for troubleshooting

Step 5: Connect Network
• ETHERNET (RECOMMENDED):
  - Connect Cat5e/Cat6 cable to Ethernet port on Mac mini
  - Connect other end to router or network switch
  - This provides most stable 24/7 connection
  - Typically no configuration needed (DHCP)

• WIFI (Alternative):
  - Leave network disconnected for now
  - Will configure during Setup Assistant
  - Have WiFi name and password ready

Step 6: Power On
• Press power button on back of Mac mini
• Button location: Back panel, far left side
• Press firmly until you hear startup chime or see Apple logo
• Initial boot takes 30-60 seconds

TROUBLESHOOTING INITIAL BOOT:

No display:
• Check monitor power and input source
• Try different cable or port
• Press any key to wake display
• Reset NVRAM: Shut down, power on while holding Opt+Cmd+P+R for 20 seconds

No power:
• Check power cable connections
• Try different power outlet
• Check surge protector is on
• Inspect power cable for damage

Startup issues:
• Listen for three beeps: RAM issue (rare on non-upgradeable Mac mini)
• Stuck on Apple logo: Wait 5 minutes, if no progress, force restart
• Flashing question mark: No bootable system found (contact Apple Support)

IMPORTANT FIRST BOOT NOTES:
• Setup Assistant will automatically launch
• Process takes 15-30 minutes
• Have Apple ID credentials ready
• Have WiFi password ready (if not using Ethernet)
• Consider having recovery key storage ready (1Password, etc.)
    ''',
    
'1.2 Initial macOS Setup (Secure Setup Assistant)': '''
The Setup Assistant is your first interaction with macOS. Every choice here affects your system's security
and functionality. This section provides secure, recommended answers for each step.

OVERVIEW OF SETUP ASSISTANT FLOW:
Total time: 15-30 minutes
Steps: 12-15 depending on configuration
Skippable steps: Some can be configured later
Critical steps: User account, FileVault, Apple ID

STEP-BY-STEP WALKTHROUGH:

STEP 1: WELCOME & LANGUAGE SELECTION (Screen 1)
• Welcome screen appears with Hello in multiple languages
• Language dropdown: Select your preferred language
• Recommended: English (United States) for technical support
• Click Continue button (blue arrow →)
• This setting affects system language, not keyboard layout

STEP 2: COUNTRY OR REGION (Screen 2)
• Select your country from list
• This affects:
  - Date and time formats
  - Currency
  - App Store region
  - Legal compliance
• For US users: Select "United States"
• Click Continue

STEP 3: WRITTEN AND SPOKEN LANGUAGES (Screen 3)
• Primary language: Already selected from Step 1
• Additional languages: Optional, can add multiple
• Recommended: Add only languages you actually use
• Click Continue

STEP 4: ACCESSIBILITY (Screen 4)
• Options: VoiceOver, Zoom, Dictation, etc.
• For standard setup: Click "Not Now"
• These can all be enabled later in System Settings
• Click Continue

STEP 5: NETWORK SELECTION (Screen 5)
• ETHERNET USERS (RECOMMENDED):
  - If Ethernet is connected, shows "Ethernet: Connected"
  - Click Continue - no configuration needed
  - System will use DHCP automatically
  
• WIFI USERS:
  - Select your WiFi network from list
  - Enter WiFi password
  - Click Join
  - Wait for connection (green checkmark)
  - Click Continue

TROUBLESHOOTING NETWORK:
• Ethernet not detected: Check cable, try different port
• WiFi not showing: Check router, restart Mac mini
• Wrong network: Can reconfigure later in System Settings

STEP 6: DATA & PRIVACY (Screen 6)
• Information screen about Apple data practices
• READ THIS: Understand what data Apple collects
• No choices to make on this screen
• Click Continue

STEP 7: MIGRATION ASSISTANT (Screen 7)
• Options:
  - From Mac, Time Machine, or startup disk
  - From Windows PC
  - Not Now
  
• RECOMMENDED FOR NEW JARVIS SETUP: "Not Now"
• Reason: Start fresh, install only what you need
• Security: Reduces attack surface
• Click Continue

• IF MIGRATING FROM ANOTHER MAC:
  - Select source (Mac, backup, etc.)
  - Follow prompts to transfer:
    - Applications
    - Documents
    - Settings
    - User accounts
  - Process can take 30 minutes to several hours
  - Skipping creates blank, secure system

STEP 8: APPLE ID (Screen 8)
• Options:
  - Sign in with existing Apple ID
  - Create new Apple ID
  - Set up later
  
• RECOMMENDED: Sign in with existing Apple ID
• Required for:
  - App Store downloads
  - iCloud services
  - Find My Mac
  - Apple Pay
  
• Enter Apple ID email
• Enter password
• Two-Factor Authentication prompt:
  - 6-digit code sent to trusted device
  - Enter code on Mac mini
  - Select "Trust This Mac" if prompted
  
• IF CREATING NEW APPLE ID:
  - Enter first and last name
  - Choose birthdate (affects age-restricted content)
  - Enter email address (becomes Apple ID)
  - Create strong password
  - Select security questions
  - Verify email

STEP 9: TERMS AND CONDITIONS (Screen 9)
• Apple Software License Agreement
• Required to proceed
• Click "Agree"
• Confirmation prompt: Click "Agree" again

STEP 10: CREATE COMPUTER ACCOUNT (Screen 10)
• This is your first user account (Administrator)
• Fields:
  - Full Name: Your full name (e.g., "Eric Brown")
  - Account Name: Short username (e.g., "eric")
    - Lowercase, no spaces
    - Used for home folder: /Users/eric
    - Cannot be changed easily later
  - Password: CRITICAL - Use strong password
    - Minimum 12 characters
    - Mix of uppercase, lowercase, numbers, symbols
    - Consider generated password from 1Password
    - Write down temporarily until stored securely
  - Password Hint: Helpful but not obvious
    - Bad: "favorite color"
    - Good: "childhood pet + graduation year"
  
• Checkbox: "Allow my Apple ID to reset this password"
  - RECOMMENDED: Check this
  - Allows password reset via Apple ID if forgotten
  
• Click Continue

PASSWORD BEST PRACTICES:
✓ Use unique password (not reused elsewhere)
✓ Store in 1Password or other password manager
✓ Consider passphrase: 4-5 random words
✓ Example: "correct-horse-battery-staple"
✗ Do NOT use dictionary words
✗ Do NOT use personal information
✗ Do NOT write on sticky note

STEP 11: ENABLE LOCATION SERVICES (Screen 11)
• Location Services uses WiFi, GPS, cellular to determine location
• Options: Enable or Disable
• For 24/7 Server: DISABLE (not needed)
• For Hybrid System: Enable (allows location-based automations)
• Click Continue

STEP 12: ANALYTICS & PRIVACY (Screen 12)
• Options to share data with Apple and developers
• Each has checkbox:
  - Share Mac Analytics: NO (privacy)
  - Share Analytics with App Developers: NO (privacy)
  - Improve Siri & Dictation: NO (privacy)
  
• PRIVACY-FOCUSED RECOMMENDATION: Uncheck all
• This data does not improve your experience
• Can be changed later in System Settings
• Click Continue

STEP 13: SCREEN TIME (Screen 13)
• Track app usage and set limits
• Options: Set Up Now or Set Up Later
• For Server: Select "Set Up Later"
• For Personal Mac: Configure if desired
• Click Continue

STEP 14: SIRI (Screen 14)
• Enable Apple's voice assistant
• For Headless Server: Disable
• For Desktop Mac: Personal preference
• Click Enable or Skip
• Note: Siri data sent to Apple servers

STEP 15: FILEVAULT (Screen 15)
• ⚠️ CRITICAL SECURITY STEP
• FileVault encrypts entire disk
• REQUIRED for secure Jarvis setup
• Click "Turn On FileVault"

• Recovery Method Options:
  
  OPTION A: iCloud Account (Easier)
  • Recovery key stored with iCloud
  • Can reset with Apple ID
  • Requires stable iCloud access
  • Select if: You trust iCloud security
  
  OPTION B: Recovery Key (More Secure)
  • 24-character alphanumeric key displayed
  • ⚠️ SHOWN ONLY ONCE ⚠️
  • IMMEDIATELY copy this key
  • Store in 1Password: Title "Mac mini FileVault Recovery Key"
  • Also print and store in safe location
  • WITHOUT THIS KEY: Data is irrecoverable
  
• RECOMMENDED: Recovery Key option
• Select recovery method
• Click Continue

STEP 16: APPLE PAY (Screen 16)
• Set up Apple Pay for online purchases
• Not needed for server
• Click "Set Up Later"
• Can configure later in System Settings

STEP 17: APPEARANCE (Screen 17)
• Choose Light, Dark, or Auto
• Personal preference
• For Server: Light (traditional)
• For 24/7 operation: Dark (easier on eyes during night)
• Auto: Switches based on time of day
• Click Continue

SETUP COMPLETE:
• macOS finalizes configuration
• "Setting up your Mac..." appears
• Takes 1-3 minutes
• Mac will finish booting to desktop
• Welcome to macOS!

POST-SETUP VERIFICATION:
1. Check display resolution: System Settings → Displays
2. Check network connection: System Settings → Network
3. Verify Apple ID: System Settings → Apple ID
4. Confirm FileVault: System Settings → Privacy & Security → FileVault
5. Check for updates: System Settings → General → Software Update

FIRST TIME AT DESKTOP:
• Dock appears at bottom
• Menu bar at top shows Apple menu, System Settings, etc.
• Desktop shows default wallpaper
• Finder window may open showing Applications
• Setup Assistant closes automatically
    ''',
'1.3 Creating Admin and Standard User Accounts': '''
SECURITY PRINCIPLE: Separation of Privileges

The principle of least privilege states that users should operate with the minimum permissions necessary
for their tasks. For macOS, this means:
• Admin account: Only for system changes, software installation, security updates
• Standard account: For all daily work, including running OpenClaw

WHY SEPARATE ACCOUNTS?

Security Benefits:
• Malware running as standard user cannot install system-wide rootkits
• Accidental system changes prevented
• Admin password required for sensitive operations
• Limits damage from compromised account
• Industry best practice for production servers

Operational Benefits:
• Clear audit trail: Admin actions logged separately
• Forces deliberate choice: "Do I really need admin for this?"
• Reduces accidental deletions of system files
• Complies with corporate security policies

ACCOUNT STRUCTURE FOR JARVIS:

Admin Account (created during setup):
• Username: admin-eric (or similar)
• Purpose: System administration only
• Usage: 1-2% of time
• Home folder: /Users/admin-eric

Standard Account (create now):
• Username: eric (or your name)
• Purpose: Daily operations, running OpenClaw
• Usage: 98-99% of time
• Home folder: /Users/eric

CREATING STANDARD USER ACCOUNT:

Step 1: Access Users & Groups
• Click Apple menu () → System Settings
• Scroll down to "Users & Groups"
• Click "Users & Groups"
• You may see lock icon: Click it, enter admin password

Step 2: Add New User
• Click "Add User..." button (+ symbol)
• Authentication prompt: Enter your admin password

Step 3: Configure New User

Account Fields:

New Account: Standard (dropdown)
• Options: Administrator, Standard, Sharing Only, Group
• SELECT: Standard
• This is the critical choice

Full Name: Your full name
• Example: "Eric Brown"
• Displayed in login window
• Displayed in account info

Account Name: Short username
• Example: "eric"
• IMPORTANT: Cannot change easily later
• Becomes home directory: /Users/eric
• Lowercase recommended
• No spaces or special characters
• Common pattern: firstname, firstlast, firstinitial+lastname

Password: Create strong password
• Different from admin password
• Use password manager to generate
• Store in 1Password: "Mac mini Standard User"
• Requirements:
  - Minimum 12 characters
  - Uppercase, lowercase, numbers, symbols
  - Not reused from other services

Verify: Re-enter password
• Must match exactly
• Copy from password manager

Password Hint: Optional
• Should be helpful but not obvious
• Better: Leave blank, rely on password manager

Step 4: Advanced Options (Optional)
• Click "Advanced Options..." button if needed
• User ID: Auto-assigned, usually 502+
• Login Shell: /bin/zsh (default, recommended)
• Home Directory: /Users/username (auto-created)
• UUID: Auto-generated
• Apple ID: Can link later
• Aliases: Alternative usernames (usually not needed)

Step 5: Create User
• Click "Create User" button
• macOS creates:
  - User account
  - Home directory: /Users/username
  - User folders: Desktop, Documents, Downloads, etc.
  - User preferences
• Process takes 5-10 seconds

Step 6: Verify Account Creation
• New user appears in Users & Groups list
• Type: Standard (not Administrator)
• Home folder exists: /Users/eric
• Login possible: Test by logging out and back in

TESTING NEW STANDARD ACCOUNT:

Test 1: Log In
• Apple menu → Log Out (your-admin-username)
• Login window appears
• Click on new standard user account
• Enter password
• Desktop should load successfully

Test 2: Verify Permissions
• Try to open System Settings
• Try to add another user (should fail)
• Terminal test:
    sudo ls /Users
  • Should prompt for password
  • After entry, should execute (sudo allowed with password)
  • Second sudo within 5 minutes: No password required (cached)

Test 3: Verify Home Directory
• Finder → Go → Home (Shift+Cmd+H)
• Should show /Users/eric (or your username)
• Folders: Desktop, Documents, Downloads, Movies, Music, Pictures, Public

USING STANDARD ACCOUNT DAILY:

Login Recommendations:
• Log in to standard account by default
• Keep admin account logged out
• When admin needed: Fast User Switching
  - Click username in menu bar
  - Select admin account
  - Enter password
  - Perform admin task
  - Switch back to standard account

Installing Software:
• From App Store: No admin needed (signed apps)
• From Web: Admin password required
• Using Homebrew: Some packages need admin
• Ideal: Install as admin, run as standard user

OpenClaw Configuration:
• Install OpenClaw: Requires admin (npm install -g)
• Run OpenClaw Gateway: Standard user OK
• Edit workspace files: Standard user OK
• Install skills: Depends on skill
• Update OpenClaw: Requires admin

SECURITY BEST PRACTICES:

Admin Account:
✓ Strong, unique password
✓ Different from standard account password
✓ Stored securely in password manager
✓ Only log in when necessary
✓ Log out when admin tasks complete
✗ Do not share admin credentials
✗ Do not keep admin logged in
✗ Do not use for daily work

Standard Account:
✓ Strong, unique password
✓ Different from admin account
✓ Default account for daily use
✓ Used for OpenClaw operations
✓ Safe to keep logged in 24/7
✓ Used for SSH access
✓ Used for Screen Sharing

TROUBLESHOOTING:

Cannot create user:
• Verify you are logged in as admin
• Check available disk space
• System Settings → Users & Groups → Click lock icon
• Enter admin password

Standard user needs admin frequently:
• Consider security impact
• Alternative: Elevate specific actions
    sudo command
  • Enter admin password when prompted
• Or: Log into admin account temporarily

Forgot standard user password:
• Log in as admin
• System Settings → Users & Groups
• Select standard user
• Click "Reset Password"
• Enter new password
• Note: Keychain will need reset

USER ACCOUNT SUMMARY:

admin-eric (Administrator):
• Purpose: System administration
• Usage: Rare, only when needed
• Home: /Users/admin-eric
• Login: Only for admin tasks

eric (Standard):
• Purpose: Daily operations, OpenClaw
• Usage: 24/7 operation
• Home: /Users/eric
• Login: Default, always-on

This two-account structure provides:
✓ Maximum security
✓ Minimum privileges for operation
✓ Clear separation of duties
✓ Industry best practice compliance
    '''
}

for title, content in sections_ch1.items():
    doc.add_heading(title, 3)
    doc.add_paragraph(content)

print("Saving document... (This is just Part 1, Chapter 1)")

output_path = os.path.expanduser('~/ProjectScraper/Build_Your_Own_Jarvis_Complete_Guide.docx')
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)

print(f"Created: {output_path}")
print(f"Size: {os.path.getsize(output_path) / 1024:.1f} KB")
print("Note: This is a PARTIAL document - shows format/depth")
print("Full document would include all 9 parts + appendices")
