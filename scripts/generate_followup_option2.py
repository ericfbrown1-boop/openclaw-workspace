"""Generate Follow-Up Tracker Option 2 Execution Plan Word Doc."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)


def h1(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0, 63, 128)


def h2(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0, 100, 160)


def tbl(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hd in enumerate(headers):
        c = table.rows[0].cells[j]
        c.text = hd
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = table.rows[i + 1].cells[j]
            c.text = str(val)
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(9)
    doc.add_paragraph()


def bul(t):
    p = doc.add_paragraph(t, style="List Bullet")
    p.runs[0].font.size = Pt(10)


def bod(t):
    p = doc.add_paragraph(t)
    p.runs[0].font.size = Pt(10)


def code(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.font.size = Pt(9)
    r.font.name = "Courier New"
    r.font.color.rgb = RGBColor(50, 50, 50)


# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PROJECT FOLLOW-UP TRACKER")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(0, 63, 128)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Option 2: Claude CoWork + M365 Outlook Connector")
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(80, 80, 80)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run(f"Detailed Execution Plan | {datetime.now().strftime('%B %d, %Y')}")
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(120, 120, 120)
r3.italic = True

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("Prepared by Jarvis AI | For: Eric Brown, CFO & COO, Cohesity")
r4.font.size = Pt(9)
r4.font.color.rgb = RGBColor(150, 150, 150)
r4.italic = True

doc.add_page_break()

# Prerequisites
h1("1. Prerequisites Checklist")
tbl(["Requirement", "Status", "Action if Missing"],
    [("Claude Desktop app on work PC", "Check", "Download from https://claude.ai/download"),
     ("Claude Pro/Max/Team/Enterprise plan", "Check", "Upgrade at claude.ai/settings/billing"),
     ("M365 Outlook with Eric.brown@cohesity.com", "Ready", "Already configured"),
     ("Admin rights on work PC", "Check", "IT request if locked down"),
     ("Internet access on work PC", "Ready", "Standard corporate network")])

# Step 1
h1("2. Step-by-Step Implementation")

h2("Step 1: Install & Enable Claude Desktop CoWork")
bod("If Claude Desktop is not already on your work PC:")
bul("Go to https://claude.ai/download and download for your OS")
bul("Run installer, sign in with your Anthropic account")
bul("Open Settings (profile icon, bottom-left) > Cowork")
bul('Toggle Cowork ON -- "Cowork" appears in the left sidebar')
bul("CoWork requires Desktop app -- browser version does not support it")

h2("Step 2: Connect Microsoft 365 / Outlook Connector")
bod("This gives Claude read access to your Cohesity Outlook emails:")
bul("Settings > Cowork > Connectors")
bul('Find "Microsoft 365" or "Outlook" connector')
bul("Click Connect / Authorize")
bul("Sign in with Eric.brown@cohesity.com at the Microsoft OAuth page")
bul("Review permissions: Mail.Read, Mail.ReadBasic (read-only)")
bul("Click Accept / Grant Access")
bul('Verify connector shows "Connected" in Settings')

bod("SECURITY: This is READ-ONLY. Claude cannot send emails or modify your mailbox.")

bod("If M365 connector not yet available: It is rolling out to Enterprise/Team plans "
    "(March-April 2026). Check for Claude Desktop updates. As fallback, connect Gmail "
    "connector with ericfbrown1@gmail.com while waiting.")

h2("Step 3: Connect Gmail Connector (for outbound capability)")
bod("Since M365 is read-only, Gmail provides the outbound follow-up channel:")
bul("Settings > Cowork > Connectors > Gmail")
bul("Click Connect, sign in with ericfbrown1@gmail.com")
bul("Grant read + search permissions")
bul("This lets CoWork draft follow-up emails you can review and send")

h2("Step 4: Create the Follow-Up Tracker Project")
bod('In CoWork sidebar, click "+ New Project" and name it "Follow-Up Tracker"')
bod("Add these project instructions (paste exactly):")
code(
    "PROJECT: Automated Follow-Up Tracker\n\n"
    "PURPOSE: Monitor emails sent by Eric Brown with 'Follow Up Requested' in the "
    "subject line. Track all action items, owners, and deadlines to completion.\n\n"
    "TRIGGER: Any email from Eric.brown@cohesity.com with subject containing "
    "'Follow Up Requested'\n\n"
    "WHEN TRIGGERED:\n"
    "1. Read full email body\n"
    "2. Extract: project name, owners (To/CC), deliverables, due dates, criteria\n"
    "3. Create tracker file: [ProjectName]_tracker.md with columns:\n"
    "   Task | Owner | Due Date | Status | Last Activity | Days Silent | "
    "Next Action | Nudge Count\n"
    "4. Status values: NOT STARTED / IN PROGRESS / BLOCKED / COMPLETE\n\n"
    "DAILY MONITORING:\n"
    "1. Search Outlook for replies in each tracked thread\n"
    "2. Analyze replies for progress, blockers, commitments\n"
    "3. Update tracker; flag YELLOW (2+ days silent) or RED (4+ days)\n"
    "4. Draft escalation emails for Eric's review (never auto-send)\n\n"
    "COMPLETION: Mark COMPLETE when all deliverables done. Generate summary."
)

h2("Step 5: Create Daily Monitoring Scheduled Task")
bod("Inside the Follow-Up Tracker project:")
bul('Click "+ New Task"')
bul("Enter this prompt:")
code(
    "Scan my Outlook inbox for emails I sent with 'Follow Up Requested' in the "
    "subject (last 30 days). For each: check for new replies, update tracker, "
    "identify at-risk (2+ days no response) or stalled (4+ days) items, "
    "draft follow-ups for stalled items, and give me a summary: "
    "GREEN / YELLOW / RED for each project."
)
bul("Type /schedule in the chat")
bul("Set: Daily at 9:00 AM PT")
bul('Name: "Daily Follow-Up Scan"')
bul("Click Create Schedule")

h2("Step 6: Create New Project Detection Task")
bod("Create a second scheduled task to catch new trigger emails:")
bul("Same project, new task")
bul("Prompt:")
code(
    "Search my Outlook sent folder for NEW emails with 'Follow Up Requested' "
    "sent in the last 2 hours, not already tracked. Parse each: extract project "
    "name, owners, deliverables, timeline. Create a new tracker file and confirm."
)
bul("Schedule: Every 2 hours, 8 AM - 6 PM PT")
bul('Name: "New Project Detection"')

h2("Step 7: Test with a Sample Email")
bod("From Cohesity Outlook, compose:")
bul("To: yourself (or a colleague)")
bul("Subject: Follow Up Requested: Test Project - Quarterly Report")
bul("Body: Assign 2-3 deliverables with due dates")
bul("Send it")
bul("Wait for the next detection scan (or run manually in CoWork)")
bul("Check the project folder for the new tracker file")
bul("Verify extracted data is accurate; refine instructions if needed")

# Daily workflow
h1("3. Your Daily Workflow (Once Live)")
tbl(["Time", "What Happens", "Your Action"],
    [("Any time", 'You send a "Follow Up Requested" email', "Just send normally"),
     ("Within 2 hours", "CoWork detects and parses the email", "Automatic"),
     ("9:00 AM daily", "CoWork scans all active threads", "Review daily summary"),
     ("If items stalled", "CoWork drafts follow-up emails", "Review, edit, send"),
     ("When complete", "CoWork marks done, generates summary", "Archive")])

# Escalation
h1("4. Escalation Ladder")
tbl(["Trigger", "Action", "Tone"],
    [("Day 1: No response", "Grace period -- no action", "-"),
     ("Day 2: Still silent", "CoWork drafts gentle check-in for your review", "Friendly"),
     ("Day 3: No progress", "CoWork drafts firmer follow-up", "Professional"),
     ("Day 4+: Stalled", "RED alert in your daily summary", "Direct"),
     ("Past due date", "CoWork drafts overdue notice for your review", "Urgent")])
bod("Important: CoWork DRAFTS all emails for your review. You decide whether to send, "
    "edit, or skip. No auto-sending.")

# Enhancements
h1("5. Optional Enhancements")
h2("Slack Connector")
bul("Settings > Cowork > Connectors > Slack")
bul("Authorize Cohesity Slack workspace")
bul("Add to scheduled task: post alerts to #follow-ups channel")

h2("Google Drive Connector")
bul("Connect Google Drive to detect when deliverables are uploaded")
bul("Auto-mark tasks COMPLETE when files appear in specified folder")

h2("Weekly Summary Report")
bul("Create weekly scheduled task (Friday 4 PM)")
bul('Prompt: "Generate weekly summary of all active Follow Up Requested projects"')

# Troubleshooting
h1("6. Troubleshooting")
tbl(["Issue", "Solution"],
    [("M365 connector not showing", "Update Claude Desktop; may need Team/Enterprise plan"),
     ("CoWork not available", "Must use Desktop app, not browser; toggle in Settings"),
     ("Scheduled tasks not running", "Claude Desktop must be open + computer awake"),
     ("Connector disconnects", "Re-authorize in Settings > Cowork > Connectors"),
     ("Missing emails in scan", "Check permissions; ensure Mail.Read scope granted")])

# Plan limits
h1("7. Plan Limits")
tbl(["Feature", "Pro ($20/mo)", "Team ($30/seat)", "Enterprise"],
    [("Active connectors", "Up to 5", "Unlimited", "Unlimited"),
     ("Recurring tasks", "10", "50 per user", "Unlimited"),
     ("M365 connector", "Rolling out", "Available", "Available"),
     ("Gmail connector", "Available", "Available", "Available"),
     ("Admin controls", "No", "Yes", "Yes + SSO/audit")])

# Timeline
h1("8. Implementation Timeline")
tbl(["Day", "Action"],
    [("Day 1", "Install Claude Desktop, enable CoWork, connect M365 + Gmail"),
     ("Day 2", "Create Follow-Up Tracker project with instructions"),
     ("Day 3", "Set up both scheduled tasks (daily scan + detection)"),
     ("Day 4", "Send test email, verify detection and parsing works"),
     ("Day 5", "Refine prompts based on test results"),
     ("Day 6-7", 'Go live with first real "Follow Up Requested" email')])

# Dropbox link
h1("9. Dropbox Link")
bod("The full execution plan (this document) plus the markdown version with "
    "copy-pasteable prompts is stored in Dropbox:")
p = doc.add_paragraph()
r = p.add_run(
    "https://www.dropbox.com/scl/fi/py2p0jup84cpd7jg7u3la/"
    "FollowUp_Tracker_Option2_Execution_Plan.md"
    "?rlkey=5p14klqfp9rdq1uxs7a6iy116&dl=0"
)
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0, 100, 200)
r.underline = True

bod("The markdown version has all prompts and instructions ready to copy-paste "
    "directly into Claude Desktop CoWork on your work PC.")

# Footer
doc.add_paragraph()
pf = doc.add_paragraph(
    "Prepared by Jarvis AI (Research + Planning Agents). "
    "Ready to assist with testing and refinement once you begin setup.\n\n"
    "Sent by Jarvis - AI assistant to Eric Brown"
)
for r in pf.runs:
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(150, 150, 150)
    r.italic = True

os.makedirs(os.path.expanduser("~/Documents/reports"), exist_ok=True)
out = os.path.expanduser(
    "~/Documents/reports/FollowUp_Tracker_Option2_Plan_20260405.docx"
)
doc.save(out)
print(f"Saved: {out}")
