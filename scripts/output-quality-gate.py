#!/usr/bin/env python3
"""
output-quality-gate.py — Automated output quality validation.
Verifies deliverables contain real content, not placeholders.
Run after any pipeline generates a deliverable.

Usage:
    python3 output-quality-gate.py --docx <path>     # Validate .docx report
    python3 output-quality-gate.py --url <url>        # Validate API endpoint
    python3 output-quality-gate.py --email <msg-id>   # Validate email was sent
    python3 output-quality-gate.py --file <path>      # Validate any text file
"""

import argparse
import sys
import json
from pathlib import Path


def validate_docx(path: str) -> tuple[bool, list[str]]:
    """Validate .docx report has real content."""
    errors = []
    try:
        from docx import Document
    except ImportError:
        errors.append("python-docx not installed — run: pip install python-docx")
        return False, errors

    if not Path(path).exists():
        errors.append(f"File not found: {path}")
        return False, errors

    doc = Document(path)
    text = " ".join(p.text for p in doc.paragraphs)
    char_count = len(text.strip())

    # Gate 1: Minimum content length
    if char_count < 500:
        errors.append(f"Report too short: {char_count} chars (minimum 500)")

    # Gate 2: Placeholder detection
    placeholders = [
        "not available",
        "not found",
        "no data",
        "N/A",
        "placeholder",
        "TODO",
        "TBD",
    ]
    for ph in placeholders:
        count = text.lower().count(ph.lower())
        if count > 3:
            errors.append(f"Too many '{ph}' occurrences: {count} (max 3)")

    # Gate 3: Section headers present (basic structure check)
    headings = [
        p.text for p in doc.paragraphs if p.style and "Heading" in str(p.style.name)
    ]
    if len(headings) < 2:
        errors.append(f"Report has only {len(headings)} headings (expected at least 2)")

    if errors:
        return False, errors
    return True, [
        f"PASS: {char_count} chars, {len(headings)} sections, placeholders OK"
    ]


def validate_url(url: str) -> tuple[bool, list[str]]:
    """Validate API endpoint returns real content."""
    errors = []
    try:
        import urllib.request

        req = urllib.request.Request(url, headers={"User-Agent": "quality-gate/1.0"})
        with urllib.request.urlopen(req, timeout=30) as resp:
            status = resp.status
            body = resp.read().decode("utf-8", errors="replace")

        if status != 200:
            errors.append(f"HTTP {status} (expected 200)")

        if len(body.strip()) < 10:
            errors.append(f"Response body too short: {len(body)} chars")

        # Try JSON parsing
        try:
            data = json.loads(body)
            if isinstance(data, dict):
                if data.get("error"):
                    errors.append(f"Response contains error: {data['error']}")
                if data.get("status") == "error":
                    errors.append("Response status is 'error'")
        except json.JSONDecodeError:
            pass  # Not JSON, that's OK

    except Exception as e:
        errors.append(f"Request failed: {e}")

    if errors:
        return False, errors
    return True, [f"PASS: HTTP {status}, {len(body)} chars"]


def validate_file(path: str) -> tuple[bool, list[str]]:
    """Validate any text file has real content."""
    errors = []

    if not Path(path).exists():
        errors.append(f"File not found: {path}")
        return False, errors

    text = Path(path).read_text(errors="replace")
    char_count = len(text.strip())

    if char_count < 100:
        errors.append(f"File too short: {char_count} chars (minimum 100)")

    placeholders = ["not available", "not found", "placeholder", "TODO", "TBD"]
    for ph in placeholders:
        count = text.lower().count(ph.lower())
        if count > 5:
            errors.append(f"Too many '{ph}' occurrences: {count} (max 5)")

    if errors:
        return False, errors
    return True, [f"PASS: {char_count} chars, placeholders OK"]


def validate_email(msg_id: str) -> tuple[bool, list[str]]:
    """Validate email was sent by checking for Gmail message ID."""
    errors = []
    if not msg_id or msg_id.strip() == "":
        errors.append("No Gmail message ID provided")
        return False, errors

    if len(msg_id) < 10:
        errors.append(f"Message ID looks invalid: '{msg_id}' (too short)")
        return False, errors

    return True, [f"PASS: Gmail message ID present ({msg_id[:20]}...)"]


def main():
    parser = argparse.ArgumentParser(description="Output Quality Gate")
    parser.add_argument("--docx", help="Validate .docx report")
    parser.add_argument("--url", help="Validate API endpoint")
    parser.add_argument("--email", help="Validate email message ID")
    parser.add_argument("--file", help="Validate any text file")
    parser.add_argument("--json", action="store_true", help="Output as JSON")
    args = parser.parse_args()

    if not any([args.docx, args.url, args.email, args.file]):
        parser.print_help()
        sys.exit(1)

    results = []

    if args.docx:
        passed, msgs = validate_docx(args.docx)
        results.append(
            {"type": "docx", "path": args.docx, "passed": passed, "messages": msgs}
        )

    if args.url:
        passed, msgs = validate_url(args.url)
        results.append(
            {"type": "url", "url": args.url, "passed": passed, "messages": msgs}
        )

    if args.email:
        passed, msgs = validate_email(args.email)
        results.append(
            {"type": "email", "msg_id": args.email, "passed": passed, "messages": msgs}
        )

    if args.file:
        passed, msgs = validate_file(args.file)
        results.append(
            {"type": "file", "path": args.file, "passed": passed, "messages": msgs}
        )

    all_passed = all(r["passed"] for r in results)

    if args.json:
        print(json.dumps({"passed": all_passed, "results": results}, indent=2))
    else:
        for r in results:
            status = "PASS" if r["passed"] else "FAIL"
            print(
                f"[{status}] {r['type']}: {r.get('path') or r.get('url') or r.get('msg_id')}"
            )
            for msg in r["messages"]:
                print(f"  {msg}")

    sys.exit(0 if all_passed else 1)


if __name__ == "__main__":
    main()
