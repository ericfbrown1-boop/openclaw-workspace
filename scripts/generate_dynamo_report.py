"""Generate Project Dynamo M&A Analysis Word Document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)


def add_h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 63, 128)


def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 100, 160)


def tbl(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    doc.add_paragraph()


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.runs[0].font.size = Pt(10)


def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10)


# ── CONFIDENTIAL HEADER ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CONFIDENTIAL - PROJECT DYNAMO")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(180, 0, 0)

# ── TITLE ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PROJECT DYNAMO")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0, 63, 128)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Cohesity Acquisition of Dell PowerProtect Data Domain")
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(80, 80, 80)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run(
    f"Strategic M&A Analysis & Skills Framework | {datetime.now().strftime('%B %d, %Y')}"
)
run3.font.size = Pt(11)
run3.font.color.rgb = RGBColor(120, 120, 120)
run3.italic = True

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run(
    "Prepared by Jarvis AI | For: Eric Brown, CFO & COO, Cohesity"
)
run4.font.size = Pt(9)
run4.font.color.rgb = RGBColor(150, 150, 150)
run4.italic = True

doc.add_page_break()

# ── 1. EXECUTIVE SUMMARY ──
add_h1("1. Executive Summary")
body(
    "Project Dynamo is the potential acquisition by Cohesity of Dell Technologies' PowerProtect "
    "Data Domain business unit via a carve-out from Dell's Infrastructure Solutions Group (ISG). "
    "Data Domain is the world's leading purpose-built backup appliance platform, holding the #1 "
    "position in the Gartner Magic Quadrant for 20 consecutive years with an estimated installed "
    "base of 60,000+ enterprise systems globally."
)
body(
    "At its core, Project Dynamo represents a bet on consolidation: Cohesity (post-Veritas) "
    "would become the undisputed leader in enterprise data protection with a combined market share "
    "potentially exceeding 30%. The financial case hinges on (1) the ability to retain the Data "
    "Domain installed base, (2) successful carve-out execution from Dell's deeply integrated "
    "organization, and (3) Cohesity's capacity to absorb a second major acquisition while still "
    "integrating Veritas."
)

add_h2("Preliminary Deal Indicators")
tbl(
    ["Dimension", "Assessment", "Signal"],
    [
        ("Strategic Fit", "High - eliminates #1 competitor, massive installed base", "Favorable"),
        ("Financial Accretion", "Conditional - depends on price and retention rate", "Needs Model"),
        ("Carve-Out Complexity", "HIGH - deeply embedded in Dell sales, IT, channel", "Risk"),
        ("Integration Capacity", "Risky - Veritas integration only 4 months old", "Risk"),
        ("Regulatory Risk", "Elevated - 30%+ combined market share; HSR/EU review", "Monitor"),
        ("Seller Motivation", "High - Dell pivoting to AI infrastructure; DD non-core", "Favorable"),
        ("Timing", "Opportunistic - but Cohesity bandwidth is stretched", "Caution"),
    ],
)

# ── 2. TARGET PROFILE ──
add_h1("2. Target Profile: Dell PowerProtect Data Domain")

add_h2("Product Overview")
tbl(
    ["Attribute", "Detail"],
    [
        ("Product Family", "PowerProtect Data Domain (DD) - purpose-built backup appliances"),
        ("Heritage", "Acquired by EMC (2009, $2.1B) then Dell via Dell-EMC merger (2016)"),
        ("Market Position", "Gartner Magic Quadrant Leader - 20 consecutive years (through 2025)"),
        ("Installed Base", "Estimated 60,000+ active systems globally, Fortune 500 heavy"),
        ("Key Differentiator", "Inline deduplication (up to 65:1), immutable snapshots, DD Boost"),
        ("Form Factors", "Physical appliances (DD3300-DD9900), virtual (DDVE), cloud (DDMC)"),
        ("Cloud Integration", "Native tiering to AWS, Azure, GCP; WORM/compliance retention lock"),
        ("Refresh Cycle", "5-7 years - long, sticky customer relationships"),
        ("Primary Customers", "Financial services, healthcare, government, telco, large enterprise"),
    ],
)

add_h2("Estimated Financials (Dell ISG-Derived)")
body(
    "Note: Dell does not break out Data Domain separately. These are derived estimates "
    "based on Dell ISG storage sub-segment disclosures, IDC market share data, and "
    "industry analyst estimates."
)
tbl(
    ["Metric", "Estimate", "Basis"],
    [
        ("Total Revenue", "$2.0-3.0B", "~15-18% of Dell ISG storage; IDC PBBA tracker"),
        ("Recurring Revenue", "40-50% of total", "Industry standard for mature appliance businesses"),
        ("Gross Margin (blended)", "55-65%", "Hardware + software + high-margin support mix"),
        ("EBITDA Margin (standalone)", "25-35%", "Pre-corporate allocation; carve-out adjustments TBD"),
        ("Revenue Growth", "-2% to +1%", "Mature PBBA market; all-flash refresh partially offsetting"),
        ("PBBA Market Share", "~55-65% units", "IDC Worldwide Quarterly PBBA Tracker"),
    ],
)

# ── 3. STRATEGIC RATIONALE ──
add_h1("3. Strategic Rationale Assessment")

add_h2("3.1 Market Consolidation")
tbl(
    ["Entity", "Est. Market Share", "Combined"],
    [
        ("Cohesity (pre-Veritas)", "~5%", "-"),
        ("Veritas (acquired Dec 2025)", "~14%", "~19%"),
        ("Dell Data Domain (Project Dynamo)", "~12%", "~31%"),
        ("Rubrik", "~8%", "-"),
        ("Commvault", "~7%", "-"),
        ("Veeam", "~10%", "-"),
    ],
)
body(
    "A 31%+ combined market share would create a dominant #1 position with a 3x lead over "
    "the next largest single competitor. This level of consolidation in a $30B+ data protection "
    "market creates a formidable defensive moat and a compelling IPO/valuation narrative."
)

add_h2("3.2 Product Portfolio Complementarity")
tbl(
    ["Product", "Architecture", "Customers", "Revenue Model"],
    [
        ("Cohesity DataProtect", "Cloud-native, software-defined", "Enterprise, cloud-first", "Subscription + cloud"),
        ("Veritas NetBackup", "Legacy on-prem, broad OS/DB", "Large F500, complex envs", "Perpetual + maintenance"),
        ("Data Domain", "Purpose-built HW appliance, dedup IP", "F500, compliance-heavy", "Hardware + multi-yr support"),
    ],
)
body(
    "The three portfolios are largely complementary rather than cannibalistic: each serves "
    "different infrastructure philosophies. The strategic vision would be to gradually migrate "
    "DD customers to Cohesity's platform over a 5-7 year horizon while capturing near-term "
    "maintenance renewal revenue from the installed base."
)

add_h2("3.3 Dell's Motivation to Sell")
for b in [
    "AI infrastructure pivot: Dell ISG increasingly dominated by AI servers; data protection is low-growth vs. $10B+ AI server backlog",
    "Portfolio simplification: Dell has precedent for major divestitures (VMware spinoff 2021, $61B)",
    "Revenue declining: Dell ISG storage revenue declined ~5% recently, driven partly by data protection softness",
    "Capital allocation: Proceeds from a $4-8B deal fund AI infrastructure investment",
    "Partnership retention: Dell could retain a reseller/OEM relationship with Cohesity post-close",
]:
    bullet(b)

# ── 4. VALUATION ──
add_h1("4. Valuation Framework")

add_h2("4.1 Comparable Transactions")
tbl(
    ["Transaction", "Year", "Deal Value", "Rev Multiple", "EBITDA Multiple", "Notes"],
    [
        ("Cohesity / Veritas", "2024-25", "~$7B implied", "~2.5x", "~12x", "Most comparable - same market"),
        ("Thoma Bravo / Veritas", "2023", "$1.5B", "~1.2x", "~8x", "Distressed; lower multiple"),
        ("Broadcom / Symantec", "2019", "$10.7B", "~4x", "~16x", "Security premium; not direct comp"),
        ("OpenText / Micro Focus", "2023", "$5.8B", "~2.0x", "~9x", "Mature software; carve-out"),
        ("Insight / Veeam", "2019", "$5.0B", "~8x ARR", "~25x", "High-growth SaaS; premium"),
    ],
)

add_h2("4.2 Implied Valuation Range")
tbl(
    ["Scenario", "Revenue", "Multiple", "Implied EV", "Rationale"],
    [
        ("Bear Case", "$2.0B", "2.0x", "$4.0B", "Revenue declining; integration risk"),
        ("Base Case", "$2.5B", "2.5x", "$6.25B", "Stable recurring; partial synergies"),
        ("Bull Case", "$3.0B", "3.0x", "$9.0B", "Strong base; full synergies; Dell eager"),
        ("EV/EBITDA Bear", "$600M", "8x", "$4.8B", "Low end; execution risk"),
        ("EV/EBITDA Base", "$750M", "10x", "$7.5B", "Mid; standalone + partial synergies"),
        ("EV/EBITDA Bull", "$900M", "12x", "$10.8B", "Full synergies; scarcity premium"),
    ],
)
body("Working assumption: $6.0-6.5B purchase price (Base Case). Sensitivity +/-$1B.")

# ── 5. ACCRETION / DILUTION ──
add_h1("5. Accretion / Dilution Analysis Framework")

add_h2("5.1 Pro Forma Model Structure")
body(
    "The following uses estimated inputs. Run with Cohesity actual internals for real model."
)
tbl(
    ["Input", "Cohesity Standalone", "Data Domain", "Pro Forma Combined"],
    [
        ("Revenue", "~$2.5B", "~$2.5B", "~$5.0B"),
        ("Gross Margin", "~70%", "~60%", "~65%"),
        ("EBITDA", "~$600M (24%)", "~$700M (28%)", "~$1.3B + synergies"),
        ("EBITDA Margin", "~24%", "~28%", "~26% (pre-synergies)"),
    ],
)

add_h2("5.2 Synergy Assumptions")
tbl(
    ["Synergy Category", "Year 1", "Year 2", "Year 3", "Notes"],
    [
        ("Cost - G&A elimination", "$50M", "$100M", "$150M", "Duplicate corporate overhead"),
        ("Cost - R&D consolidation", "$75M", "$150M", "$200M", "Overlapping engineering"),
        ("Cost - Sales/mktg overlap", "$25M", "$75M", "$100M", "Field sales rationalization"),
        ("Revenue - Cross-sell", "$25M", "$75M", "$150M", "Cohesity upsell to DD customers"),
        ("Revenue - Migration svc", "$10M", "$30M", "$75M", "DD to Cohesity migration"),
        ("Total Synergies", "$185M", "$430M", "$675M", "Run-rate ~$700M+ by Y3"),
    ],
)

add_h2("5.3 Integration & Carve-Out Costs (One-Time)")
tbl(
    ["Cost Category", "Estimate", "Timing"],
    [
        ("IT separation & migration", "$100-200M", "Year 1-2"),
        ("HR restructuring & retention", "$75-150M", "Year 1"),
        ("Facilities rationalization", "$25-50M", "Year 1-3"),
        ("TSA costs (Dell, 18 months)", "$75-150M", "Year 1-2"),
        ("Rebranding & channel", "$25-50M", "Year 1"),
        ("Legal, banking, advisory", "$50-100M", "At close"),
        ("Total One-Time Costs", "$350-700M", "Primarily Y1-Y2"),
    ],
)

add_h2("5.4 Accretion Assessment (Preliminary)")
body("At $6.25B purchase price, funded 50% debt / 50% equity:")
tbl(
    ["Metric", "Year 1", "Year 2", "Year 3"],
    [
        ("Incremental EBITDA from DD", "$700M", "$730M", "$760M"),
        ("Less: integration costs", "($400M)", "($200M)", "($100M)"),
        ("Plus: synergies", "$185M", "$430M", "$675M"),
        ("Less: interest expense (7%)", "($219M)", "($219M)", "($219M)"),
        ("Less: intangibles amort.", "($150M)", "($150M)", "($150M)"),
        ("Net EBITDA contribution", "~$116M", "~$591M", "~$966M"),
        ("Accretion / (Dilution)", "DILUTIVE Y1", "ACCRETIVE Y2", "STRONGLY ACCRETIVE Y3"),
    ],
)
body(
    "PRELIMINARY CONCLUSION: Deal is likely DILUTIVE in Year 1 due to one-time integration "
    "and carve-out costs. Becomes ACCRETIVE in Year 2 as synergies ramp. Full accretion "
    "in Year 3. This is typical for carve-out acquisitions of this complexity."
)

# ── 6. RISK MATRIX ──
add_h1("6. Carve-Out Risk Assessment")

add_h2("Risk Matrix")
tbl(
    ["Risk", "Likelihood", "Impact", "Rating", "Mitigation"],
    [
        ("Stranded costs exceed estimate", "Medium", "High", "HIGH", "Detailed shared services audit; conservative model"),
        ("TSA failures from Dell", "Medium", "High", "HIGH", "SLA with penalties; parallel standalone build"),
        ("Customer churn in transition", "High", "High", "HIGH", "Retention programs; contractual lock-ins; early comms"),
        ("Key employee defection", "Medium", "Medium", "MEDIUM", "Retention packages; clear career path"),
        ("IP separation challenges", "Low", "High", "MEDIUM", "Clean IP audit pre-sign; escrow disputed IP"),
        ("Channel partner disruption", "Medium", "Medium", "MEDIUM", "Partner incentive program; Dell co-sell agreement"),
        ("Antitrust/regulatory delay", "High", "Medium", "MEDIUM", "Early HSR; prepare remedies; 12-18mo close"),
        ("Integration overload (Veritas+DD)", "High", "High", "CRITICAL", "Dedicated PMO; phased approach; 6mo delay option"),
        ("Revenue decline post-close", "Medium", "High", "HIGH", "Conservative case; invest in DD roadmap immediately"),
    ],
)

# ── 7. GO/NO-GO ──
add_h1("7. Go / No-Go Decision Framework")

add_h2("Weighted Scorecard")
tbl(
    ["Criterion", "Weight", "Score (1-5)", "Weighted", "Comments"],
    [
        ("Strategic fit", "20%", "4.5", "0.90", "Dominant market position; perfect fit"),
        ("Financial accretion (Y2+)", "20%", "3.5", "0.70", "Accretive Y2+ at base case"),
        ("Customer base quality", "15%", "3.0", "0.45", "High-quality F500; churn risk real"),
        ("Technology / IP value", "10%", "4.0", "0.40", "Dedup IP, DD Boost, 20yr MQ leader"),
        ("Carve-out complexity", "15%", "2.0", "0.30", "Deeply embedded; HIGH risk"),
        ("Regulatory risk", "10%", "2.5", "0.25", "30%+ share; DOJ/EU scrutiny"),
        ("Integration capacity", "10%", "1.5", "0.15", "Veritas only 4mo old; bandwidth"),
        ("TOTAL", "100%", "-", "3.15", "CONDITIONAL PROCEED"),
    ],
)
body(
    "PRELIMINARY VERDICT: CONDITIONAL PROCEED (Score 3.15/5.0). Strong strategic rationale "
    "partially offset by carve-out complexity and integration bandwidth. Recommend formal DD "
    "process starting no earlier than Q2 2026 to allow Veritas stabilization."
)

# ── 8. NEXT STEPS ──
add_h1("8. Recommended Next Steps")

steps = [
    ("Immediate (0-30 days)", [
        "Engage M&A counsel and investment bankers for NDA/exploration",
        "Initiate preliminary conversations with Dell Corporate Development",
        "Commission independent market sizing of Data Domain standalone revenue",
        "Internal bandwidth assessment: when can integration team absorb DD?",
    ]),
    ("Short-term (30-90 days)", [
        "Execute NDA; obtain Dell's confidential information memorandum (CIM)",
        "Build detailed standalone financial model using DD-specific data",
        "Run accretion/dilution model with Cohesity actual internal financials",
        "Antitrust counsel pre-assessment: model likely remedies",
        "Begin TSA scope planning based on Dell shared services inventory",
    ]),
    ("Due Diligence (90-180 days)", [
        "Full financial, legal, IP, and operational due diligence",
        "Customer interviews (top 50 DD accounts) - retention intent",
        "Key employee retention planning and org design",
        "Negotiate TSA schedules and stranded cost allocations",
        "File HSR (Hart-Scott-Rodino) pre-merger notification",
    ]),
]

for phase, items in steps:
    add_h2(phase)
    for item in items:
        bullet(item)
    doc.add_paragraph()

# ── APPENDIX: SKILL.MD ──
add_h1("Appendix: Project Dynamo SKILL.md")
body(
    "The SKILL.md file is installed at ~/.openclaw/workspace/skills/project-dynamo/SKILL.md. "
    "It enables Jarvis to execute structured Project Dynamo analysis on demand."
)
doc.add_paragraph()

skill_path = os.path.expanduser(
    "~/.openclaw/workspace/skills/project-dynamo/SKILL.md"
)
with open(skill_path) as f:
    skill_content = f.read()
p = doc.add_paragraph()
run = p.add_run(skill_content)
run.font.size = Pt(8)
run.font.name = "Courier New"
run.font.color.rgb = RGBColor(50, 50, 50)

# ── FOOTER ──
doc.add_paragraph()
p_disc = doc.add_paragraph(
    "CONFIDENTIAL - Project Dynamo. Prepared for Eric Brown, CFO & COO, Cohesity. "
    "Do not distribute. Financial estimates based on public information and analyst "
    "estimates; due diligence required to validate. Not investment advice.\n"
    "Sent by Jarvis - AI assistant to Eric Brown"
)
for run in p_disc.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.italic = True

os.makedirs(os.path.expanduser("~/Documents/reports"), exist_ok=True)
outpath = os.path.expanduser(
    "~/Documents/reports/ProjectDynamo_MA_Analysis_20260404.docx"
)
doc.save(outpath)
print(f"Saved: {outpath}")
