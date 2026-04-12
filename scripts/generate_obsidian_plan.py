#!/usr/bin/env python3
"""Generate the Obsidian Second-Brain + AutoResearch integration plan as .docx.

Deliverable: ~/openclaw-workspace/outputs/obsidian-second-brain-plan.docx

Atomic write: builds in memory, saves to .tmp, then os.replace (Auditor patch #5).
"""
from __future__ import annotations

import hashlib
import os
import shutil
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = Path.home() / "openclaw-workspace" / "outputs"
OUTPUT_PATH = OUTPUT_DIR / "obsidian-second-brain-plan.docx"


def build_document() -> Document:
    """Build and return the complete plan document."""
    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ==================== TITLE PAGE ====================
    title = doc.add_heading("Obsidian Second-Brain + AutoResearch Integration Plan", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Prepared for Eric Brown — Jarvis AI Assistant Infrastructure")
    doc.add_paragraph("Task ID: obsidian-second-brain-plan")
    doc.add_paragraph("Date: April 2026")
    doc.add_paragraph(
        "This document details a four-phase plan to integrate Obsidian as Jarvis's "
        "persistent knowledge layer on Eric's MacBook Pro, incorporating an adapted "
        "version of karpathy/autoresearch's self-improving agent instructions pattern."
    )

    # ==================== 1. EXECUTIVE SUMMARY ====================
    doc.add_heading("1. Executive Summary", level=1)

    doc.add_paragraph(
        "This plan describes how to transform Obsidian into Jarvis's second brain — "
        "a structured markdown vault that Jarvis reads from for context injection into "
        "pipeline stages, writes to for daily briefings and research outputs, and uses "
        "as the substrate for cross-run learning. The integration replaces the current "
        "flat memory file tree (~/.claude/projects/*/memory/*.md) with a richer, "
        "queryable, link-graph-enabled knowledge base while maintaining backward "
        "compatibility with the existing Librarian memory system."
    )

    doc.add_paragraph(
        "The plan draws inspiration from karpathy/autoresearch, an autonomous ML "
        "experimentation framework with 42k GitHub stars. While autoresearch itself "
        "is designed for GPU training loops (not personal knowledge management), its "
        "core design pattern — program.md as self-modifying agent instructions that "
        "improve through measured iteration — is directly applicable to Jarvis. We "
        "adapt this pattern as a Phase 4 meta-capability where Jarvis proposes "
        "improvements to its own operational instructions based on measured outcomes, "
        "subject to Eric's review and approval."
    )

    doc.add_paragraph(
        "Estimated total effort: approximately 12 hours of agent coding time across "
        "four phases over two weeks, plus Eric's review checkpoints at each phase gate."
    )

    # ==================== 2. THE AUTORESEARCH DISCREPANCY ====================
    doc.add_heading("2. The karpathy/autoresearch Discrepancy and Resolution", level=1)

    doc.add_heading("2.1 What autoresearch Actually Is", level=2)
    doc.add_paragraph(
        "The karpathy/autoresearch project (github.com/karpathy/autoresearch) is an "
        "autonomous ML experimentation framework. Per its README: 'Give an AI agent a "
        "small but real LLM training setup and let it experiment autonomously overnight. "
        "It modifies the code, trains for 5 minutes, checks if the result improved, "
        "keeps or discards, and repeats.' It consists of three core files: prepare.py "
        "(data preparation), train.py (the code the agent modifies), and program.md "
        "(the human programs the agent via this markdown file). There is no knowledge "
        "graph, no note-taking, and no PKM functionality."
    )

    doc.add_heading("2.2 Why It Matters for Jarvis", level=2)
    doc.add_paragraph(
        "The valuable insight from autoresearch is not GPU training — it is the "
        "iteration loop: the agent reads instructions from program.md, executes a "
        "task, measures the result, keeps or discards based on improvement, and "
        "proposes modifications to the instructions. This maps directly to Jarvis: "
        "a jarvis-program.md file that Jarvis reads at daily briefing time, acts on, "
        "logs results to the Obsidian vault, and proposes improvements to its own "
        "operational instructions for Eric's review."
    )

    doc.add_heading("2.3 Recommended Approach: Option 2 — Obsidian + Adapted Pattern", level=2)
    doc.add_paragraph(
        "We recommend building the Obsidian second-brain (Phases 1-3) and then "
        "adapting autoresearch's program.md pattern as a Jarvis meta-capability "
        "(Phase 4). This delivers both the PKM value Eric wants and the self-improving "
        "agent pattern that makes autoresearch interesting, while being honest about "
        "what autoresearch actually is versus what it is not."
    )

    # ==================== 3. ARCHITECTURE ====================
    doc.add_heading("3. Architecture Overview", level=1)

    doc.add_heading("3.1 Vault Location and Access", level=2)
    doc.add_paragraph(
        "The Obsidian vault will be located at ~/Documents/jarvis-vault/ on Eric's "
        "MacBook Pro. This is a standard macOS Documents location that is iCloud-eligible "
        "but should have iCloud Drive sync disabled for the vault folder to prevent "
        "sync conflicts during automated writes (set via xattr or Finder). Programmatic "
        "access is through obsidian-cli, already installed at /opt/homebrew/bin/obsidian-cli. "
        "All cron-invoked commands must use the absolute path /opt/homebrew/bin/obsidian-cli "
        "to avoid the launchd PATH bug documented in jarvis_launchd_path_bug.md."
    )

    doc.add_heading("3.2 Vault Structure (PARA Method)", level=2)
    doc.add_paragraph(
        "The vault follows the PARA organizational method with additional Jarvis-specific "
        "directories. Top-level folders: Projects/ (active time-bound work, auto-created "
        "per pipeline task), Areas/ (ongoing domains like Jarvis-Operations, PowerSpec, "
        "Financial-Reports), Resources/ (reference material including Incidents/ and "
        "Research/), Archive/ (completed or inactive items), Daily/ (daily briefing "
        "notes in YYYY-MM-DD.md format), Templates/ (Templater templates for structured "
        "note creation), and Meta/ (self-improvement infrastructure including "
        "jarvis-program.md and proposed-amendments/)."
    )

    doc.add_heading("3.3 Data Flow", level=2)
    doc.add_paragraph(
        "Data flows bidirectionally between Jarvis and the vault. Write flows: the "
        "daily briefing cron (6am PT) creates Daily/YYYY-MM-DD.md notes via obsidian-cli "
        "create; the librarian_weekly.py script writes incident summaries to "
        "Resources/Incidents/; the pipeline post-conductor hook copies outputs to "
        "Projects/<task-id>/. Read flows: external_context_hook.py queries the vault "
        "via obsidian-cli search-content and injects relevant notes into pipeline "
        "Researcher and Planner stages; the daily cron reads Meta/jarvis-program.md "
        "for operational instructions."
    )

    doc.add_heading("3.4 Integration with Existing Systems", level=2)
    doc.add_paragraph(
        "The vault supplements but does not replace the existing Librarian memory "
        "system (~/.claude/projects/*/memory/*.md). Both systems coexist: Librarian "
        "memory files continue to be injected into pipeline stages via "
        "load_librarian_memory(), while vault content is injected via the "
        "external_context_hook.py extension point. This maintains backward compatibility "
        "and allows a gradual migration where proven vault patterns eventually supersede "
        "flat memory files."
    )

    # ==================== 4. PIPELINE INTEGRATION ====================
    doc.add_heading("4. Pipeline Integration via external_context_hook.py", level=1)

    doc.add_paragraph(
        "The jarvis_pipeline.py orchestrator already includes an external_context_hook.py "
        "extension point designed for exactly this purpose. The hook's fetch_context(task, "
        "stage) function is called during Researcher and Planner stages. For vault "
        "integration, the implementation will: extract search terms from the task "
        "objective, query the vault via obsidian-cli search-content, read the top 5 "
        "most relevant notes via obsidian-cli print, concatenate results up to the "
        "400KB soft cap (within the 500KB EXTERNAL_CONTEXT_MAX_BYTES limit), and "
        "return the formatted markdown."
    )

    doc.add_paragraph(
        "Error handling follows the existing fail-safe pattern documented in "
        "jarvis_pipeline_runner.md: any exception in the hook is logged to "
        "~/openclaw-workspace/logs/jarvis-pipeline-hooks.log, and the pipeline "
        "continues with empty context. The hook respects the task-level opt-out "
        "(externalContext.enabled = false) for trivial tasks that do not benefit "
        "from vault context injection."
    )

    doc.add_paragraph(
        "The hook will use absolute paths for all CLI invocations "
        "(/opt/homebrew/bin/obsidian-cli) per the launchd PATH bug lesson, even though "
        "the hook runs inside the pipeline process rather than directly from launchd. "
        "Defense-in-depth costs nothing and prevents future regressions if the execution "
        "context changes."
    )

    # ==================== 5. OBSIDIAN PLUGIN RECOMMENDATIONS ====================
    doc.add_heading("5. Obsidian Plugin Recommendations", level=1)

    doc.add_heading("5.1 Core Plugins (Phase 1)", level=2)
    doc.add_paragraph(
        "Dataview: the de facto standard for querying an Obsidian vault as a database. "
        "Enables TABLE, LIST, and TASK queries in notes. Essential for the Jarvis "
        "Operations dashboard that lists recent daily notes, open tasks, and incidents. "
        "Over 2 million downloads; actively maintained."
    )
    doc.add_paragraph(
        "Templater: structured note templates with dynamic variables (date, title, "
        "frontmatter). More powerful than Obsidian's core Templates plugin. Supports "
        "tp.date.now, tp.file.title, and custom JavaScript functions. Used by the "
        "daily briefing cron to create notes with consistent structure."
    )
    doc.add_paragraph(
        "Obsidian Git: automatic version control for the vault. Configurable auto-commit "
        "interval (recommended: 5 minutes), auto-push to GitHub for off-machine backup, "
        "and merge-based sync to handle concurrent edits. Provides rollback capability "
        "if Jarvis writes corrupt a note."
    )

    doc.add_heading("5.2 Enhancement Plugins (Phase 2-3)", level=2)
    doc.add_paragraph(
        "Excalidraw: visual diagrams, architecture sketches, and mind maps embedded "
        "directly in notes. Useful for pipeline architecture diagrams and system "
        "overviews that Jarvis can reference. Calendar: a calendar view of daily "
        "notes that pairs with Templater for auto-creation of day pages. Smart "
        "Connections or a Claude-Obsidian plugin: AI-powered semantic search and "
        "auto-tagging for notes, enabling richer context injection beyond keyword "
        "matching."
    )

    # ==================== 6. DAILY BRIEFING VAULT WRITES ====================
    doc.add_heading("6. Daily Briefing Vault Integration", level=1)

    doc.add_paragraph(
        "The existing daily briefing cron job (6am PT) will be extended to write a "
        "structured daily note to the vault at Daily/YYYY-MM-DD.md. The note includes "
        "sections for: calendar events (next 24 hours via gog calendar list, with "
        "graceful failure handling per gog_oauth_scopes.md), pipeline status (via "
        "jarvis_pipeline.py list), open tasks (from tasks.json), incidents since "
        "yesterday (from incidents.jsonl), weather summary (via wttr.in), and a "
        "free-form notes section for Eric's manual additions."
    )

    doc.add_paragraph(
        "The cron script will read Meta/jarvis-program.md at startup and follow its "
        "embedded checklist instructions. This is the read path of the autoresearch-"
        "inspired self-improvement loop: the program file tells Jarvis what to do each "
        "morning, and Jarvis logs outcomes to the vault for later analysis."
    )

    doc.add_paragraph(
        "All vault write commands use the absolute path /opt/homebrew/bin/obsidian-cli "
        "to prevent the launchd PATH bug. The daily briefing template uses Templater "
        "syntax for dates and dynamic content insertion."
    )

    # ==================== 7. AUTORESEARCH SELF-IMPROVEMENT LOOP ====================
    doc.add_heading("7. AutoResearch-Inspired Self-Improvement Loop (Phase 4)", level=1)

    doc.add_heading("7.1 The program.md Pattern", level=2)
    doc.add_paragraph(
        "In karpathy/autoresearch, program.md is the file where the human programs "
        "the agent. The agent reads it, executes the instructions, measures results, "
        "and proposes modifications. We adapt this as jarvis-vault/Meta/jarvis-program.md "
        "— a markdown file containing Jarvis's operational instructions that evolve "
        "over time through measured iteration."
    )

    doc.add_heading("7.2 The Iteration Cycle", level=2)
    doc.add_paragraph(
        "Monday through Thursday: the 6am cron reads jarvis-program.md, executes the "
        "daily checklist (calendar, pipeline status, incidents, tasks), and logs metrics "
        "to Meta/metrics.md. Friday: the cron additionally runs a weekly self-review — "
        "counting tasks completed, incidents logged, and pipeline halts for the week; "
        "comparing to last week's metrics; identifying what went well, what failed, and "
        "what was slow; and generating 1-3 proposed amendments to jarvis-program.md in "
        "Meta/proposed-amendments/YYYY-MM-DD.md. Each proposed amendment must cite "
        "measured evidence, not speculation."
    )

    doc.add_heading("7.3 Human-in-the-Loop Safety", level=2)
    doc.add_paragraph(
        "The critical safety guardrail: Jarvis NEVER edits jarvis-program.md directly. "
        "All proposed changes are written to Meta/proposed-amendments/ and require "
        "Eric's manual review and approval before being applied. This is the key "
        "difference from autoresearch's fully autonomous loop — Karpathy's use case "
        "is overnight GPU training where full autonomy makes sense; Jarvis manages "
        "Eric's real infrastructure where human oversight is non-negotiable. The cron "
        "script includes an explicit guard: if the write target is jarvis-program.md, "
        "abort immediately."
    )

    doc.add_heading("7.4 Metrics and Measurement", level=2)
    doc.add_paragraph(
        "Meta/metrics.md maintains a structured log of weekly measurements: tasks "
        "completed, incidents logged, pipeline halts, vault notes created, vault size "
        "in megabytes, and external context injection rates. This data drives the "
        "self-improvement proposals — without measurement, there is no basis for "
        "claiming improvement. The metrics table is appended weekly by the Friday "
        "self-review step and referenced in all proposed amendments."
    )

    # ==================== 8. PHASED ROLLOUT ====================
    doc.add_heading("8. Phased Rollout Plan", level=1)

    doc.add_heading("Phase 0: Prerequisites (Day 1, ~30 minutes)", level=2)
    doc.add_paragraph(
        "Verify installations: Obsidian.app (confirmed at /Applications/Obsidian.app), "
        "obsidian-cli (confirmed at /opt/homebrew/bin/obsidian-cli), python-docx "
        "(confirmed v1.2.0). Create vault directory structure with mkdir -p commands "
        "for all PARA folders. Register vault with obsidian-cli via set-default. "
        "Eric opens Obsidian.app once to create .obsidian/ config. Enable Community "
        "Plugins in Obsidian settings (one-time manual step). Install core plugins: "
        "Dataview, Templater, Obsidian Git. Clone autoresearch for reference: "
        "git clone --depth 1 https://github.com/karpathy/autoresearch.git into "
        "~/openclaw-workspace/vendor/autoresearch/. Disable iCloud sync for the vault "
        "folder to prevent sync conflicts. Checkpoint: obsidian-cli list --vault "
        "jarvis-vault returns the directory structure."
    )

    doc.add_heading("Phase 1: Vault Population — Jarvis Writes (Days 2-3, ~4h)", level=2)
    doc.add_paragraph(
        "Create Templater templates in Templates/ for daily briefings, incident reports, "
        "research briefs, and project kickoffs. Write Meta/jarvis-program.md with initial "
        "operational instructions (daily checklist, weekly self-review protocol, current "
        "focus areas, and safety constraints). Modify the daily briefing cron script to "
        "write Daily/YYYY-MM-DD.md via obsidian-cli create using absolute paths. Modify "
        "librarian_weekly.py to also write Resources/Incidents/week-of-YYYY-MM-DD.md. "
        "Add post-conductor hook to jarvis_pipeline.py to copy pipeline outputs to vault. "
        "Run /simplify on all modified scripts per feedback_quality_audit.md. Checkpoint: "
        "after 2 days, obsidian-cli list shows at least 10 notes; Daily/ contains at "
        "least 2 daily briefing notes."
    )

    doc.add_heading("Phase 2: Vault Reading — Jarvis Queries (Days 4-5, ~3h)", level=2)
    doc.add_paragraph(
        "Implement external_context_hook.py with vault query logic: copy from .example "
        "template, add obsidian-cli search-content integration, respect the 500KB "
        "EXTERNAL_CONTEXT_MAX_BYTES cap, handle missing vault gracefully (return empty "
        "string, log warning). Add externalContext configuration to relevant tasks in "
        "tasks.json. Test with a pipeline run and verify externalContextBytes > 0 in "
        "pipeline-state.json. Run /simplify on the hook implementation. Checkpoint: "
        "pipeline stage prompts contain vault-sourced content visible in "
        "pipeline-outputs/."
    )

    doc.add_heading("Phase 3: Enhanced Features (Days 6-8, ~3h)", level=2)
    doc.add_paragraph(
        "Create Dataview dashboard note at Areas/Jarvis-Operations/dashboard.md with "
        "queries for recent daily notes, open tasks, and recent incidents. Set up "
        "Obsidian Git with 5-minute auto-commit interval and GitHub remote for "
        "off-machine backup. Optional: Excalidraw architecture diagram for the Jarvis "
        "pipeline. Optional: Dropbox weekly backup script. Optional: Apple Notes bridge "
        "via memo CLI export. Checkpoint: dashboard renders correctly in Obsidian with "
        "live Dataview queries; git log shows auto-commits."
    )

    doc.add_heading("Phase 4: Self-Improvement Loop (Days 9-12, ~5h)", level=2)
    doc.add_paragraph(
        "Implement Friday self-review logic in the daily briefing cron: read "
        "Meta/jarvis-program.md for review instructions, read Meta/metrics.md for "
        "current week's data, generate Meta/proposed-amendments/YYYY-MM-DD.md with "
        "proposals. Create Meta/metrics.md with the initial schema (weekly rows for "
        "tasks completed, incidents, pipeline halts, notes created, vault size). Add "
        "metrics collection to the daily cron (append row weekly). Implement the "
        "safety guardrail: explicit check that Jarvis never writes to "
        "jarvis-program.md directly. Run /simplify on all Phase 4 code. Checkpoint: "
        "after first Friday, proposed-amendments/ contains at least one file; "
        "metrics.md has at least one data row; jarvis-program.md is byte-identical "
        "to its initial content."
    )

    # ==================== 9. RISKS AND MITIGATIONS ====================
    doc.add_heading("9. Risks and Mitigations", level=1)

    risks = [
        (
            "launchd PATH resolution failure",
            "High", "High",
            "All cron-invoked commands use absolute paths (/opt/homebrew/bin/obsidian-cli). "
            "This is a known bug class documented in jarvis_launchd_path_bug.md where "
            "launchd's minimal PATH (/usr/bin:/bin:/usr/sbin:/sbin) excludes Homebrew "
            "binaries. Defense-in-depth: even scripts invoked via /bin/zsh -lc get "
            "explicit EnvironmentVariables.PATH in their plist."
        ),
        (
            "Unbounded vault growth",
            "Medium", "Medium",
            "Archive notes older than 90 days via a weekly cleanup script. Monitor vault "
            "size in Meta/metrics.md. Target: under 500 MB after 30 days."
        ),
        (
            "external_context_hook.py crashes pipeline",
            "Low", "Low",
            "Hook runs with full try/except per jarvis_pipeline_runner.md contract. All "
            "exceptions logged to jarvis-pipeline-hooks.log. Pipeline continues with "
            "empty context on any failure."
        ),
        (
            "Obsidian Git conflicts with concurrent writes",
            "Medium", "Medium",
            "Use atomic file writes (write to temp, rename). Obsidian Git's 5-minute "
            "auto-commit interval and merge-based sync handle most conflicts. Document "
            "manual resolution for edge cases."
        ),
        (
            "gog OAuth token expires during daily briefing",
            "Low", "Medium",
            "Daily briefing handles gog failures gracefully: skip calendar section, log "
            "incident to incidents.jsonl, continue with remaining sections. The OAuth "
            "scope declaration fix (gog_oauth_scopes.md) should prevent this once "
            "re-auth completes."
        ),
        (
            "Self-improvement loop proposes bad changes",
            "Medium", "Low",
            "All amendments go to proposed-amendments/ for Eric's review. Jarvis never "
            "edits jarvis-program.md directly. The human-in-the-loop is non-negotiable."
        ),
        (
            "iCloud sync conflicts with vault",
            "Medium", "Medium",
            "Disable iCloud Drive sync for ~/Documents/jarvis-vault/ during Phase 0 "
            "setup using xattr or Finder settings. The vault is local-first; cloud "
            "backup goes through Obsidian Git to GitHub and optionally Dropbox."
        ),
        (
            "Pipeline quality stage failures (recurring pattern)",
            "Medium", "Medium",
            "All pipeline tasks have verificationCmd. Hybrid Quality gate (deterministic "
            "plus LLM review) is enforced. Modified scripts go through /simplify plus "
            "external audit per feedback_quality_audit.md. Recurring pipeline pattern "
            "from librarian_weekly_patterns.md treated as HIGH severity."
        ),
    ]

    for name, likelihood, impact, mitigation in risks:
        doc.add_heading(f"Risk: {name}", level=2)
        doc.add_paragraph(f"Likelihood: {likelihood} | Impact: {impact}")
        doc.add_paragraph(f"Mitigation: {mitigation}")

    # ==================== 10. SUCCESS METRICS ====================
    doc.add_heading("10. Success Metrics", level=1)

    metrics = [
        ("Vault note count", "5+ notes/day week 1; 10+/day by week 4",
         "obsidian-cli list --vault jarvis-vault -r | wc -l"),
        ("Daily briefing reliability", "100% of mornings have Daily/YYYY-MM-DD.md by 6:15am",
         "Check Daily/ directory for gaps"),
        ("External context injection", "externalContextBytes > 0 for 80%+ of pipeline runs",
         "Query pipeline-state.json"),
        ("Eric vault usage", "3+ manual accesses/week within 30 days",
         "Eric self-reports (privacy-respecting)"),
        ("Self-improvement cycle", "1+ proposed-amendments/ file generated and reviewed",
         "Check Meta/proposed-amendments/ directory"),
        ("Pipeline quality improvement", "50%+ reduction in quality-stage halts",
         "Compare incidents.jsonl pipeline category: 2 weeks before vs after"),
        ("Vault size manageable", "Under 500 MB after 30 days",
         "du -sh ~/Documents/jarvis-vault/"),
        ("Dataview queries functional", "Dashboard renders 3+ live query results",
         "Manual check in Obsidian.app"),
    ]

    for name, target, method in metrics:
        doc.add_paragraph(f"{name}: Target — {target}. Measured via: {method}.")

    # ==================== 11. CONFIGURATION APPENDIX ====================
    doc.add_heading("11. Configuration and Implementation Appendix", level=1)

    doc.add_heading("11.1 Templater Daily Briefing Template", level=2)
    doc.add_paragraph(
        "The daily-briefing.md template uses Templater syntax for dynamic date "
        "insertion (tp.date.now) and structured sections: Calendar (Next 24h), "
        "Pipeline Status, Open Tasks, Incidents (Last 24h), Weather, and Notes. "
        "Each section is populated by the daily cron script via obsidian-cli create "
        "with the --content flag, passing the rendered template content."
    )

    doc.add_heading("11.2 Dataview Dashboard Queries", level=2)
    doc.add_paragraph(
        "The Areas/Jarvis-Operations/dashboard.md note contains three Dataview TABLE "
        "queries: recent daily briefings (from Daily/, sorted by date descending, "
        "limited to 7), open tasks (from Projects/, filtered by status != completed, "
        "sorted by priority), and recent incidents (from Resources/Incidents/, sorted "
        "by date descending, limited to 5). These queries execute live in Obsidian "
        "and update automatically as new notes are created."
    )

    doc.add_heading("11.3 external_context_hook.py Implementation", level=2)
    doc.add_paragraph(
        "The hook implementation follows the contract in "
        "scripts/external_context_hook.py.example: a fetch_context(task, stage) function "
        "that returns a string. For the Obsidian integration, it extracts search terms "
        "from the task objective (truncated to 200 characters), queries the vault via "
        "/opt/homebrew/bin/obsidian-cli search-content, reads the top 5 results via "
        "obsidian-cli print, and concatenates them with markdown headers. The total "
        "output respects EXTERNAL_CONTEXT_MAX_BYTES (500 KB). All operations are wrapped "
        "in try/except with logging to jarvis-pipeline-hooks.log."
    )

    doc.add_heading("11.4 Obsidian Git Configuration", level=2)
    doc.add_paragraph(
        "The Obsidian Git plugin is configured with a 5-minute auto-save interval, "
        "30-minute auto-push interval, merge-based sync method, and commit messages "
        "in the format 'vault: auto-backup {{date}}'. The git binary path is set to "
        "/usr/bin/git (system git, not Homebrew, since Obsidian Git runs inside "
        "Obsidian.app's context which may not have Homebrew PATH)."
    )

    doc.add_heading("11.5 jarvis-program.md Initial Content", level=2)
    doc.add_paragraph(
        "The initial jarvis-program.md contains: a Daily Briefing Checklist (check "
        "calendar, pipeline state, incidents, tasks, write summary to vault), a Weekly "
        "Self-Review protocol (count metrics, compare to last week, identify wins and "
        "failures, propose 1-3 amendments with evidence), Current Focus Areas (pipeline "
        "reliability, auth lifecycle monitoring, vault population), and Constraints "
        "(never auto-approve amendments, never modify the file directly, measure before "
        "claiming improvement). This file is inspired by karpathy/autoresearch's "
        "program.md pattern and evolves over time through the measured amendment cycle."
    )

    doc.add_heading("11.6 Dropbox Backup Strategy", level=2)
    doc.add_paragraph(
        "Weekly vault backup to Dropbox at /Jarvis Reports/jarvis-vault-backup/ via "
        "the existing dropbox-cli.py which handles token refresh internally using the "
        "refresh token from the Tokens Google Doc (tokens_doc.md). The backup script "
        "zips the vault directory, uploads via the Dropbox API, and logs success or "
        "failure to incidents.jsonl. Failures are non-blocking — the vault is "
        "local-first with Obsidian Git providing the primary off-machine backup."
    )

    doc.add_heading("11.7 Apple Notes Bridge (Optional)", level=2)
    doc.add_paragraph(
        "For users who have existing knowledge in Apple Notes, the memo CLI (OpenClaw "
        "apple-notes skill) can export selected notes to the vault. This is a one-time "
        "migration step, not an ongoing sync. The export preserves markdown formatting "
        "and places notes in Resources/ or the appropriate PARA folder."
    )

    # ==================== 12. PAST LESSONS APPLIED ====================
    doc.add_heading("12. Past Lessons Applied", level=1)

    lessons = [
        ("jarvis_launchd_path_bug.md",
         "All cron and launchd-invoked commands use absolute paths to Homebrew binaries. "
         "The EnvironmentVariables.PATH block is included in any new plist. This is the "
         "single most impactful lesson — it prevented a 2-day silent failure in the "
         "session monitor."),
        ("feedback_quality_audit.md",
         "Every phase includes a /simplify plus external audit checkpoint before code "
         "enters the repo. The pipeline hybrid Quality gate (deterministic verificationCmd "
         "plus LLM review) provides the external audit layer."),
        ("gog_oauth_scopes.md",
         "Daily briefing vault writes include gog calendar list which could fail if the "
         "OAuth token expires. The plan requires graceful failure handling: skip the "
         "calendar section, log an incident, and continue with remaining sections."),
        ("jarvis_orchestration_gap.md",
         "All vault writes happen in deterministic Python scripts (cron, Librarian, "
         "post-conductor hooks), not via sessions_spawn which is affected by the one-turn "
         "CLI limit. The plan is structured for execution via jarvis_pipeline.py."),
        ("jarvis_pipeline_runner.md",
         "The external_context_hook.py implementation follows the documented contract. "
         "The tasks.json entry follows the required schema. Pipeline execution uses "
         "jarvis_pipeline.py with proper verificationCmd."),
        ("librarian_weekly_patterns.md",
         "The recurring pipeline quality-failure pattern (2 incidents in 7 days) is "
         "treated as HIGH severity. All tasks have verificationCmd, quality gates are "
         "enforced, and modified scripts go through review."),
        ("tokens_doc.md",
         "Dropbox upload uses dropbox-cli.py which relies on the refresh token from "
         "the Tokens Google Doc. Token refresh is handled internally by the CLI."),
    ]

    for name, application in lessons:
        doc.add_paragraph(f"{name}: {application}")

    return doc


def generate() -> Path:
    """Generate the .docx file with atomic write."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = build_document()

    # Atomic write: save to temp file, then os.replace (Auditor patch #5)
    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx", dir=str(OUTPUT_DIR))
    os.close(tmp_fd)
    try:
        doc.save(tmp_path)
        os.replace(tmp_path, str(OUTPUT_PATH))
    except Exception:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise

    # Compute and print SHA256 for verification
    sha = hashlib.sha256(OUTPUT_PATH.read_bytes()).hexdigest()
    print(f"Generated: {OUTPUT_PATH}")
    print(f"SHA256: {sha}")
    return OUTPUT_PATH


if __name__ == "__main__":
    generate()
