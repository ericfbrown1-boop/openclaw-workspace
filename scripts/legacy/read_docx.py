#!/usr/bin/env python3
import sys

try:
    from docx import Document
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "-q", "python-docx"])
    from docx import Document

def read_word_doc(filepath):
    """Read and display Word document content"""
    print(f"Reading: {filepath}\n")
    print("="*80)
    
    doc = Document(filepath)
    
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
            full_text.append(para.text)
    
    # Check for tables
    if doc.tables:
        print("\n[Document contains tables]")
        for i, table in enumerate(doc.tables):
            print(f"\nTable {i+1}:")
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                print("  |  ".join(row_data))
    
    print("="*80)
    print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    
    # Save text version
    text_file = filepath.replace('.docx', '_content.txt')
    with open(text_file, 'w') as f:
        f.write('\n'.join(full_text))
    print(f"\nText version saved to: {text_file}")
    
    return doc, full_text

if __name__ == "__main__":
    doc, text = read_word_doc("NARR memo Feb 2026.docx")
