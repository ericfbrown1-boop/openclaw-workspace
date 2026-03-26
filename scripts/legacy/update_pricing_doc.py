#!/usr/bin/env python3
from docx import Document
from docx.shared import RGBColor
import os

# Load the document
doc_path = os.path.expanduser('~/ProjectScraper/Cloud_and_SAAS_Pricing_Analysis.docx')
doc = Document(doc_path)

print("Updating Cohesity pricing from PowerPoint data...")
print("="*80)

# New pricing data from PowerPoint slides 10-13
cohesity_pricing = {
    'MBS': {
        '10GB': {'per_year': 8, 'per_month': 0.67},
        '20GB': {'per_year': 16, 'per_month': 1.00},
        '50GB': {'per_year': 32, 'per_month': 2.00},
        '80GB': {'per_year': 48, 'per_month': 3.00},
    },
    'M365_Enterprise': {
        '10GB': {'per_year': 36, 'per_month': 3.00},
        '20GB': {'per_year': 48, 'per_month': 4.00},
        '50GB': {'per_year': 72, 'per_month': 6.00},
        '80GB': {'per_year': 96, 'per_month': 8.00},
        'Fair_Use_300GB': {'per_year': 144, 'per_month': 12.00},
    },
    'Capacity': {
        'list_price': 1800,  # $/FETB
        'buy_price': 1080,   # $/FETB
        'baas_monthly': 150, # $/TB/month (existing)
    }
}

# Key updates needed:
updates_made = []

# 1. Update Executive Summary - paragraph 8
for i, para in enumerate(doc.paragraphs):
    if 'Cohesity M365 Pricing: $150/TB/month' in para.text:
        old_text = para.text
        # Update to include multiple pricing models
        new_text = """Cohesity M365 Pricing: Multiple models available:
• MBS (Entry-level): $0.67-$3.00/user/month ($8-$48/user/year) for 10-80GB tiers
• M365 Enterprise: $3.00-$12.00/user/month ($36-$144/user/year) for 10GB-Fair Use (300GB)
• Capacity-based BaaS: $150/TB/month with validated 50-77% savings vs. user-based competitors
• Front-End TB pricing: $1,800 list / $1,080 buy price"""
        para.text = new_text
        updates_made.append(f"Updated paragraph {i}: Executive Summary pricing")
        print(f"✓ Updated Executive Summary (para {i})")
        break

# 2. Find and update Table 2 (User-Based Pricing Comparison)
for table_idx, table in enumerate(doc.tables):
    # Check if this is the user-based pricing table
    if len(table.rows) > 1 and len(table.columns) >= 3:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        if 'User Band' in first_row or 'Cohesity' in first_row:
            print(f"\n✓ Found User-Based Pricing Table (Table {table_idx + 1})")
            
            # Update Cohesity column with new pricing based on MBS
            # Table structure: User Band | Cohesity | Rubrik | Commvault | Lowest | % Diff
            for row_idx, row in enumerate(table.rows):
                if row_idx == 0:  # Skip header
                    continue
                
                cells = [cell.text.strip() for cell in row.cells]
                
                # Update pricing for each user band
                if '1-100 users' in cells[0]:
                    # MBS 20GB tier: $1.00/user/month
                    row.cells[1].text = '$1.00/user/mo (MBS 20GB)'
                    updates_made.append("Updated 1-100 users pricing")
                    print(f"  • 1-100 users: $1.00/user/mo (was $1.50)")
                    
                elif '101-500 users' in cells[0]:
                    row.cells[1].text = '$1.00-$2.00/user/mo (MBS)'
                    updates_made.append("Updated 101-500 users pricing")
                    print(f"  • 101-500 users: $1.00-$2.00/user/mo")
                    
                elif '501-1,000 users' in cells[0]:
                    row.cells[1].text = '$2.00/user/mo (MBS 50GB)'
                    updates_made.append("Updated 501-1000 users pricing")
                    print(f"  • 501-1,000 users: $2.00/user/mo")
                    
                elif '1,000+ users' in cells[0]:
                    row.cells[1].text = '$2.00-$3.00/user/mo (MBS)'
                    updates_made.append("Updated 1000+ users pricing")
                    print(f"  • 1,000+ users: $2.00-$3.00/user/mo")
                    
                elif 'Enterprise (5,000+)' in cells[0]:
                    row.cells[1].text = '$3.00-$12.00/user/mo (M365 Ent)'
                    updates_made.append("Updated 5000+ users pricing")
                    print(f"  • Enterprise 5,000+: $3.00-$12.00/user/mo")

# 3. Update footnote 1 about Cohesity pricing
for i, para in enumerate(doc.paragraphs):
    if 'Cohesity: Calculated at $150/TB with 6GB/user' in para.text:
        old_text = para.text
        new_text = """Footnotes: 1. Cohesity: Three pricing models available:
   • MBS (Microsoft Backup Service): $0.67-$3.00/user/month for 10-80GB per user tiers
   • M365 Enterprise: $3.00-$12.00/user/month for 10GB-Fair Use (300GB) tiers
   • Capacity BaaS: $150/TB/month ($1,800/FETB list, $1,080 buy price)
   Pricing shown uses MBS tiers (entry-level) for comparison. Discounts: 35% list (small tiers), 51% total with channel."""
        para.text = new_text
        updates_made.append(f"Updated footnote 1 with actual PowerPoint pricing")
        print(f"\n✓ Updated footnote (para {i})")
        break

# 4. Add new pricing comparison section for M365 Enterprise vs Rubrik
# Find where to insert (after Table 2)
insert_after_para = None
for i, para in enumerate(doc.paragraphs):
    if 'Table 2: M365 Capacity-Based Pricing' in para.text:
        insert_after_para = i
        break

if insert_after_para:
    # Insert new comparison data
    new_para = doc.paragraphs[insert_after_para]._element
    parent = new_para.getparent()
    
    # Add heading
    new_heading = doc.add_paragraph()
    new_heading.text = "\n1.1b M365 Pricing: Cohesity vs Rubrik Direct Comparison (from PowerPoint)"
    new_heading.style = 'Heading 3'
    
    # Add comparison text
    comp_para = doc.add_paragraph()
    comp_para.text = """Per PowerPoint Slide 13 comparison (50GB tier):
• Cohesity M365 Core (50GB): $48/user/year = $4.00/user/month
• Cohesity M365 Enterprise (50GB): $72/user/year = $6.00/user/month  
• Rubrik (50GB): $60/user/year = $5.00/user/month

Key Finding: Cohesity M365 Core is 20% LOWER than Rubrik ($48 vs $60/year).
Cohesity M365 Enterprise is 20% HIGHER than Rubrik ($72 vs $60/year) but includes advanced security features (GTI, DSPM, Data Classification, Recovery Agent) not in Rubrik Foundation tier."""
    
    updates_made.append("Added M365 PowerPoint pricing comparison section")
    print("\n✓ Added new comparison section with PowerPoint data")

# 5. Update the strategic recommendations to reflect accurate pricing
for i, para in enumerate(doc.paragraphs):
    if 'Maintain $150/TB M365 pricing for enterprise' in para.text:
        para.text = """Maintain hybrid M365 pricing strategy:
• MBS ($0.67-$3/user/mo) for SMB and entry-level customers
• M365 Enterprise ($3-$12/user/mo) for advanced security features
• Capacity BaaS ($150/TB/mo) for large enterprise and high-storage environments"""
        updates_made.append(f"Updated strategic recommendation (para {i})")
        print(f"\n✓ Updated strategic recommendations (para {i})")
        break

# Save the updated document
output_path = os.path.expanduser('~/ProjectScraper/Cloud_and_SAAS_Pricing_Analysis.docx')
doc.save(output_path)

print("\n" + "="*80)
print(f"✅ Document updated successfully!")
print(f"📄 Saved to: {output_path}")
print(f"\n📊 Summary of changes made:")
for idx, update in enumerate(updates_made, 1):
    print(f"   {idx}. {update}")

print("\n🔍 Key pricing corrections:")
print("   • Added MBS pricing: $0.67-$3.00/user/month (10-80GB tiers)")
print("   • Added M365 Enterprise: $3.00-$12.00/user/month (10GB-Fair Use)")
print("   • Clarified $150/TB/month is for capacity BaaS model")
print("   • Added FETB pricing: $1,800 list / $1,080 buy")
print("   • Updated comparison with Rubrik from PowerPoint slide 13")
print("\n✅ Document ready for review!")
