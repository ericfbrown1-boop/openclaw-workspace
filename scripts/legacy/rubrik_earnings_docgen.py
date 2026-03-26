#!/usr/bin/env python3
"""Generate Rubrik Q4 FY2026 Earnings Analysis Word Document"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Rubrik, Inc. (NYSE: RBRK)', level=0)
subtitle = doc.add_heading('Q4 FY2026 Earnings Analysis — Sell-Side Perspective', level=1)
doc.add_paragraph('Report Date: March 12, 2026 | Reporting Period: Q4 FY2026 (ended January 31, 2026)')
doc.add_paragraph('Fiscal Calendar: February 1 – January 31 | IPO: April 2024 at $32/share')
doc.add_paragraph('')

# Executive Summary
doc.add_heading('Executive Summary', level=2)
doc.add_paragraph(
    'Rubrik delivered a blowout Q4 FY2026, materially beating consensus on all key metrics. '
    'Revenue of $377.7M (+46% YoY) crushed the Street estimate of $342.4M by 10.3%. Non-GAAP EPS '
    'of $0.04 beat consensus of -$0.11 by $0.15, marking the company\'s first quarter of non-GAAP '
    'profitability. Subscription ARR reached $1.46B (+34% YoY) with a record $115M in net new '
    'subscription ARR. Free cash flow was $70.1M in Q4 and $237.8M for the full year — a dramatic '
    'improvement from $21.6M in FY2025. The stock traded down ~6% after hours to ~$51, likely '
    'reflecting macro headwinds and profit-taking after the beat, despite guidance that also '
    'exceeded expectations.'
)

# Q4 Highlights
doc.add_heading('Q4 FY2026 Key Metrics vs. Consensus', level=2)
t = doc.add_table(rows=6, cols=4)
t.style = 'Light Grid Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Metric', 'Actual', 'Consensus', 'Beat/Miss']
for i, h in enumerate(headers):
    t.rows[0].cells[i].text = h
data = [
    ['Total Revenue', '$377.7M', '$342.4M', '+10.3% Beat'],
    ['Non-GAAP EPS', '$0.04', '-$0.11', '+$0.15 Beat'],
    ['Subscription ARR', '$1.46B', '~$1.40B', 'Beat'],
    ['Net New Sub ARR', '$115M (record)', '~$95M', 'Significant Beat'],
    ['Free Cash Flow', '$70.1M', 'N/A', 'Strong'],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        t.rows[r+1].cells[c].text = val

doc.add_paragraph('')

# Revenue Trend
doc.add_heading('Total Revenue — Last 5 Quarters YoY Growth', level=2)
t2 = doc.add_table(rows=6, cols=5)
t2.style = 'Light Grid Accent 1'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
headers2 = ['Quarter', 'Period Ended', 'Revenue ($M)', 'YoY Growth', 'QoQ Growth']
for i, h in enumerate(headers2):
    t2.rows[0].cells[i].text = h
rev_data = [
    ['Q4 FY2025', 'Jan 31, 2025', '$258.1', '+47%', '+9.3%'],
    ['Q1 FY2026', 'Apr 30, 2025', '$278.5', '+49%', '+7.9%'],
    ['Q2 FY2026', 'Jul 31, 2025', '$309.9', '+51%', '+11.3%'],
    ['Q3 FY2026', 'Oct 31, 2025', '$350.2', '+48%', '+13.0%'],
    ['Q4 FY2026', 'Jan 31, 2026', '$377.7', '+46%', '+7.9%'],
]
for r, row_data in enumerate(rev_data):
    for c, val in enumerate(row_data):
        t2.rows[r+1].cells[c].text = val

doc.add_paragraph(
    'Revenue growth has sustained in the mid-to-high 40s% range over the last five quarters, '
    'a remarkable trajectory for a company at this scale. Full FY2026 revenue was $1.32B (+48% YoY), '
    'up from $886.5M in FY2025. Subscription revenue was $1.26B (+53% YoY). Note: ~$18M of Q4 '
    'revenue came from material rights recognition; excluding this, revenue growth was 43% YoY.'
)

# ARR Trend
doc.add_heading('Subscription ARR — Last 5 Quarters YoY Growth', level=2)
t3 = doc.add_table(rows=6, cols=5)
t3.style = 'Light Grid Accent 1'
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
headers3 = ['Quarter', 'As-of Date', 'Sub ARR ($M)', 'YoY Growth', 'Net New ARR ($M)']
for i, h in enumerate(headers3):
    t3.rows[0].cells[i].text = h
arr_data = [
    ['Q4 FY2025', 'Jan 31, 2025', '$1,092.6', '+39%', '~$101'],
    ['Q1 FY2026', 'Apr 30, 2025', '$1,180', '+38%', '~$87'],
    ['Q2 FY2026', 'Jul 31, 2025', '$1,252', '+36%', '~$71'],
    ['Q3 FY2026', 'Oct 31, 2025', '$1,350', '+34%', '~$98'],
    ['Q4 FY2026', 'Jan 31, 2026', '$1,460', '+34%', '$115 (record)'],
]
for r, row_data in enumerate(arr_data):
    for c, val in enumerate(row_data):
        t3.rows[r+1].cells[c].text = val

doc.add_paragraph(
    'While YoY ARR growth has decelerated from 39% to 34% (natural at scale), the Q4 net new ARR '
    'of $115M was a record, showing re-acceleration in absolute dollar terms. This is a critical '
    'signal — it demonstrates the flywheel is accelerating even as the denominator grows. '
    'The sequential acceleration from Q2\'s $71M net new ARR trough to Q4\'s $115M is impressive.'
)

# Guidance
doc.add_heading('FY2027 & Q1 FY2027 Guidance vs. Consensus', level=2)
t4 = doc.add_table(rows=8, cols=4)
t4.style = 'Light Grid Accent 1'
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
headers4 = ['Metric', 'Guidance', 'Consensus (Pre-Report)', 'Commentary']
for i, h in enumerate(headers4):
    t4.rows[0].cells[i].text = h
guide_data = [
    ['Q1 FY27 Revenue', '$365M–$367M', '~$350M est.', 'Midpoint $366M = ~31% YoY'],
    ['Q1 FY27 Non-GAAP EPS', '-$0.04 to -$0.02', '-$0.23', 'Dramatically above Street'],
    ['FY27 Revenue', '$1,597M–$1,607M', '~$1,600M', 'Midpoint $1,602M = ~21% YoY'],
    ['FY27 Sub ARR', '$1,829M–$1,839M', 'N/A', 'Midpoint $1,834M = ~26% YoY'],
    ['FY27 Non-GAAP EPS', '$0.07–$0.27', 'Negative', 'First profitable FY expected'],
    ['FY27 FCF', '$265M–$275M', 'N/A', 'Midpoint $270M, ~17% margin'],
    ['FY27 Sub ARR Contrib Margin', '~13%', 'N/A', 'Up from 11.6% in Q4'],
]
for r, row_data in enumerate(guide_data):
    for c, val in enumerate(row_data):
        t4.rows[r+1].cells[c].text = val

doc.add_paragraph(
    'Guidance implies a step-down in revenue growth from 46-48% to ~21% for FY2027. However, Q1 '
    'guidance of $366M at midpoint represents ~31% YoY growth (vs. Q1 FY2026\'s $278.5M), suggesting '
    'the full-year guide is conservative and typical of Rubrik\'s beat-and-raise pattern. The shift '
    'to non-GAAP profitability for the full year is a milestone — FY2026 was breakeven at -$0.01 '
    'non-GAAP EPS.'
)

# Margin Trends
doc.add_heading('Profitability & Margin Trends', level=2)
t5 = doc.add_table(rows=6, cols=5)
t5.style = 'Light Grid Accent 1'
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
headers5 = ['Quarter', 'Non-GAAP GM', 'Sub ARR Contrib Margin', 'FCF ($M)', 'FCF Margin']
for i, h in enumerate(headers5):
    t5.rows[0].cells[i].text = h
margin_data = [
    ['Q4 FY2025', '79.7%', '2.1%', '$75.2', '29.1%'],
    ['Q1 FY2026', '80.5%', 'N/R', '$33.3', '12.0%'],
    ['Q2 FY2026', '81.6%', '~9%', '$57.5', '18.6%'],
    ['Q3 FY2026', '83.0%', '10.3%', '$76.9', '22.0%'],
    ['Q4 FY2026', '83.7%', '11.6%', '$70.1', '18.6%'],
]
for r, row_data in enumerate(margin_data):
    for c, val in enumerate(row_data):
        t5.rows[r+1].cells[c].text = val

doc.add_paragraph(
    'Non-GAAP gross margins have expanded steadily from 79.7% to 83.7% over five quarters, driven '
    'by subscription mix shift, cloud hosting efficiencies, and scale. Subscription ARR contribution '
    'margin — Rubrik\'s key operating leverage metric — improved from 2.1% to 11.6%, demonstrating '
    'the unit economics are scaling. Full-year FY2026 FCF of $237.8M (18% margin) was a dramatic '
    'improvement from $21.6M in FY2025. FY2027 FCF guidance of $265-275M implies continued expansion.'
)

# Full Year Summary
doc.add_heading('Full-Year FY2026 Summary', level=2)
t6 = doc.add_table(rows=7, cols=3)
t6.style = 'Light Grid Accent 1'
t6.alignment = WD_TABLE_ALIGNMENT.CENTER
headers6 = ['Metric', 'FY2026', 'FY2025']
for i, h in enumerate(headers6):
    t6.rows[0].cells[i].text = h
fy_data = [
    ['Total Revenue', '$1,316M (+48%)', '$886.5M'],
    ['Subscription Revenue', '$1,264M (+53%)', '$828.7M'],
    ['GAAP Gross Margin', '80.1%', '70.0%'],
    ['Non-GAAP Gross Margin', '82.3%', '78.0%'],
    ['Free Cash Flow', '$237.8M', '$21.6M'],
    ['Non-GAAP EPS', '-$0.01', '-$1.57'],
]
for r, row_data in enumerate(fy_data):
    for c, val in enumerate(row_data):
        t6.rows[r+1].cells[c].text = val

# Valuation
doc.add_heading('Valuation Snapshot', level=2)
doc.add_paragraph(
    'As of market close March 12, 2026: $54.43 (-3.8% on the day). After-hours: ~$56.30 (+3.4% post-earnings).'
)
t7 = doc.add_table(rows=8, cols=2)
t7.style = 'Light Grid Accent 1'
t7.alignment = WD_TABLE_ALIGNMENT.CENTER
val_data = [
    ['Market Cap', '~$10.9B'],
    ['EV/Revenue (NTM FY2027, $1.60B)', '~6.8x'],
    ['EV/ARR (Current $1.46B)', '~7.5x'],
    ['EV/ARR (FY27 Midpoint $1.83B)', '~6.0x'],
    ['EV/FCF (FY27 Guide $270M)', '~40x'],
    ['52-Week Range', '$46.36 – $103.00'],
    ['Analyst Consensus', 'Moderate Buy | PT $98.68'],
    ['Cash & ST Investments', '$1.68B'],
]
for r, row_data in enumerate(val_data):
    t7.rows[r].cells[0].text = row_data[0]
    t7.rows[r].cells[1].text = row_data[1]

doc.add_paragraph(
    'At ~$55, RBRK trades at ~6.8x NTM revenue and ~40x NTM FCF. This is well below '
    'its historical average and the ~$98.68 consensus price target. The stock is down ~47% from '
    'its 52-week high of $103, reflecting broader tech selloffs and macro uncertainty. For a company '
    'delivering 46% revenue growth, 83%+ gross margins, and positive FCF, this valuation appears '
    'compressed. The risk is deceleration, but Q4\'s record net new ARR argues against that narrative.'
)

# Customer Metrics
doc.add_heading('Customer & Go-to-Market Metrics', level=2)
doc.add_paragraph(
    '• Customers with $100K+ ARR: 2,805 (+25% YoY)\n'
    '• $100K+ customers represent ~85% of total subscription ARR\n'
    '• DBNR: >120% (stable)\n'
    '• Insiders own 32.4% of shares\n'
    '• Institutional ownership: 49.5%'
)

# Product & Strategic Highlights
doc.add_heading('Product & Strategic Highlights', level=2)
doc.add_paragraph(
    '• Rubrik Agent Cloud (RAC): GA\'d with Amazon Bedrock AgentCore and Microsoft Copilot Studio '
    'integrations — positions Rubrik as the governance layer for enterprise AI agents\n'
    '• Rubrik Security Cloud Sovereign: Data sovereignty for regulated industries\n'
    '• Intelligent Business Recovery for M365: Automated recovery across Exchange, OneDrive, '
    'SharePoint, Teams\n'
    '• DevOps Protection: Azure DevOps and GitHub pipeline protection\n'
    '• McLaren Racing partnership: Multi-year technology partnership with F1 team\n'
    '• Jesse Green promoted to CRO to lead global revenue organization'
)

# Competitive Context
doc.add_heading('Competitive Context (Cohesity Peer View)', level=2)
doc.add_paragraph(
    'Rubrik continues to aggressively position against both legacy vendors (Commvault, Veritas) '
    'and next-gen competitors (Cohesity, Veeam). Key competitive dynamics:\n\n'
    '• Rubrik\'s 34% ARR growth and 46% revenue growth outpace Commvault (18% ARR growth)\n'
    '• The "AI operations" and "cyber resilience" narrative is resonating — management referenced '
    'displacement of legacy vendors in multiple customer wins\n'
    '• Agent Cloud (RAC) and AI governance positioning is ahead of competitors\n'
    '• SBC remains elevated at ~$329M for FY2026 (~25% of revenue) but declining as a percentage\n'
    '• GAAP net loss was $(1.78)/share for FY2026 vs. $(7.48) in FY2025 — improving but still red'
)

# Risks
doc.add_heading('Key Risks & Watchpoints', level=2)
doc.add_paragraph(
    '1. ARR growth deceleration: 39% → 34% over five quarters; FY2027 guide implies ~26%\n'
    '2. Material rights revenue: ~$18M in Q4 from non-recurring items inflates growth\n'
    '3. Stock-based compensation: $329M in FY2026 (~25% of revenue)\n'
    '4. Convertible debt: ~$1.1B outstanding\n'
    '5. Macro/IT spending: Enterprise budget compression in uncertain macro\n'
    '6. Competition: Cohesity-Veritas merger creates formidable competitor\n'
    '7. Valuation reset: Stock down 47% from highs despite strong fundamentals\n'
    '8. Rosenblatt downgraded expectations pre-earnings (contrarian signal?)'
)

# Analyst Sentiment
doc.add_heading('Wall Street Analyst Sentiment', level=2)
doc.add_paragraph(
    '• 22 Buy ratings, 1 Strong Buy, 1 Hold, 1 Sell\n'
    '• Consensus rating: Moderate Buy\n'
    '• Consensus price target: $98.68 (~80% upside from current)\n'
    '• Notable moves: Goldman Sachs cut PT from $120 to $80 (Buy); Mizuho cut from $97 to $80 (Outperform)\n'
    '• Deutsche Bank and UBS reiterated Buy ratings\n'
    '• Weiss Ratings: Sell (D-) — contrarian'
)

# Bottom Line
doc.add_heading('Bottom Line', level=2)
doc.add_paragraph(
    'Rubrik delivered an exceptional Q4 and full FY2026. The combination of 46% revenue growth, '
    '83.7% non-GAAP gross margins, record $115M net new ARR, and $237.8M annual FCF makes this '
    'one of the strongest prints in the data protection/cyber resilience space. The pivot to '
    'non-GAAP profitability is a significant milestone. At ~6.8x NTM revenue and ~$55, the stock '
    'appears significantly undervalued relative to growth and the $98.68 consensus target. The '
    'primary debate is whether growth can sustain above 25-30% as the ARR base scales past $1.5B. '
    'Q4\'s record net new ARR is a strong data point that it can. This is a name to watch closely '
    'for competitive intelligence purposes.'
)

doc.add_paragraph('')
doc.add_paragraph('')
p = doc.add_paragraph('Sent by Jarvis - AI assistant to Eric Brown')
p.runs[0].italic = True

output_path = os.path.expanduser('~/rubrik_earnings_analysis.docx')
doc.save(output_path)
print(f'Document saved to {output_path}')
