from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create new document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title
title = doc.add_heading('Cryoseisms: A Comprehensive Analysis of Snow Quakes', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Physical Mechanics, Thermodynamic Implications, and Case Studies')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(14)
subtitle_format.font.italic = True

doc.add_paragraph()
doc.add_paragraph('Prepared: February 2026')
doc.add_paragraph()

print("Document created successfully")
doc.save('/Users/ericbrown/.openclaw/workspace/Snow_Quakes_Report.docx')
