#!/usr/bin/env python3
"""Create Word document with top 15 CoWork skills for CFO/COO."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('Top 15 Claude CoWork Skills for CFO/COO', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Email Triage & Executive Productivity Skills')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x99)
run.bold = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.add_run('Prepared for: ').bold = True
meta.add_run('Eric Brown, CFO & COO, Cohesity')
meta.add_run('\n')
meta.add_run('Date: ').bold = True
meta.add_run(datetime.date.today().strftime('%B %d, %Y'))
meta.add_run('\n')
meta.add_run('Sources: ').bold = True
meta.add_run('GitHub (anthropics/skills, VoltAgent/awesome-agent-skills, alirezarezvani/claude-skills, K-Dense-AI), FindSkill.ai, Snyk, Reddit r/ClaudeAI, Substack')

doc.add_page_break()

# ==========================================================
# TABLE OF CONTENTS
# ==========================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    'PART A: Top 5 Email Triage & Review Skills',
    '  1. GOG Email Triage (Anthropic PR #299)',
    '  2. Inbox Classifier (LongRacks Labs)',
    '  3. GOG Email Draft & Send Suite',
    '  4. Internal Comms (Anthropic Official)',
    '  5. Custom CFO Email Triage Skill',
    '',
    'PART B: Top 10 CFO/COO Productivity Skills',
    '  6. Excel/XLSX Skill (Anthropic Official)',
    '  7. PowerPoint/PPTX Skill (Anthropic Official)',
    '  8. Word/DOCX Skill (Anthropic Official)',
    '  9. CEO Strategic Advisor (alirezarezvani)',
    '  10. Anthropic Financial Modeling Suite',
    '  11. Cost-Benefit Analysis (FindSkill.ai)',
    '  12. Competitive Intelligence Monitor',
    '  13. Board Deck Review & Preparation',
    '  14. Contract & Vendor Review',
    '  15. Executive Meeting Prep & Briefing',
]
for item in toc_items:
    if item == '':
        doc.add_paragraph()
    elif item.startswith('PART'):
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.bold = True
    else:
        doc.add_paragraph(item)

doc.add_page_break()

# ==========================================================
# PART A: EMAIL SKILLS
# ==========================================================
doc.add_heading('PART A: Top 5 Email Triage & Review Skills', level=1)

# --- SKILL 1 ---
doc.add_heading('1. GOG Email Triage', level=2)
p = doc.add_paragraph()
p.add_run('Source: ').bold = True
p.add_run('github.com/anthropics/skills (PR #299 by meticulo3366)')
p.add_run('\n')
p.add_run('Popularity: ').bold = True
p.add_run('Official Anthropic skills repo — PR submitted with 7,050+ lines of documentation')
p.add_run('\n')
p.add_run('What it does: ').bold = True
p.add_run('Inbox review and prioritization with urgency classification using the GOG CLI (Google Workspace). Triages unread emails into P0-P3 priority levels, generates daily summaries, and flags action items.')

doc.add_paragraph('Skill Content (paste into Skill Creator):')
skill1 = '''---
name: gog-email-triage
description: >
  Inbox review and prioritization with urgency classification.
  Use this skill when the user asks to check email, triage inbox,
  review unread messages, prioritize emails, or create a daily
  email summary. Requires GOG CLI v0.9.0+ authenticated with Gmail.
argument-hint: "[timeframe] e.g., 'today', 'last 24h', 'this week'"
---

# GOG Email Triage

## Purpose
Review and prioritize inbox emails using urgency classification.

## Priority Classification System
- **P0 — Immediate**: Revenue impact, legal/compliance, CEO/Board requests,
  security incidents, customer escalations
- **P1 — Today**: Direct reports' blockers, partner/vendor deadlines,
  time-sensitive approvals, meeting prep needed today
- **P2 — This Week**: FYI from leadership, non-urgent approvals,
  project updates, internal announcements
- **P3 — When Available**: Newsletters, marketing, social notifications,
  informational CCs, automated reports

## Workflow

### Step 1: Fetch unread emails
```bash
gog gmail search "is:unread" --max 50 --format json
```

### Step 2: For each email, classify priority
Evaluate based on:
- Sender (executive team, direct reports, external partners, automated)
- Subject line keywords (urgent, action required, approval, deadline, FYI)
- Thread participation (am I TO or CC?)
- Time sensitivity (dates mentioned, deadlines referenced)

### Step 3: Generate triage report
```
## 📧 Email Triage Report — [Date]
**Unread:** [count] | **Needs Action:** [count]

### 🔴 P0 — Immediate (Act Now)
- [Sender]: [Subject] — [Why it's P0] — [Suggested action]

### 🟠 P1 — Today
- [Sender]: [Subject] — [Key ask] — [Suggested action]

### 🟡 P2 — This Week
- [Sender]: [Subject] — [Summary]

### 🟢 P3 — When Available
- [count] newsletters, [count] notifications, [count] automated reports
```

### Step 4: Ask user which P0/P1 items to act on
For each action item, offer:
- Draft a reply
- Forward to someone
- Create a calendar event
- Add to task list
- Archive/snooze

## Safety
- NEVER delete emails without explicit confirmation
- NEVER send replies without "YES, SEND" confirmation
- Log all actions to ~/.gog-assistant/audit.log
- No email bodies in logs — metadata only

## Dynamic Context
Inject current state at start:
```bash
gog gmail search "is:unread" --max 5 --format brief
gog calendar today --format brief
```
'''
p = doc.add_paragraph()
run = p.add_run(skill1)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_page_break()

# --- SKILL 2 ---
doc.add_heading('2. Inbox Classifier', level=2)
p = doc.add_paragraph()
p.add_run('Source: ').bold = True
p.add_run('github.com/longrackslabs/inbox-classifier')
p.add_run('\n')
p.add_run('Popularity: ').bold = True
p.add_run('Recommended on Reddit r/ClaudeAI — "after a week of training, doing a pretty good job"')
p.add_run('\n')
p.add_run('What it does: ').bold = True
p.add_run('AI-powered email classification that learns your priorities over time. Automatically categorizes emails into custom folders/labels. Self-improving — gets better the more you use it.')

doc.add_paragraph('Skill Content:')
skill2 = '''---
name: inbox-classifier
description: >
  Classify and categorize incoming emails into priority buckets using
  learned patterns. Use when the user wants to organize their inbox,
  set up email rules, or automatically sort messages. Learns from
  user corrections to improve over time.
---

# Inbox Classifier

## Classification Categories (Customize for your role)

### Executive Categories
1. **Board & Investor** — Board members, investors, legal counsel
2. **Revenue Critical** — Customer escalations, deal approvals, renewals
3. **Direct Reports** — Your team's updates, requests, blockers
4. **Cross-Functional** — Other C-suite, department heads, strategic initiatives
5. **External Partners** — Vendors, consultants, industry contacts
6. **Informational** — Newsletters, reports, FYIs, automated notifications
7. **Personal** — Non-work items that made it to work inbox

### Per-Category Rules
For each email, determine:
- **Category** (from above)
- **Urgency**: Immediate / Today / This Week / Whenever
- **Action Type**: Reply needed / Approve / Review / FYI only / Delegate
- **Delegate To**: [Name/role if applicable]

## Learning Loop
After each triage session:
1. User corrects any misclassifications
2. Update classification rules in `~/.inbox-classifier/rules.json`
3. Track accuracy metrics over time

## Output Format
```
📥 INBOX CLASSIFICATION — [Date] [Time]
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

🔴 BOARD & INVESTOR (2 emails)
  → [Sender] — [Subject] — REPLY NEEDED by [date]
  → [Sender] — [Subject] — FYI (earnings-related)

🟠 REVENUE CRITICAL (3 emails)
  → [Sender] — [Subject] — APPROVAL needed ($X deal)

🟡 DIRECT REPORTS (5 emails)
  → [Summarized as group]

📊 Accuracy: 94% (last 7 days) | Processed: 847 emails
```
'''
p = doc.add_paragraph()
run = p.add_run(skill2)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_page_break()

# --- SKILL 3 ---
doc.add_heading('3. GOG Email Draft & Send Suite', level=2)
p = doc.add_paragraph()
p.add_run('Source: ').bold = True
p.add_run('github.com/anthropics/skills (PR #299 — companion to Email Triage)')
p.add_run('\n')
p.add_run('What it does: ').bold = True
p.add_run('Composes executive-quality email replies with tone variants (concise vs. warmer), handles follow-up tracking, and requires explicit confirmation before sending. Includes audit logging.')

doc.add_paragraph('Skill Content:')
skill3 = '''---
name: gog-email-draft
description: >
  Draft and send emails with executive-appropriate tone variants.
  Use when the user asks to reply to an email, compose a message,
  draft a response, or send a follow-up. Provides concise and
  warmer tone options for each draft.
---

# GOG Email Draft & Send

## Drafting Workflow

### Step 1: Understand context
- Read the original email thread: `gog gmail read [id]`
- Identify key asks, deadlines, stakeholders
- Note the sender's tone and formality level

### Step 2: Generate two tone variants
**Concise (CFO voice):**
- Direct, professional, action-oriented
- Lead with the decision/answer
- 3-5 sentences max for routine items

**Warmer (relationship-building):**
- Same substance, slightly more personal
- Acknowledge the sender's effort/situation
- Appropriate for board members, key partners, direct reports

### Step 3: Present both options
```
📝 DRAFT REPLY TO: [Sender] — [Subject]

━━ Option A: Concise ━━
[Draft text]

━━ Option B: Warmer ━━
[Draft text]

Choose A, B, or tell me what to adjust.
```

### Step 4: Send with confirmation
```bash
# Only after explicit "YES, SEND" from user
gog gmail send --to [addr] --subject [subj] --body [text] --cc [cc]
```

## Safety Rules
- ALWAYS show draft before sending
- REQUIRE "YES, SEND" exact text for confirmation
- Auto-CC ericfbrown1@gmail.com on all outgoing
- Append "Sent by [Assistant Name]" footer
- Log send events to audit trail

## Follow-up Tracking
After sending, offer to:
- Set a follow-up reminder (3 days default)
- Add to follow-up tracking list
- Schedule a calendar reminder
'''
p = doc.add_paragraph()
run = p.add_run(skill3)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_page_break()

# --- SKILL 4 ---
doc.add_heading('4. Internal Comms (Anthropic Official)', level=2)
p = doc.add_paragraph()
p.add_run('Source: ').bold = True
p.add_run('github.com/anthropics/skills/tree/main/skills/internal-comms')
p.add_run('\n')
p.add_run('Popularity: ').bold = True
p.add_run('Official Anthropic skill — ships with Claude')
p.add_run('\n')
p.add_run('What it does: ').bold = True
p.add_run('Write internal communications: status reports, 3P updates (Progress/Plans/Problems), company newsletters, FAQs, incident reports, leadership updates.')

doc.add_paragraph('Skill Content:')
skill4 = '''---
name: internal-comms
description: >
  A set of resources to help me write all kinds of internal
  communications, using the formats that my company likes to use.
  Claude should use this skill whenever asked to write some sort of
  internal communications (status reports, leadership updates,
  3P updates, company newsletters, FAQs, incident reports,
  project updates, etc.).
---

## When to use this skill
To write internal communications, use this skill for:
- 3P updates (Progress, Plans, Problems)
- Company newsletters
- FAQ responses
- Status reports
- Leadership updates
- Project updates
- Incident reports

## How to use this skill

To write any internal communication:

1. **Identify the communication type** from the request
2. **Load the appropriate guideline file** from the `examples/` directory:
    - `examples/3p-updates.md` - For Progress/Plans/Problems team updates
    - `examples/company-newsletter.md` - For company-wide newsletters
    - `examples/faq-answers.md` - For answering frequently asked questions
    - `examples/general-comms.md` - For anything else
3. **Follow the specific instructions** in that file for formatting,
   tone, and content gathering

## CFO/COO Enhancement: Add These Examples
Create an `examples/cfo-updates.md` with:
- Quarterly business review format
- Earnings call prep notes format
- Board update email template
- All-hands financial summary format
- Investor relations update template
'''
p = doc.add_paragraph()
run = p.add_run(skill4)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_page_break()

# --- SKILL 5 ---
doc.add_heading('5. Custom CFO Email Triage & Review Skill', level=2)
p = doc.add_paragraph()
p.add_run('Source: ').bold = True
p.add_run('Custom-built for Eric Brown, CFO/COO of Cohesity')
p.add_run('\n')
p.add_run('What it does: ').bold = True
p.add_run('Purpose-built email triage that understands Cohesity\'s organizational structure, prioritizes by business impact, and handles attachment review (forwarding to the CFO Document Review skill).')

doc.add_paragraph('Skill Content:')
skill5 = '''---
name: cfo-email-triage
description: >
  Executive email triage optimized for a CFO/COO role at an enterprise
  software company. Prioritizes by business impact, identifies action
  items, flags financial/legal sensitivity, and routes attachment
  reviews. Use when asked to check email, review inbox, triage
  messages, or prepare an email briefing.
argument-hint: "[timeframe] [filter] e.g., 'last 24h urgent only'"
---

# CFO Email Triage & Review

## Priority Matrix (Business Impact)

### 🔴 TIER 1 — Executive Action Required (< 2 hours)
- Board of Directors / Chairman communications
- CEO direct requests or escalations
- Legal/compliance matters with deadlines
- M&A or financing-related correspondence
- Revenue recognition or audit issues
- Security incidents affecting customers
- Analyst / investor inquiries during quiet period

### 🟠 TIER 2 — Same-Day Response (< 8 hours)
- Direct report escalations and blockers
- Customer deal approvals > $500K
- Budget/headcount approval requests
- Partner/vendor contract decisions
- Cross-functional leadership alignment requests

### 🟡 TIER 3 — This Week
- FP&A reports and variance analyses
- Project status updates from teams
- Internal policy/process reviews
- Vendor proposals and RFP responses
- Recruiting approvals

### 🟢 TIER 4 — Batch Processing
- Industry newsletters and analyst reports
- Automated system notifications
- Meeting invites (review with calendar context)
- Marketing/PR updates
- Training and development opportunities

## Attachment Handling
When emails contain attachments:
1. **Identify file type** (PPT, Excel, Word, PDF)
2. **Flag for CFO Document Review skill** if it's:
   - A board deck or investor presentation
   - A financial model or budget
   - A contract or legal document
   - An analyst report or competitive intel
3. **Summarize attachment content** in the triage report
4. **Note if review is time-sensitive** based on email context

## Sender Intelligence
Maintain a sender priority map:
- **Board members** → Always Tier 1
- **CEO, President** → Always Tier 1
- **CFO direct reports** (FP&A, Controller, Treasury, IR) → Tier 1-2
- **Other C-suite** → Tier 2
- **VP+ level** → Tier 2-3
- **External counsel, auditors** → Tier 1-2
- **Analysts, investors** → Tier 1 during earnings, else Tier 3
- **Vendors** → Tier 3-4

## Output Format
```
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📧 CFO INBOX BRIEFING — [Date] [Time]
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Unread: [N] | Action Required: [N] | Attachments: [N]

🔴 TIER 1 — EXECUTIVE ACTION ([count])
┌─────────────────────────────────────
│ FROM: [Name] ([Role])
│ SUBJ: [Subject]
│ RECEIVED: [Time]
│ ASK: [What they need from you]
│ DEADLINE: [If any]
│ ATTACHMENTS: [Type — needs review? Y/N]
│ SUGGESTED ACTION: [Reply / Approve / Delegate / Call]
└─────────────────────────────────────

🟠 TIER 2 — SAME DAY ([count])
[Similar format, condensed]

🟡 TIER 3 — THIS WEEK ([count])
[Summary list format]

🟢 TIER 4 — BATCH ([count])
[Count by category: X newsletters, Y notifications, Z vendor emails]

📎 ATTACHMENTS NEEDING REVIEW:
- [Filename] from [Sender] — [Type] — Priority: [H/M/L]

⏰ UPCOMING DEADLINES FROM EMAILS:
- [Date]: [What's due] — from [who]
```

## Actions After Triage
For each Tier 1-2 item, offer:
1. Draft a reply (concise CFO voice)
2. Forward with instructions to delegate
3. Schedule a call/meeting
4. Request more information
5. Flag for deeper document review
'''
p = doc.add_paragraph()
run = p.add_run(skill5)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_page_break()

# ==========================================================
# PART B: TOP 10 CFO/COO SKILLS
# ==========================================================
doc.add_heading('PART B: Top 10 CFO/COO Productivity Skills', level=1)

# --- SKILL 6 ---
doc.add_heading('6. Excel/XLSX Skill (Anthropic Official)', level=2)
p = doc.add_paragraph()
p.add_run('Source: ').bold = True
p.add_run('github.com/anthropics/skills/tree/main/skills/xlsx')
p.add_run('\n')
p.add_run('Stars: ').bold = True
p.add_run('Official Anthropic — ships with Claude CoWork')
p.add_run('\n')
p.add_run('Why you need it: ').bold = True
p.add_run('Professional financial model creation with industry-standard color coding (blue = inputs, black = formulas), proper number formatting, formula error prevention, and audit-ready outputs. Auto-triggers whenever .xlsx files are involved.')

doc.add_paragraph('Skill Content (truncated — full version at source):')
skill6 = '''---
name: xlsx
description: "Use this skill any time a spreadsheet file is the primary
  input or output. This means any task where the user wants to: open,
  read, edit, or fix an existing .xlsx, .xlsm, .csv, or .tsv file;
  create a new spreadsheet from scratch or from other data sources;
  or convert between tabular file formats."
---

# Requirements for Outputs

## Financial models

### Color Coding Standards (Industry Standard)
- **Blue text (0,0,255)**: Hardcoded inputs and assumptions
- **Black text (0,0,0)**: ALL formulas and calculations
- **Green text (0,128,0)**: Links from other worksheets
- **Red text (255,0,0)**: External links to other files
- **Yellow background (255,255,0)**: Key assumptions needing attention

### Number Formatting Standards
- **Years**: Text format ("2024" not "2,024")
- **Currency**: $#,##0 format; ALWAYS specify units in headers
- **Zeros**: Display as "-"
- **Percentages**: 0.0% format (one decimal)
- **Multiples**: 0.0x for valuation multiples
- **Negative numbers**: Parentheses (123) not -123

### Formula Rules
- ALL assumptions in separate assumption cells
- Cell references instead of hardcoded values
- Verify all references, check for off-by-one errors
- Test with edge cases (zero, negative)

[Full skill: github.com/anthropics/skills/tree/main/skills/xlsx]
'''
p = doc.add_paragraph()
run = p.add_run(skill6)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_page_break()

# --- SKILL 7 ---
doc.add_heading('7. PowerPoint/PPTX Skill (Anthropic Official)', level=2)
p = doc.add_paragraph()
p.add_run('Source: ').bold = True
p.add_run('github.com/anthropics/skills/tree/main/skills/pptx')
p.add_run('\n')
p.add_run('Why you need it: ').bold = True
p.add_run('Create, edit, and analyze PowerPoint decks. Includes design guidance (color palettes, layout principles), slide templates, and the ability to extract text from existing .pptx files. Essential for board deck creation and review.')

doc.add_paragraph('Key capabilities: Read/analyze decks, edit from templates, create from scratch with professional design, extract speaker notes. Includes 6 curated color palettes.')
doc.add_paragraph('Install: Already ships with Claude CoWork. Full source at github.com/anthropics/skills/tree/main/skills/pptx')

# --- SKILL 8 ---
doc.add_heading('8. Word/DOCX Skill (Anthropic Official)', level=2)
p = doc.add_paragraph()
p.add_run('Source: ').bold = True
p.add_run('github.com/anthropics/skills/tree/main/skills/docx')
p.add_run('\n')
p.add_run('Why you need it: ').bold = True
p.add_run('Create, edit, and analyze Word documents. Used for contracts, memos, reports, and any formal business documentation. Supports professional formatting, track changes, and template-based editing.')

doc.add_paragraph('Install: Ships with Claude CoWork. Pairs well with the CFO Document Review skill from the earlier email.')

# --- SKILL 9 ---
doc.add_heading('9. CEO Strategic Advisor', level=2)
p = doc.add_paragraph()
p.add_run('Source: ').bold = True
p.add_run('github.com/alirezarezvani/claude-skills/c-level-advisor/ceo-advisor')
p.add_run('\n')
p.add_run('Stars: ').bold = True
p.add_run('2,300+ (entire repo)')
p.add_run('\n')
p.add_run('Why you need it: ').bold = True
p.add_run('Executive leadership guidance for strategic decision-making, board governance, and investor relations. Includes Python tools for strategy analysis and financial scenario modeling. Covers board presentations, stakeholder communication, and organizational culture.')

doc.add_paragraph('Skill Content:')
skill9 = '''---
name: ceo-advisor
description: >
  Strategic advisory for executive leadership decisions.
  Use for: strategic planning, board preparation, financial
  scenario modeling, investor communications, M&A evaluation,
  and organizational strategy. Includes frameworks for
  executive decision-making and stakeholder management.
---

# CEO Strategic Advisor

## Core Capabilities
- Strategic planning and initiative evaluation
- Financial scenario modeling and business outcomes
- Executive decision framework (structured methodology)
- Leadership and organizational culture development
- Board governance and investor relations
- Stakeholder communication best practices

## Python Analysis Tools
- `strategy_analyzer.py` — Evaluate strategic initiatives
- `financial_scenario_analyzer.py` — Model financial scenarios

## Key Workflows
1. **Strategic Decision**: Problem → Analysis → Financial Model →
   Decision Framework → Stakeholder Communication
2. **Board Prep**: Content Prep → Presentation Design → Q&A Prep →
   Rehearsal → Follow-up
3. **M&A Evaluation**: Strategic fit → Financial analysis →
   Integration plan → Risk assessment → Board recommendation

Install: npx ai-agent-skills install alirezarezvani/claude-skills/c-level-advisor/ceo-advisor
'''
p = doc.add_paragraph()
run = p.add_run(skill9)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_page_break()

# --- SKILL 10 ---
doc.add_heading('10. Anthropic Financial Modeling Suite', level=2)
p = doc.add_paragraph()
p.add_run('Source: ').bold = True
p.add_run('github.com/anthropics/claude-cookbooks (Financial Skills Collection)')
p.add_run('\n')
p.add_run('Stars: ').bold = True
p.add_run('32,682')
p.add_run('\n')
p.add_run('Why you need it: ').bold = True
p.add_run('Professional-grade DCF models, Monte Carlo simulations, sensitivity analysis, and scenario planning. The most starred financial modeling collection in the Claude ecosystem.')

doc.add_paragraph('Skill Content:')
skill10 = '''---
name: financial-modeling
description: >
  Build DCF models, run Monte Carlo simulations, perform sensitivity
  analysis, and create financial scenario plans. Use when asked to
  value a company, model cash flows, stress-test assumptions, or
  build a financial forecast. Designed for institutional-quality output.
---

# Financial Modeling Suite

## Model Types
1. **DCF (Discounted Cash Flow)**
   - Revenue build-up (bottom-up or top-down)
   - WACC calculation with beta, risk-free rate, equity premium
   - Terminal value (Gordon Growth or Exit Multiple)
   - Sensitivity tables on WACC and growth rate

2. **Comparable Company Analysis**
   - EV/Revenue, EV/EBITDA, P/E multiples
   - Growth-adjusted multiples
   - Public comp set with quartile ranges

3. **LBO (Leveraged Buyout)**
   - Sources & uses, debt tranches
   - Cash flow sweep mechanics
   - IRR sensitivity to entry/exit multiples

4. **Monte Carlo Simulation**
   - Define probability distributions for key inputs
   - 10,000+ iteration simulation
   - Confidence interval visualization
   - Tornado chart for sensitivity ranking

5. **Three-Statement Model**
   - Income statement → Balance sheet → Cash flow
   - Circular reference handling (interest expense)
   - Working capital model

## Output Standards
- All models delivered as .xlsx with proper formatting
- Assumption cells clearly labeled (blue font)
- Sensitivity tables on key variables
- Executive summary tab with key metrics
- Charts for visual presentation

Install: From claude-cookbooks financial skills collection
'''
p = doc.add_paragraph()
run = p.add_run(skill10)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_page_break()

# --- SKILL 11 ---
doc.add_heading('11. Cost-Benefit Analysis', level=2)
p = doc.add_paragraph()
p.add_run('Source: ').bold = True
p.add_run('FindSkill.ai (1,236+ skill templates)')
p.add_run('\n')
p.add_run('Why you need it: ').bold = True
p.add_run('Rigorous cost-benefit analyses with NPV, ROI, IRR calculations, sensitivity testing, and intangible valuation for investment and policy decisions. Essential for evaluating vendor proposals, capex requests, and strategic investments.')

doc.add_paragraph('Skill Content:')
skill11 = '''---
name: cost-benefit-analysis
description: >
  Conduct rigorous cost-benefit analyses with NPV, ROI, IRR
  calculations, sensitivity testing, and intangible valuation
  for investment and policy decisions. Use when evaluating
  vendor proposals, technology investments, build-vs-buy
  decisions, headcount requests, or any spend > $100K.
---

# Cost-Benefit Analysis Framework

## Analysis Structure

### 1. Define the Decision
- What are we evaluating? (Investment, project, vendor, policy)
- What's the baseline (do nothing)?
- Time horizon (typically 3-5 years for tech investments)
- Discount rate (use company WACC or hurdle rate)

### 2. Quantify Costs
- **Direct costs**: License, implementation, hardware, services
- **Indirect costs**: Training, change management, productivity dip
- **Ongoing costs**: Maintenance, support, headcount, renewals
- **Opportunity costs**: What else could this money fund?
- **Risk-adjusted costs**: Probability-weighted overrun scenarios

### 3. Quantify Benefits
- **Revenue impact**: New revenue, retention improvement, upsell
- **Cost savings**: Headcount, efficiency, vendor consolidation
- **Risk reduction**: Compliance, security, business continuity
- **Intangible benefits**: Employee satisfaction, brand, speed-to-market
  (quantify using proxy metrics where possible)

### 4. Calculate Metrics
- **NPV** at company hurdle rate
- **IRR** (should exceed hurdle rate)
- **Payback period** (simple and discounted)
- **ROI** (lifetime benefit / total cost)
- **Break-even point**

### 5. Sensitivity Analysis
- Best case / Base case / Worst case
- Key variable tornado chart
- Monte Carlo if high uncertainty

### 6. Recommendation
- Go / No-Go / Conditional
- Key risks and mitigations
- Implementation timeline
- Success metrics to track

## Output: Executive Summary Format
```
## 💰 COST-BENEFIT ANALYSIS: [Project Name]
**Date:** [Today] | **Analyst:** Claude AI | **Reviewer:** CFO

| Metric | Value |
|--------|-------|
| NPV (5yr) | $X.XM |
| IRR | XX% |
| Payback | X.X years |
| ROI | XXX% |

**Recommendation:** [GO / NO-GO / CONDITIONAL]
**Key Risk:** [Primary risk and mitigation]
```
'''
p = doc.add_paragraph()
run = p.add_run(skill11)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_page_break()

# --- SKILL 12 ---
doc.add_heading('12. Competitive Intelligence Monitor', level=2)
p = doc.add_paragraph()
p.add_run('Source: ').bold = True
p.add_run('Custom — combining patterns from K-Dense-AI and marketing-skill/intelligence pod')
p.add_run('\n')
p.add_run('Why you need it: ').bold = True
p.add_run('Track Rubrik, Commvault, and Veeam in real-time. Monitor earnings, product launches, pricing changes, and leadership moves. Generate competitive briefings for board meetings.')

doc.add_paragraph('Skill Content:')
skill12 = '''---
name: competitive-intelligence
description: >
  Monitor and analyze competitors in the data protection/backup
  market. Use when asked about competitor news, competitive analysis,
  market positioning, earnings comparison, or preparing competitive
  briefings. Covers Rubrik (RBRK), Commvault (CVLT), Veeam, Veritas,
  and Dell APEX.
---

# Competitive Intelligence Monitor

## Target Competitors
| Company | Ticker | Status | Key Metrics |
|---------|--------|--------|-------------|
| Rubrik | RBRK | Public | ARR, Subscription Rev |
| Commvault | CVLT | Public | Revenue, Metallic ARR |
| Veeam | Private | Insight Partners | ARR (reported selectively) |
| Veritas | Private | Cohesity merger target | Integration status |
| Dell | DELL | Public | APEX Backup segment |

## Intelligence Categories

### Financial Intelligence (Quarterly)
- Earnings releases and guidance
- ARR growth rates and comparisons
- Margin trends (gross, operating, FCF)
- Customer count and NRR metrics
- Analyst consensus and revisions

### Product Intelligence (Ongoing)
- New product launches and features
- Pricing and packaging changes
- Technology partnerships and integrations
- Patent filings and acquisitions
- Cloud marketplace listings (AWS, Azure, GCP)

### Go-to-Market Intelligence
- Key customer wins and losses
- Channel partner changes
- Leadership hires and departures
- Geographic expansion moves
- Marketing messaging shifts

## Output: Competitive Briefing
```
## 🎯 COMPETITIVE BRIEFING — [Date]

### Headlines This Week
- [Competitor]: [News item] — Impact: [H/M/L]

### Earnings Comparison (Latest Quarter)
| Metric | Cohesity | Rubrik | Commvault | Veeam |
|--------|----------|--------|-----------|-------|
| ARR | $X | $X | $X | $X |
| YoY Growth | X% | X% | X% | X% |
| Gross Margin | X% | X% | X% | N/A |

### Competitive Implications for Cohesity
1. [Insight and recommended response]
2. [Insight and recommended response]

### Sources
- [URL 1], [URL 2], ...
```
'''
p = doc.add_paragraph()
run = p.add_run(skill12)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_page_break()

# --- SKILL 13 ---
doc.add_heading('13. Board Deck Review & Preparation', level=2)
p = doc.add_paragraph()
p.add_run('Source: ').bold = True
p.add_run('Custom — combining Anthropic pptx skill + CEO Advisor + financial modeling patterns')
p.add_run('\n')
p.add_run('Why you need it: ').bold = True
p.add_run('End-to-end board deck preparation: review existing decks for quality/accuracy, build financial summary slides, prepare appendix materials, and draft speaker notes. Includes a board-readiness checklist.')

doc.add_paragraph('Skill Content:')
skill13 = '''---
name: board-deck-review
description: >
  Review, prepare, and enhance board of directors presentation
  materials. Use when preparing for board meetings, reviewing
  board decks, creating financial summary slides, or drafting
  speaker notes for board presentations.
---

# Board Deck Review & Preparation

## Board Deck Structure (Best Practice)
1. **Cover Slide** — Meeting date, confidentiality notice
2. **Agenda** — Topics with time allocation
3. **CEO Update** — Strategic highlights, key wins, challenges
4. **CFO Update** — Financial performance vs. plan
5. **Key Metrics Dashboard** — ARR, revenue, margins, customer metrics
6. **Competitive Landscape** — Market position, competitor moves
7. **Product & Engineering** — Roadmap progress, key releases
8. **GTM Update** — Pipeline, bookings, channel performance
9. **People & Culture** — Headcount, attrition, key hires
10. **Risk & Compliance** — Legal, regulatory, security
11. **Strategic Initiatives** — M&A, partnerships, new markets
12. **Q&A / Discussion Topics** — Open items, decisions needed
13. **Appendix** — Supporting data, detailed financials

## Review Checklist
- [ ] Numbers internally consistent across all slides?
- [ ] All charts have proper labels, units, time periods?
- [ ] YoY and QoQ comparisons are apples-to-apples?
- [ ] Forward-looking statements have safe harbor language?
- [ ] Competitive data sourced and dated?
- [ ] Headcount and budget numbers match approved plan?
- [ ] Prior quarter action items addressed?
- [ ] Confidentiality markings on every slide?
- [ ] Consistent formatting (fonts, colors, alignment)?
- [ ] Speaker notes complete for each slide?
- [ ] Backup slides for anticipated questions?

## Financial Slide Standards
- ARR bridge (beginning → new → expansion → churn → ending)
- Revenue waterfall vs. plan and prior year
- Margin trends (trailing 4-8 quarters)
- Cash flow and balance sheet highlights
- Guidance vs. actual scorecard
- Key KPIs with red/yellow/green status

## Output: Board Readiness Score
Rate 1-5 on each dimension, overall READY / NEEDS WORK / NOT READY
'''
p = doc.add_paragraph()
run = p.add_run(skill13)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_page_break()

# --- SKILL 14 ---
doc.add_heading('14. Contract & Vendor Review', level=2)
p = doc.add_paragraph()
p.add_run('Source: ').bold = True
p.add_run('Custom — inspired by FindSkill.ai legal/compliance templates and business-growth/contracts-proposals')
p.add_run('\n')
p.add_run('Why you need it: ').bold = True
p.add_run('Review vendor contracts, SaaS agreements, and partnership proposals from a CFO perspective. Flags financial terms, auto-renewal traps, liability exposure, and benchmarks pricing against market.')

doc.add_paragraph('Skill Content:')
skill14 = '''---
name: contract-vendor-review
description: >
  Review vendor contracts, SaaS agreements, and partnership
  proposals from a CFO/COO perspective. Use when evaluating
  new vendor contracts, renewal terms, partnership agreements,
  or any document with financial/legal obligations.
---

# Contract & Vendor Review

## Review Framework

### Financial Terms Audit
- **Total contract value** (TCV) and annualized cost
- **Payment terms** — Net 30/60/90, upfront vs. installments
- **Price escalation** — Annual increases, CPI-linked, fixed
- **Volume commitments** — Minimums, overages, true-up mechanics
- **Discount structure** — Multi-year, volume, bundle discounts
- **Currency** — Fixed or floating FX, which party bears risk

### Risk Flags (Auto-Flag These)
🚩 Auto-renewal clauses (especially > 1 year)
🚩 Unlimited liability or uncapped indemnification
🚩 Exclusivity provisions
🚩 Non-compete restrictions
🚩 Change of control triggers
🚩 Most-favored-nation (MFN) clauses
🚩 Data ownership ambiguity
🚩 Termination penalties > 1x annual value
🚩 Audit rights limited to vendor

### Operational Terms
- **SLA commitments** — Uptime %, response times, penalties
- **Data handling** — Ownership, portability, deletion on termination
- **Integration requirements** — APIs, SSO, compliance certifications
- **Support model** — Named contacts, response tiers, escalation path
- **Insurance requirements** — Cyber, E&O, professional liability

### Benchmarking
Compare pricing against:
- Industry standard rates for similar services
- Prior year contract (if renewal)
- Competitive alternatives
- Internal cost of build/self-service

## Output Format
```
## 📋 CONTRACT REVIEW: [Vendor] — [Agreement Type]
**Date:** [Today] | **TCV:** $[X] | **Term:** [X years]

### 💰 Financial Summary
| Term | Value | Market Benchmark | Assessment |
|------|-------|-----------------|------------|
| Annual cost | $X | $X-Y range | Fair/High/Low |
| Escalation | X% | 3-5% typical | Acceptable/Concern |

### 🚩 Risk Flags
1. [Flag] — [Impact] — [Recommended negotiation point]

### ✅ Acceptable Terms
- [List of standard/favorable terms]

### 📝 Negotiation Recommendations
1. [Highest priority: what to push back on]
2. [Second priority]
3. [Nice to have]

### Verdict: APPROVE / NEGOTIATE / REJECT
```
'''
p = doc.add_paragraph()
run = p.add_run(skill14)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_page_break()

# --- SKILL 15 ---
doc.add_heading('15. Executive Meeting Prep & Briefing', level=2)
p = doc.add_paragraph()
p.add_run('Source: ').bold = True
p.add_run('Custom — combining calendar skills, email context, and executive advisory patterns')
p.add_run('\n')
p.add_run('Why you need it: ').bold = True
p.add_run('Generates pre-meeting briefings by pulling calendar context, recent email threads with attendees, relevant documents, and talking points. Includes post-meeting action item capture.')

doc.add_paragraph('Skill Content:')
skill15 = '''---
name: meeting-prep
description: >
  Prepare executive briefings for upcoming meetings. Use when
  asked to prep for a meeting, create a briefing, review upcoming
  calendar events, or generate talking points. Pulls context from
  calendar, email, and documents to create comprehensive prep materials.
---

# Executive Meeting Prep & Briefing

## Pre-Meeting Briefing Workflow

### Step 1: Calendar Context
```bash
gog calendar show [event-id]  # Get meeting details
gog calendar today --format detailed  # Today's full schedule
```
Extract: Title, attendees, location/link, agenda (if attached), duration

### Step 2: Attendee Intelligence
For each external attendee:
- Company, title, LinkedIn summary
- Recent news about their company
- History of interactions (search email for past threads)
- Any open action items with them

For each internal attendee:
- Role and reporting line
- Recent email threads relevant to meeting topic
- Any pending approvals or decisions involving them

### Step 3: Topic Preparation
Based on meeting title/agenda:
- Pull relevant recent documents
- Summarize current status of discussed projects
- Identify open issues and potential questions
- Prepare data points and metrics they might ask about

### Step 4: Generate Briefing
```
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📋 MEETING BRIEFING
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📅 [Day, Date] | ⏰ [Time] | ⏱️ [Duration]
📍 [Location/Link]
📌 [Meeting Title]

👥 ATTENDEES
- [Name] — [Title, Company] — [Context note]
- [Name] — [Title, Company] — [Context note]

🎯 LIKELY AGENDA & YOUR POSITION
1. [Topic] — [Your recommended stance/talking point]
2. [Topic] — [Key data point to reference]
3. [Topic] — [Decision needed: your recommendation]

📊 KEY METRICS TO HAVE READY
- [Metric]: [Current value] ([trend])
- [Metric]: [Current value] ([trend])

⚠️ WATCH OUTS
- [Potential sensitive topic and how to handle]
- [Question they might ask and suggested answer]

📎 REFERENCE DOCUMENTS
- [Doc name] — [Where to find it]

✅ YOUR ASKS (what you need from this meeting)
1. [Decision or approval needed]
2. [Information to gather]
```

### Step 5: Post-Meeting Actions
After meeting, prompt for:
- Key decisions made
- Action items (who, what, by when)
- Follow-up emails to send
- Calendar holds to create
- Documents to update
'''
p = doc.add_paragraph()
run = p.add_run(skill15)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_page_break()

# ==========================================================
# INSTALLATION GUIDE
# ==========================================================
doc.add_heading('Quick Installation Guide', level=1)

doc.add_paragraph('Two ways to install any skill from this document:')

doc.add_heading('Option A: Skill Creator (Easiest)', level=2)
steps = [
    'Open Claude Desktop → Cowork tab',
    'Type /skill-creator',
    'Paste the skill content from this document',
    'Claude walks you through naming and testing',
    'Skill auto-saves to ~/.claude/skills/',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Option B: Manual Install', level=2)
steps = [
    'Create folder: ~/.claude/skills/[skill-name]/',
    'Create SKILL.md file inside with the skill content',
    'Add any supporting files (scripts, examples, references)',
    'Restart Claude Desktop',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Option C: From GitHub Repos', level=2)
repos = [
    ('Anthropic Official', 'github.com/anthropics/skills', 'Pre-installed with CoWork'),
    ('VoltAgent Awesome', 'github.com/VoltAgent/awesome-agent-skills', '6,532 stars, curated collection'),
    ('alirezarezvani', 'github.com/alirezarezvani/claude-skills', '2,300 stars, 169 skills, C-level advisory'),
    ('K-Dense-AI Scientific', 'github.com/K-Dense-AI/claude-scientific-skills', '8,241 stars, 140 data/finance skills'),
]
for name, url, notes in repos:
    p = doc.add_paragraph()
    run = p.add_run(f'• {name}')
    run.bold = True
    p.add_run(f' — {url} — {notes}')

doc.add_paragraph()
doc.add_heading('Recommended Install Order for CFO/COO', level=2)
order = [
    'CFO Email Triage (#5) — Daily inbox management',
    'Excel XLSX (#6) — Already installed, verify it works',
    'PowerPoint PPTX (#7) — Already installed',
    'Board Deck Review (#13) — Next board meeting prep',
    'Financial Modeling (#10) — Valuation and forecasting',
    'Competitive Intelligence (#12) — Ongoing market monitoring',
    'Contract Review (#14) — Next vendor negotiation',
    'Meeting Prep (#15) — Daily meeting preparation',
    'Cost-Benefit Analysis (#11) — Next capex/investment decision',
    'CEO Advisor (#9) — Strategic planning sessions',
]
for i, item in enumerate(order, 1):
    doc.add_paragraph(f'{i}. {item}')

# Save
output_path = '/Users/ericbrown/.openclaw/workspace/Top_15_CoWork_Skills_CFO.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
