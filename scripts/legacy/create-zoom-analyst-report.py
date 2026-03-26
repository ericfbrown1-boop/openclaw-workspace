#!/usr/bin/env python3
"""Create Zoom Financial Analyst Report — sell-side analyst framework."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_table_row(table, row_idx, values, bold_first=True):
    for col_idx, val in enumerate(values):
        cell = table.cell(row_idx, col_idx)
        cell.text = str(val)
        if bold_first and col_idx == 0:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

# Title
title = doc.add_heading('Zoom Communications, Inc. (NASDAQ: ZM)', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Sell-Side Equity Research — Financial Analysis')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x24, 0x40, 0xB3)
run.bold = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Q4 & Full Year FY2026 Results (Fiscal Year Ended January 31, 2026)\n').bold = True
meta.add_run(f'Report Date: {datetime.date.today().strftime("%B %d, %Y")}\n')
meta.add_run('Prepared for: Eric Brown, CFO & COO, Cohesity\n')
meta.add_run('Analyst POV: Expert Sell-Side Coverage\n')
meta.add_run('Classification: Confidential')

doc.add_page_break()

# ==========================================================
# COMPANY SNAPSHOT
# ==========================================================
doc.add_heading('Company Snapshot', level=1)

snap_table = doc.add_table(rows=10, cols=2)
snap_table.style = 'Light Grid Accent 1'
snap_data = [
    ('Ticker / Exchange', 'ZM / NASDAQ'),
    ('Sector', 'Technology — Application Software (UCaaS / CCaaS)'),
    ('Fiscal Year End', 'January 31'),
    ('Reporting Period', 'Q4 FY2026 (quarter ended January 31, 2026)'),
    ('Stock Price (Mar 6, 2026)', '~$78.05'),
    ('Market Cap', '~$24B'),
    ('52-Week Range', '$58 — $96.22'),
    ('Consensus Rating', 'Hold (47% Hold, 29% Strong Buy, 18% Buy)'),
    ('Consensus Price Target', '$95.32 (22% upside)'),
]
snap_table.cell(0, 0).text = 'Attribute'
snap_table.cell(0, 1).text = 'Detail'
for i, (k, v) in enumerate(snap_data, 1):
    snap_table.cell(i, 0).text = k
    snap_table.cell(i, 1).text = v

doc.add_page_break()

# ==========================================================
# QUARTERLY REVENUE & YoY GROWTH
# ==========================================================
doc.add_heading('Revenue Performance: Quarterly Trend & YoY Growth', level=1)

doc.add_paragraph(
    'Zoom\'s fiscal year ends January 31. The table below shows trailing quarterly revenue '
    'with YoY growth rates, emphasizing the revenue re-acceleration trend through FY2026.'
)

# Revenue trend table
rev_table = doc.add_table(rows=10, cols=5)
rev_table.style = 'Light Grid Accent 1'
rev_headers = ['Quarter', 'Total Revenue', 'YoY Growth', 'Enterprise Rev', 'Online Rev']
for i, h in enumerate(rev_headers):
    rev_table.cell(0, i).text = h

# Data: FY2025 Q1-Q4 and FY2026 Q1-Q4
# FY2025 (ended Jan 31, 2025): Total $4,666.3M
# FY2026 (ended Jan 31, 2026): Total $4,868.8M
rev_rows = [
    ('Q1 FY2025 (Apr 2024)', '$1,141.2M', '+3.2%', '$676.2M', '$465.0M'),
    ('Q2 FY2025 (Jul 2024)', '$1,162.5M', '+2.1%', '$689.3M', '$473.2M'),
    ('Q3 FY2025 (Oct 2024)', '$1,178.7M', '+3.6%', '$698.2M', '$480.5M'),
    ('Q4 FY2025 (Jan 2025)', '$1,183.9M', '+2.7%', '$707.1M', '$476.8M'),
    ('Q1 FY2026 (Apr 2025)', '$1,174.3M', '+2.9%', '$702.8M', '$471.5M'),
    ('Q2 FY2026 (Jul 2025)', '$1,218.1M', '+4.8%', '$725.7M', '$492.4M'),
    ('Q3 FY2026 (Oct 2025)', '$1,229.4M', '+4.3%', '$748.3M', '$481.1M'),
    ('Q4 FY2026 (Jan 2026)', '$1,247.0M', '+5.3%', '$757.3M', '$489.7M'),
    ('FY2026 Full Year', '$4,868.8M', '+4.4%', '$2,934.1M', '$1,934.7M'),
]
for i, row in enumerate(rev_rows, 1):
    for j, val in enumerate(row):
        rev_table.cell(i, j).text = val

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Key Observations:').bold = True
observations = [
    'Revenue re-accelerated through FY2026: from +2.9% in Q1 to +5.3% in Q4 — the strongest quarterly growth in 2+ years.',
    'Enterprise revenue growth accelerated to +7.1% YoY in Q4 (from +6.5% FY2026 full year), driven by Contact Center and AI.',
    'Online revenue stabilized at +2.6% in Q4 (up from +1.2% full year), with monthly churn at historic low 2.7%.',
    'Full year enterprise revenue of $2,934.1M (+6.5%) now represents 60.3% of total — up from 59.1% in FY2025.',
    'Revenue growth acceleration of 130bps (FY2025 +3.1% → FY2026 +4.4%) breaks the post-pandemic deceleration trend.',
]
for o in observations:
    doc.add_paragraph(o, style='List Bullet')

doc.add_page_break()

# ==========================================================
# GUIDANCE vs ACTUALS — NEXT QUARTER ANALYSIS
# ==========================================================
doc.add_heading('Guidance Analysis: Q1 FY2027 & Full Year FY2027', level=1)

doc.add_paragraph(
    'Per USER.md framework: midpoint of next-quarter guidance, implied YoY growth, '
    'compared against the last five quarters of actuals.'
)

guide_table = doc.add_table(rows=9, cols=4)
guide_table.style = 'Light Grid Accent 1'
g_headers = ['Period', 'Revenue', 'YoY Growth', 'Non-GAAP EPS']
for i, h in enumerate(g_headers):
    guide_table.cell(0, i).text = h

g_rows = [
    ('Q4 FY2025 (Actual)', '$1,183.9M', '+2.7%', '$1.41'),
    ('Q1 FY2026 (Actual)', '$1,174.3M', '+2.9%', '$1.35'),
    ('Q2 FY2026 (Actual)', '$1,218.1M', '+4.8%', '$1.39'),
    ('Q3 FY2026 (Actual)', '$1,229.4M', '+4.3%', '$1.73*'),
    ('Q4 FY2026 (Actual)', '$1,247.0M', '+5.3%', '$1.44'),
    ('', '', '', ''),
    ('Q1 FY2027 (Guidance Mid)', '$1,222.5M', '+4.1%', '$1.41'),
    ('FY2027 (Guidance)', '>$5,060M', '~3.9%', '$5.79 (mid)'),
]
for i, row in enumerate(g_rows, 1):
    for j, val in enumerate(row):
        guide_table.cell(i, j).text = val

doc.add_paragraph('* Q3 FY2026 EPS included one-time items; normalized ~$1.48')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Guidance Assessment:').bold = True
guide_points = [
    'Q1 FY2027 revenue midpoint of $1,222.5M implies +4.1% YoY growth — in line with trailing 5-quarter average of +4.3%.',
    'Full year FY2027 guidance of >$5.06B implies ~3.9% growth — slight deceleration from FY2026\'s 4.4%. Conservative or realistic?',
    'EPS guidance of $5.77-$5.81 implies slight decline from FY2026 $5.92 — likely reflects increased AI/Contact Center investment.',
    'The $5B milestone is psychologically important but represents only modest acceleration at current run rates.',
    'Guidance does NOT embed any inflection from AI monetization — potential upside catalyst if Custom AI Companion scales.',
]
for g in guide_points:
    doc.add_paragraph(g, style='List Bullet')

doc.add_page_break()

# ==========================================================
# MARGIN TRENDS
# ==========================================================
doc.add_heading('Profitability & Margin Analysis', level=1)

doc.add_heading('Operating Margins — Quarterly Trend', level=2)

margin_table = doc.add_table(rows=6, cols=5)
margin_table.style = 'Light Grid Accent 1'
m_headers = ['Period', 'GAAP Gross Margin', 'Non-GAAP Gross Margin', 'GAAP Op Margin', 'Non-GAAP Op Margin']
for i, h in enumerate(m_headers):
    margin_table.cell(0, i).text = h

m_rows = [
    ('Q4 FY2025', '75.6%', '78.9%', '19.0%', '39.5%'),
    ('Q1 FY2026', '75.8%', '79.2%', '19.4%', '39.8%'),
    ('Q2 FY2026', '76.2%', '79.5%', '24.7%', '40.5%'),
    ('Q3 FY2026', '76.5%', '80.0%', '25.3%', '41.5%'),
    ('Q4 FY2026', '76.8%', '80.1%', '20.0%', '39.3%'),
]
for i, row in enumerate(m_rows, 1):
    for j, val in enumerate(row):
        margin_table.cell(i, j).text = val

doc.add_paragraph()

doc.add_heading('Full Year Margin Comparison', level=2)
fy_margin_table = doc.add_table(rows=7, cols=3)
fy_margin_table.style = 'Light Grid Accent 1'
fy_m_headers = ['Metric', 'FY2025', 'FY2026']
for i, h in enumerate(fy_m_headers):
    fy_margin_table.cell(0, i).text = h

fy_m_rows = [
    ('Non-GAAP Gross Margin', '78.9%', '79.7%'),
    ('GAAP Operating Margin', '17.4%', '23.1%'),
    ('Non-GAAP Operating Margin', '39.4%', '40.4%'),
    ('GAAP Net Margin', '21.6%', '39.0%'),
    ('Non-GAAP Net Income Margin', '~24.5%', '~24.7%'),
    ('FCF Margin', '~38.8%', '~39.5%'),
]
for i, row in enumerate(fy_m_rows, 1):
    for j, val in enumerate(row):
        fy_margin_table.cell(i, j).text = val

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Margin Analysis:').bold = True
margin_notes = [
    'Non-GAAP gross margin expanded 80bps to 79.7% — approaching 80% threshold, reflecting software-mix improvement.',
    'GAAP operating margin expanded 570bps to 23.1% — driven by cost discipline and reduced SBC as percentage of revenue.',
    'Non-GAAP operating margin of 40.4% (+100bps) is elite among enterprise SaaS companies at this scale.',
    'FCF margin of ~39.5% on $4.87B revenue = $1.92B free cash flow — exceptional cash generation.',
    'Q4 non-GAAP op margin of 39.3% dipped slightly (down 20bps YoY) — management investing in AI/Contact Center growth.',
    'GAAP net margin inflated by ~$970M unrealized gains on strategic investments — non-operational, non-cash.',
]
for n in margin_notes:
    doc.add_paragraph(n, style='List Bullet')

doc.add_page_break()

# ==========================================================
# CASH FLOW & BALANCE SHEET
# ==========================================================
doc.add_heading('Cash Flow & Capital Allocation', level=1)

cf_table = doc.add_table(rows=8, cols=3)
cf_table.style = 'Light Grid Accent 1'
cf_table.cell(0, 0).text = 'Metric'
cf_table.cell(0, 1).text = 'FY2025'
cf_table.cell(0, 2).text = 'FY2026'
cf_rows = [
    ('Operating Cash Flow', '$1,855.0M', '$1,989.4M'),
    ('Free Cash Flow', '$1,808.4M', '$1,924.1M'),
    ('FCF Margin', '~38.8%', '~39.5%'),
    ('Cash & Marketable Securities', '$7,790M', '$7,817M'),
    ('Shares Repurchased (FY)', '15.9M shares', '20.4M shares'),
    ('Buyback Spend (FY)', '~$1,050M', '~$1,665M'),
    ('Remaining Buyback Auth', '~$1.5B', '~$1.0B'),
]
for i, (m, fy25, fy26) in enumerate(cf_rows, 1):
    cf_table.cell(i, 0).text = m
    cf_table.cell(i, 1).text = fy25
    cf_table.cell(i, 2).text = fy26

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Capital Allocation Commentary:').bold = True
cap_notes = [
    'FCF grew 6.4% YoY to $1.92B despite only 4.4% revenue growth — operating leverage working.',
    'Cash position essentially flat at $7.8B DESPITE $2.7B in cumulative buybacks (FY2025+FY2026) — the business generates enough cash to fund buybacks AND maintain the fortress.',
    'Share count declining: 20.4M shares repurchased in FY2026 (vs. 15.9M in FY2025) = accelerating buybacks.',
    '$1.0B remaining buyback authorization — likely to be refreshed given strong FCF generation.',
    'No debt on balance sheet — entirely equity-financed. Unusual and highly conservative for a $5B revenue company.',
    'Strategic investment portfolio generating unrealized gains (~$970M in FY2026) but this is volatile and non-operational.',
]
for n in cap_notes:
    doc.add_paragraph(n, style='List Bullet')

doc.add_page_break()

# ==========================================================
# PRODUCT/USE-CASE CALLOUTS
# ==========================================================
doc.add_heading('Product & Use-Case Drivers', level=1)

doc.add_paragraph(
    'Management\'s Q4 FY2026 earnings call highlighted several product catalysts driving performance:'
)

drivers = [
    ('Zoom Contact Center (CCaaS)', 
     'Fastest-growing segment. ARR growing "high double digits." 1,100+ customers, doubled YoY. '
     '7 of top 10 Q4 deals were competitive displacements of legacy vendors (NICE, Genesys, etc.). '
     'This is now the primary AI monetization vehicle.'),
    ('Zoom Phone', 
     'Surpassed 10M paid seats (milestone reached Oct 2025). 6 of 10 largest Contact Center deals '
     'also included Phone — proving the platform cross-sell. Displaced Microsoft Teams and Cisco '
     'calling at two major US financial institutions in Q4.'),
    ('AI Companion & Custom AI Companion', 
     'Base AI Companion adoption surged 4x YoY; MAU tripled YoY in Q4. Custom AI Companion '
     '($12/user/month) — paid tier for enterprise workflow integration. Contributed to "significant '
     'deal closures" in Q3 and Q4. Prices at 40% of Microsoft Copilot ($30).'),
    ('Zoom Virtual Agent 3.0', 
     'AI-powered automated customer service. Included in 4 of top 10 CX deals. Nearly 7-figure ARR '
     'deal with leading US retailer (1,100+ locations). This is the "digital labor" play — automating '
     'call center agents.'),
    ('Workvivo (Employee Experience)', 
     '1,225+ customers (~70-80% YoY growth). Three $1M+ ARR deals signed. $100K+ ARR customers '
     'grew 140% YoY. Benefiting from Meta sunsetting Workplace — Workvivo is the preferred migration path.'),
    ('Net Dollar Expansion Rate', 
     'TTM enterprise NDE of 98% — below 100%, indicating some existing customer contraction. '
     'This is a yellow flag: while new products are growing, some existing enterprise accounts are '
     'shrinking or consolidating licenses. Management needs to get this above 100%.'),
]

for title_text, desc in drivers:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ==========================================================
# VALUATION
# ==========================================================
doc.add_heading('Valuation Summary', level=1)

doc.add_heading('Multiples Analysis', level=2)

val_table = doc.add_table(rows=7, cols=3)
val_table.style = 'Light Grid Accent 1'
val_table.cell(0, 0).text = 'Metric'
val_table.cell(0, 1).text = 'Current'
val_table.cell(0, 2).text = 'Peer Median (UCaaS/SaaS)'

val_rows = [
    ('EV/Revenue (NTM)', '~3.3x', '6-8x'),
    ('EV/Revenue (LTM)', '~3.5x', '7-10x'),
    ('EV/Non-GAAP Operating Income', '~8.6x', '20-25x'),
    ('EV/FCF', '~8.8x', '20-30x'),
    ('P/E (Non-GAAP, NTM)', '~13.5x', '25-35x'),
    ('Rule of 40', '~44', '35-50 (median ~40)'),
]
for i, (m, curr, peer) in enumerate(val_rows, 1):
    val_table.cell(i, 0).text = m
    val_table.cell(i, 1).text = curr
    val_table.cell(i, 2).text = peer

doc.add_paragraph()

doc.add_heading('Peer Comparison — Revenue Multiples', level=2)
peer_table = doc.add_table(rows=7, cols=5)
peer_table.style = 'Light Grid Accent 1'
peer_headers = ['Company', 'LTM Revenue', 'Rev Growth', 'FCF Margin', 'EV/Rev (NTM)']
for i, h in enumerate(peer_headers):
    peer_table.cell(0, i).text = h

peer_data = [
    ('Zoom (ZM)', '$4.87B', '+4.4%', '~39.5%', '~3.3x'),
    ('RingCentral (RNG)', '$2.4B', '+9%', '~25%', '~3.5x'),
    ('Twilio (TWLO)', '$4.5B', '+7%', '~15%', '~3.0x'),
    ('HubSpot (HUBS)', '$2.7B', '+20%', '~22%', '~12x'),
    ('ServiceNow (NOW)', '$11B', '+22%', '~32%', '~16x'),
    ('Salesforce (CRM)', '$38B', '+8%', '~30%', '~7x'),
]
for i, row in enumerate(peer_data, 1):
    for j, val in enumerate(row):
        peer_table.cell(i, j).text = val

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Valuation Assessment:').bold = True
val_notes = [
    'Zoom trades at a significant discount to SaaS peers on every metric — 3.3x EV/Revenue vs. median 6-8x.',
    'The discount reflects the "low growth" narrative (4.4% vs. peer median ~10-15%). Market is pricing Zoom as ex-growth.',
    'However, the FCF yield is exceptional: ~12% FCF yield at current price vs. 3-5% for most SaaS peers.',
    'Rule of 40 score of ~44 (4.4% growth + ~39.5% FCF margin) is ABOVE peer median — growth + profitability is good, market just doesn\'t reward it.',
    'The disconnect: market penalizes low revenue growth despite exceptional cash generation. A re-rating to even 5x revenue = ~$80B EV = ~$50/share upside from current levels.',
    'Catalysts for re-rating: AI-driven revenue acceleration above 6-8%, Contact Center achieving $500M+ ARR, net dollar expansion returning above 100%.',
]
for n in val_notes:
    doc.add_paragraph(n, style='List Bullet')

doc.add_page_break()

# ==========================================================
# ANALYST SENTIMENT
# ==========================================================
doc.add_heading('Analyst Sentiment & Price Targets', level=1)

doc.add_heading('Consensus Ratings (17 Analysts)', level=2)
rating_table = doc.add_table(rows=6, cols=3)
rating_table.style = 'Light Grid Accent 1'
rating_table.cell(0, 0).text = 'Rating'
rating_table.cell(0, 1).text = 'Count'
rating_table.cell(0, 2).text = '%'
r_data = [
    ('Strong Buy', '5', '29%'),
    ('Buy', '3', '18%'),
    ('Hold', '8', '47%'),
    ('Sell', '0', '0%'),
    ('Strong Sell', '1', '6%'),
]
for i, (r, c, pct) in enumerate(r_data, 1):
    rating_table.cell(i, 0).text = r
    rating_table.cell(i, 1).text = c
    rating_table.cell(i, 2).text = pct

doc.add_paragraph()
doc.add_heading('Notable Post-Earnings Analyst Actions', level=2)
actions = [
    ('BTIG', 'Buy', '$100', 'Cut from $115. Notes AI monetization thesis intact but needs more time to prove out.'),
    ('KeyBanc', 'Underweight', '$74', 'Raised from $69. Still skeptical on growth re-acceleration despite improving metrics.'),
    ('Goldman Sachs', 'Neutral', '$92', 'Notes strong FCF but limited near-term catalysts for multiple expansion.'),
    ('Morgan Stanley', 'Overweight', '$105', 'Bullish on CCaaS displacement and AI Companion monetization trajectory.'),
]
for firm, rating, target, comment in actions:
    p = doc.add_paragraph()
    run = p.add_run(f'{firm} ({rating}, PT ${target}): ')
    run.bold = True
    p.add_run(comment)

doc.add_page_break()

# ==========================================================
# INVESTMENT THESIS
# ==========================================================
doc.add_heading('Investment Thesis Summary', level=1)

doc.add_heading('Bull Case ($110+)', level=2)
bulls = [
    'AI Companion + Custom AI drives revenue acceleration to 7-10% in FY2028',
    'Contact Center becomes a $1B ARR business within 2 years (currently high-double-digit growth)',
    'Workvivo captures significant Workplace from Meta migration wave',
    'Share buybacks + FCF compound EPS growth to 15%+ even on modest revenue growth',
    'Re-rating from 3.3x to 5x revenue = $80/share EV upside',
]
for b in bulls:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Bear Case ($55-60)', level=2)
bears = [
    'Revenue growth stalls at 3-4% as Microsoft Teams bundling continues to erode share',
    'AI Companion commoditized — every platform offers free meeting summaries',
    'Net dollar expansion stays below 100% — existing enterprise base contracting',
    'Online segment declines as SMBs consolidate to bundled solutions',
    'Multiple compression if market rotates away from low-growth software',
]
for b in bears:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Base Case ($85-95)', level=2)
doc.add_paragraph(
    'Revenue grows 4-5% through FY2028, Contact Center contributes meaningful incremental growth, '
    'FCF remains $2B+, buybacks compress share count 3-5% annually, EPS grows 8-10% on operating '
    'leverage. Stock re-rates modestly as AI monetization evidence accumulates. 12-month target: $90-95.'
)

doc.add_page_break()

# ==========================================================
# RISKS
# ==========================================================
doc.add_heading('Key Risks', level=1)
risks = [
    ('Microsoft Bundling', 'Teams is included with M365 E3/E5 — the "free" competitor problem. '
     'Enterprises standardizing on Microsoft stack may not pay separately for Zoom.'),
    ('Revenue Growth Ceiling', 'If organic growth can\'t break above 5-6%, the stock stays cheap. '
     'The market rewards growth, not just profitability.'),
    ('NDE Below 100%', 'Net dollar expansion at 98% means existing enterprise customers are net-shrinking. '
     'Must reverse this to support the platform narrative.'),
    ('Dual-Class Control', 'Eric Yuan\'s 10x voting shares mean limited shareholder recourse if strategy falters.'),
    ('AI Hype Risk', 'If Custom AI Companion adoption is slower than expected, the growth re-acceleration thesis fails.'),
    ('Strategic Investment Volatility', 'The $970M unrealized gain could reverse in a market downturn, '
     'creating misleading GAAP earnings volatility.'),
]
for title_text, desc in risks:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— End of Report —')
run.italic = True
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Disclaimer: This report is prepared for internal use by Cohesity management. '
    'It is not investment advice. All data sourced from public filings, earnings calls, '
    'and analyst reports as of March 6, 2026.').font.size = Pt(8)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Prepared by Jarvis — AI Assistant to Eric Brown').font.size = Pt(9)

# Save
output_path = '/Users/ericbrown/ProjectScraper/Zoom_Financial_Analysis_FY2026.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
