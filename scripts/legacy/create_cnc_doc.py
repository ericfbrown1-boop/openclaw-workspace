#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Home CNC Woodworking Solutions', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Custom Tables & Game Boards', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Spacing

# Section 1: Top Recommendations by Budget
doc.add_heading('1. Top Recommendations by Budget', level=1)

doc.add_heading('Entry-Level ($2,000-$4,000)', level=2)
p = doc.add_paragraph()
p.add_run('Shapeoko 5 Pro').bold = True
doc.add_paragraph('Excellent beginner machine with strong community support', style='List Bullet')
doc.add_paragraph('Solid construction and reliable performance', style='List Bullet')
doc.add_paragraph('Great for learning CNC basics', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('X-Carve Pro').bold = True
doc.add_paragraph('User-friendly design, ideal for hobbyists', style='List Bullet')
doc.add_paragraph('Good customer support and documentation', style='List Bullet')
doc.add_paragraph('Expandable work area options', style='List Bullet')

doc.add_paragraph()

doc.add_heading('Mid-Range ($4,000-$8,000)', level=2)
p = doc.add_paragraph()
p.add_run('Onefinity Woodworker X-50').bold = True
doc.add_paragraph('50" x 50" work area - perfect for most custom tables and large game boards', style='List Bullet')
doc.add_paragraph('Extremely rigid construction for precision', style='List Bullet')
doc.add_paragraph('Built specifically for woodworking', style='List Bullet')
doc.add_paragraph('Excellent community and support', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Avid CNC').bold = True
doc.add_paragraph('Professional-grade quality', style='List Bullet')
doc.add_paragraph('Highly customizable configurations', style='List Bullet')
doc.add_paragraph('Outstanding precision and repeatability', style='List Bullet')

doc.add_paragraph()

doc.add_heading('Professional ($8,000-$15,000+)', level=2)
p = doc.add_paragraph()
p.add_run('Full 4x8 CNC Routers').bold = True
doc.add_paragraph('Complete table-sized sheets in one setup', style='List Bullet')
doc.add_paragraph('Production-level capabilities', style='List Bullet')
doc.add_paragraph('Ideal for commercial work', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('ShopBot').bold = True
doc.add_paragraph('Industry standard for professional woodworking', style='List Bullet')
doc.add_paragraph('Extensive training and support resources', style='List Bullet')
doc.add_paragraph('Proven reliability for commercial use', style='List Bullet')

doc.add_paragraph()

# Section 2: Project-Specific Recommendations
doc.add_heading('2. Project-Specific Recommendations', level=1)

doc.add_heading('For Game Boards (Chess, Backgammon, Cribbage)', level=2)
doc.add_paragraph('Work area: 24" x 24" minimum (larger is better for flexibility)', style='List Bullet')
doc.add_paragraph('Fine detail capability: 0.001-0.005" precision', style='List Bullet')
doc.add_paragraph('Ability to handle inlays and intricate patterns', style='List Bullet')
doc.add_paragraph('Recommended machines: Shapeoko 5 Pro, X-Carve Pro, Onefinity', style='List Bullet')

doc.add_paragraph()

doc.add_heading('For Custom Tables', level=2)
doc.add_paragraph('Work area: At least 48" x 48" (Onefinity X-50 ideal)', style='List Bullet')
doc.add_paragraph('For full table tops: 4x8 router recommended', style='List Bullet')
doc.add_paragraph('Heavy-duty construction for cutting hardwoods', style='List Bullet')
doc.add_paragraph('Powerful spindle (2+ HP) for efficient material removal', style='List Bullet')

doc.add_paragraph()

# Section 3: Key Features to Look For
doc.add_heading('3. Key Features to Look For', level=1)

p = doc.add_paragraph()
p.add_run('Work Area').bold = True
doc.add_paragraph('Consider your largest project + some margin', style='List Bullet')
doc.add_paragraph('Bigger isn\'t always better - larger machines need more space and maintenance', style='List Bullet')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Spindle').bold = True
doc.add_paragraph('1.5-2.5 HP for hobby work', style='List Bullet')
doc.add_paragraph('2.5+ HP for professional/commercial use', style='List Bullet')
doc.add_paragraph('Water-cooled vs air-cooled options', style='List Bullet')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Precision & Rigidity').bold = True
doc.add_paragraph('Look for heavy-duty linear rails', style='List Bullet')
doc.add_paragraph('Minimal flex in the gantry', style='List Bullet')
doc.add_paragraph('Quality ball screws or rack & pinion drive', style='List Bullet')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Software Compatibility').bold = True
doc.add_paragraph('Works with standard G-code', style='List Bullet')
doc.add_paragraph('Compatible with common CAM software (Fusion 360, VCarve, etc.)', style='List Bullet')
doc.add_paragraph('User-friendly control software', style='List Bullet')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Dust Collection').bold = True
doc.add_paragraph('Integrated dust shoe or mounting options', style='List Bullet')
doc.add_paragraph('Effective chip evacuation is essential for wood', style='List Bullet')

doc.add_paragraph()

# Section 4: What You'll Need
doc.add_heading('4. What You\'ll Need Beyond the Machine', level=1)

doc.add_heading('CAD/CAM Software', level=2)
doc.add_paragraph('Design: Fusion 360 (free for hobbyists), Inkscape, Adobe Illustrator', style='List Bullet')
doc.add_paragraph('CAM: VCarve (popular for woodworking), Fusion 360, Carbide Create', style='List Bullet')

doc.add_paragraph()

doc.add_heading('Tooling', level=2)
doc.add_paragraph('End mills: 1/4", 1/8", 1/16" for various detail levels', style='List Bullet')
doc.add_paragraph('V-bits for engraving and v-carving', style='List Bullet')
doc.add_paragraph('Ball nose bits for 3D carving', style='List Bullet')
doc.add_paragraph('Budget $200-500 for a good starter set', style='List Bullet')

doc.add_paragraph()

doc.add_heading('Dust Collection', level=2)
doc.add_paragraph('Shop vacuum (minimum) or dedicated dust collector', style='List Bullet')
doc.add_paragraph('CNC produces a LOT of fine dust', style='List Bullet')

doc.add_paragraph()

doc.add_heading('Materials', level=2)
doc.add_paragraph('Hardwoods: walnut, maple, cherry for game boards', style='List Bullet')
doc.add_paragraph('Plywood: Baltic birch for practice and prototypes', style='List Bullet')
doc.add_paragraph('MDF: Great for templates and jigs', style='List Bullet')

doc.add_paragraph()

# Section 5: Learning Resources
doc.add_heading('5. Learning Resources', level=1)

doc.add_heading('YouTube Channels', level=2)
doc.add_paragraph('Winston Moy - Excellent CNC woodworking tutorials', style='List Bullet')
doc.add_paragraph('Maker Tales - Game board and CNC projects', style='List Bullet')
doc.add_paragraph('NYC CNC - Professional techniques and tips', style='List Bullet')

doc.add_paragraph()

doc.add_heading('Forums & Communities', level=2)
doc.add_paragraph('CNCZone.com - Long-established community', style='List Bullet')
doc.add_paragraph('Reddit r/hobbycnc - Active hobbyist community', style='List Bullet')
doc.add_paragraph('Machine-specific forums (Shapeoko, Onefinity, etc.)', style='List Bullet')

doc.add_paragraph()

doc.add_heading('Tutorials', level=2)
doc.add_paragraph('Start with simple projects: coasters, signs, cutting boards', style='List Bullet')
doc.add_paragraph('Progress to inlays and more complex joinery', style='List Bullet')
doc.add_paragraph('VCarve has excellent tutorial videos', style='List Bullet')

doc.add_paragraph()

# Section 6: Top Pick Recommendation
doc.add_heading('6. Top Pick Recommendation', level=1)

recommendation = doc.add_heading('Onefinity Woodworker X-50', level=2)

highlight = doc.add_paragraph()
highlight_run = highlight.add_run('Best Overall Balance for Your Projects')
highlight_run.bold = True
highlight_run.font.size = Pt(12)
highlight_run.font.color.rgb = RGBColor(0, 102, 204)

doc.add_paragraph()

doc.add_heading('Why This Machine?', level=3)
doc.add_paragraph('50" x 50" work area handles most custom table components and large game boards', style='List Bullet')
doc.add_paragraph('Professional-grade rigidity at a mid-range price', style='List Bullet')
doc.add_paragraph('Purpose-built for woodworking (not a converted metal machine)', style='List Bullet')
doc.add_paragraph('Strong community support and extensive documentation', style='List Bullet')
doc.add_paragraph('Room to grow without immediately needing to upgrade', style='List Bullet')

doc.add_paragraph()

doc.add_heading('Price Range', level=3)
price = doc.add_paragraph('$5,000-$6,500 (depending on configuration and spindle choice)')
price.runs[0].bold = True

doc.add_paragraph()

doc.add_heading('Perfect For', level=3)
doc.add_paragraph('Game board production with intricate inlays', style='List Bullet')
doc.add_paragraph('Custom table components and decorative elements', style='List Bullet')
doc.add_paragraph('Growing from hobbyist to semi-professional work', style='List Bullet')
doc.add_paragraph('Anyone serious about CNC woodworking but not ready for industrial scale', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph()

# Footer
footer = doc.add_paragraph('___________________________________________')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_text = doc.add_paragraph('This guide provides recommendations based on your interest in creating custom tables and game boards.')
footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_text.runs[0].font.size = Pt(9)
footer_text.runs[0].font.italic = True

# Save document
output_path = os.path.expanduser('~/.openclaw/workspace/CNC_Woodworking_Recommendations.docx')
doc.save(output_path)
print(f"Document saved to: {output_path}")
