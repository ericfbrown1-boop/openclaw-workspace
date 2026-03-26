from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

# Create a new Document
doc = Document()

# Title page
title = doc.add_heading('Your Personalized Wellness Plan', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph("Based on Gary Brecka's Health Philosophy")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(16)
subtitle_format.font.color.rgb = RGBColor(64, 64, 64)

doc.add_paragraph()

# Client profile section
profile_heading = doc.add_heading('Client Profile', level=1)
profile_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

profile_text = '''
Name: Eric Brown
Age: 60 years old
Weight: Just under 190 pounds
Current Exercise: Rowing machine, 30 minutes daily
Primary Goals:
• Lower carbohydrate intake
• Optimize eating/sleep intervals (intermittent fasting & time-restricted eating)
• Enhance overall health and longevity
'''
doc.add_paragraph(profile_text)

# Executive Summary
exec_heading = doc.add_heading('Executive Summary', level=1)
exec_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

exec_summary = '''This personalized wellness plan integrates Gary Brecka's evidence-based health philosophy with your specific goals and lifestyle. Gary Brecka is a renowned human biologist and founder of the 10X Health System, known for his data-driven approach to optimal health through personalized nutrition, targeted supplementation, and lifestyle optimization.

This plan focuses on four core pillars:
1. Low-carbohydrate, nutrient-dense nutrition
2. Time-restricted eating for metabolic optimization
3. Sleep optimization for cellular repair and longevity
4. Strategic integration of daily rowing with morning protocols

The approach emphasizes sustainable, long-term wellness rather than quick fixes, supporting cellular health, reducing inflammation, and promoting longevity.'''

doc.add_paragraph(exec_summary)

# Gary Brecka's Philosophy
philosophy_heading = doc.add_heading("Gary Brecka's Core Health Philosophy", level=1)
philosophy_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

philosophy_text = '''Gary Brecka's approach is rooted in several fundamental principles:

1. PERSONALIZED BIOLOGY: Health optimization starts with understanding your unique biochemistry and genetics, not one-size-fits-all solutions.

2. CELLULAR HEALTH: Optimal health begins at the cellular level. Proper nutrition, supplementation, and lifestyle habits ensure cells receive what they need to function optimally.

3. METHYLATION SUPPORT: Many people have genetic mutations (like MTHFR) that affect how the body processes nutrients. Using methylated vitamins with high bioavailability ensures proper absorption.

4. NUTRIENT DENSITY OVER CALORIES: Focus on foods packed with essential nutrients rather than calorie counting or restrictive dieting.

5. INFLAMMATION REDUCTION: Chronic inflammation is the root cause of many age-related diseases. Diet and lifestyle should minimize inflammatory triggers.

6. SUSTAINABLE HABITS: Small, consistent daily practices compound over time to create dramatic health transformations.

7. DATA-DRIVEN DECISIONS: Use biomarkers, blood work, and measurable results to guide health decisions, not trends.'''

doc.add_paragraph(philosophy_text)

# Add page break
doc.add_page_break()

# Nutrition Plan
nutrition_heading = doc.add_heading('Part 1: Personalized Nutrition Plan', level=1)
nutrition_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

doc.add_heading('Low-Carb Dietary Framework', level=2)

nutrition_text = '''Following Gary Brecka's philosophy, your nutrition plan emphasizes:

• HIGH-QUALITY PROTEINS (30-35% of daily intake)
  - Grass-fed beef, wild-caught salmon, organic chicken
  - Pasture-raised eggs
  - Protein intake: 120-140g daily (0.6-0.7g per pound of body weight)

• HEALTHY FATS (50-60% of daily intake)
  - Avocados and extra virgin olive oil
  - Nuts and seeds (almonds, walnuts, chia, flax)
  - Fatty fish rich in omega-3s (salmon, sardines, mackerel)
  - Grass-fed butter or ghee

• NUTRIENT-DENSE VEGETABLES (unlimited)
  - Dark leafy greens: spinach, kale, Swiss chard, arugula
  - Cruciferous vegetables: broccoli, cauliflower, Brussels sprouts
  - Colorful vegetables: bell peppers, carrots, beets
  - Alliums: garlic, onions, leeks (anti-inflammatory properties)

• LIMITED CARBOHYDRATES (10-15% of daily intake)
  - Focus on fiber-rich sources: 50-75g net carbs daily
  - Quinoa, brown rice (small portions)
  - Sweet potatoes, root vegetables (moderate amounts)
  - Berries for antioxidants

FOODS TO ELIMINATE:
- Processed foods and refined sugars
- White bread, pasta, refined grains
- Artificial sweeteners and additives
- Trans fats and vegetable oils
- Excessive alcohol'''

doc.add_paragraph(nutrition_text)

# The 30/30/30 Method
doc.add_heading('The 30/30/30 Morning Method', level=2)

method_text = '''Gary Brecka's signature 30/30/30 method is perfect for integrating with your rowing routine:

Within 30 minutes of waking:
1. Consume 30 grams of high-quality protein
2. Follow with 30 minutes of low-intensity exercise (your rowing!)

WHY IT WORKS:
• Stabilizes blood sugar levels immediately upon waking
• Boosts metabolism and promotes fat burning throughout the day
• Prevents muscle catabolism (breakdown)
• Reduces cravings and hunger later in the day
• Supports hormone balance, particularly insulin and cortisol

IMPLEMENTATION FOR YOU:
• Wake up and consume a protein-rich breakfast within 30 minutes
• Examples: 3-egg omelet with vegetables, protein smoothie with spinach and berries, Greek yogurt with nuts and seeds
• Then complete your 30-minute rowing session
• Keep heart rate at steady-state (approximately 60-70% max heart rate, around 100-120 bpm for age 60)'''

doc.add_paragraph(method_text)

# Add page break
doc.add_page_break()

# Time-Restricted Eating
eating_heading = doc.add_heading('Part 2: Time-Restricted Eating Protocol', level=1)
eating_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

eating_text = '''Gary Brecka strongly advocates for time-restricted eating to optimize cellular autophagy (cellular cleanup and repair) and metabolic health.

RECOMMENDED EATING WINDOW: 8-Hour Window (16:8 Protocol)

Eating Window: 10:00 AM - 6:00 PM
Fasting Period: 6:00 PM - 10:00 AM (next day)

WHY THIS SCHEDULE:
• Aligns with circadian biology for optimal digestion
• Allows 14-16 hours overnight for cellular repair and autophagy
• Stops eating 3-4 hours before bedtime for better sleep quality
• Maximizes fat burning during fasted state
• Reduces inflammation and oxidative stress

DAILY MEAL STRUCTURE:

10:00 AM - MEAL 1 (Post-Workout)
• 30g+ protein-rich meal following your rowing session
• Examples:
  - Vegetable-packed omelet (3 eggs) with avocado slices
  - Protein smoothie: spinach, berries, protein powder, almond milk, chia seeds
  - Grilled salmon with sautéed greens and olive oil

2:00 PM - MEAL 2 (Midday)
• Balanced, nutrient-dense meal
• Examples:
  - Grilled chicken salad with mixed greens, cherry tomatoes, cucumber, olive oil dressing
  - Mediterranean quinoa bowl with vegetables and lean turkey
  - Beef and broccoli stir-fry with cauliflower rice

5:30 PM - MEAL 3 (Early Dinner)
• Complete dinner before 6:00 PM
• Examples:
  - Herb-crusted baked salmon with roasted Brussels sprouts and sweet potato
  - Grass-fed beef with roasted cauliflower and mixed green salad
  - Lemon herb chicken with steamed vegetables

HYDRATION DURING FASTING:
• Water with pinch of Celtic sea salt or Himalayan pink salt
• Black coffee or green tea (no cream or sugar)
• Herbal teas
• Electrolyte water (unsweetened)

IMPORTANT NOTES:
• First 2-3 weeks may require adjustment
• Listen to your body - flexibility is key
• Stay well-hydrated throughout fasting period
• Break fast gently if you experience discomfort'''

doc.add_paragraph(eating_text)

# Sample Weekly Meal Plan
doc.add_heading('7-Day Sample Meal Plan', level=2)

meal_plan = '''
DAY 1-3:
Meal 1 (10 AM): Spinach & mushroom omelet with avocado slices
Meal 2 (2 PM): Grilled chicken salad with leafy greens, cherry tomatoes, cucumber, olive oil dressing
Meal 3 (5:30 PM): Herb-crusted baked salmon with roasted Brussels sprouts and sweet potatoes

DAY 4-6:
Meal 1 (10 AM): Protein smoothie with spinach, frozen berries, almond milk, protein powder, chia seeds
Meal 2 (2 PM): Mediterranean quinoa veggie bowl with mixed vegetables and lean turkey
Meal 3 (5:30 PM): Spicy beef & broccoli stir-fry with cauliflower rice

DAY 7:
Meal 1 (10 AM): Greek yogurt with berries, nuts, and a drizzle of honey
Meal 2 (2 PM): Tuna lettuce wraps with chopped vegetables and homemade dressing
Meal 3 (5:30 PM): Lemon herb chicken thighs with roasted cauliflower and mixed green salad

SNACKS (if needed within eating window):
• Handful of mixed nuts (almonds, walnuts, pecans)
• Apple slices with almond butter
• Dark chocolate (70%+ cacao) with raw almonds
• Celery with guacamole
• Hard-boiled eggs
'''

doc.add_paragraph(meal_plan)

# Add page break
doc.add_page_break()

# Sleep Optimization
sleep_heading = doc.add_heading('Part 3: Sleep Optimization Protocol', level=1)
sleep_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

sleep_text = '''Gary Brecka emphasizes that sleep is the foundation of all health optimization. Without quality sleep, no protocol will work optimally.

SLEEP SCHEDULE: 7-9 Hours Consistently
Recommended: 10:00 PM - 6:00 AM (8 hours)
Then 30-minute morning routine before 10 AM breakfast

EVENING PROTOCOL (6:00 PM - 10:00 PM):

6:00 PM - Last meal completed
• Allows 4 hours for digestion before bed
• Reduces acid reflux and improves sleep quality

7:00-8:00 PM - Wind-down activities
• Light walk or gentle stretching
• Reading, journaling, or gratitude practice
• Avoid intense exercise or stimulating activities

8:00 PM - Screen curfew begins
• No phones, tablets, computers, or TV
• Blue light disrupts melatonin production
• If screens necessary, use blue-light blocking glasses

9:00 PM - Sleep preparation
• Dim lights throughout home
• Cool bedroom temperature (65-68°F optimal)
• Take sleep-support supplements (see supplement section)
• Hot shower or bath (body temperature drop after promotes sleep)

9:30 PM - Breathwork for sleep
• Box breathing technique: 4-second inhale, 4-second hold, 4-second exhale, 4-second hold
• Repeat for 5-10 minutes
• Focus solely on breath until drifting to sleep

10:00 PM - Lights out

MORNING PROTOCOL (6:00 AM - 10:00 AM):

6:00 AM - Wake at consistent time (even weekends)
• No snooze button - this disrupts sleep architecture
• Open blinds immediately for natural light

6:05 AM - Hydration ritual
• 8-10 oz water with pinch of Celtic sea salt or Himalayan pink salt
• Rehydrates body and provides essential minerals
• Can add lemon for additional benefits

6:10 AM - First light exposure (10-15 minutes)
• Get outside if possible, or near a window
• Natural sunlight exposure signals circadian rhythm
• Sets cortisol awakening response properly
• Improves mood, energy, and sleep quality that night
• Look toward the horizon (not directly at sun)

6:25 AM - Breathwork practice (5-10 minutes)
• Three rounds of 30 deep breaths (Wim Hof method)
• Or alternate nostril breathing for calmness
• Activates nervous system and increases oxygen

6:35 AM - Cold exposure (optional but highly recommended)
• 30-60 seconds cold shower
• Or splash cold water on face
• Boosts circulation, reduces inflammation, improves mental clarity
• Start with 30 seconds and gradually increase

6:40 AM - Light stretching or yoga (10 minutes)
• Gentle movement to activate muscles
• Prepares body for rowing workout

6:50 AM - Pre-workout preparation
• Prepare 30g protein breakfast
• Lay out rowing area
• Ensure proper hydration

7:00 AM - Consume protein breakfast (30g+)

7:30 AM - Begin 30-minute rowing session

8:00 AM - Post-workout hydration and recovery

SLEEP QUALITY OPTIMIZATIONS:
• Blackout curtains or eye mask (complete darkness)
• White noise machine or fan (consistent sound)
• Remove electronic devices from bedroom
• Keep bedroom cool (65-68°F)
• Invest in quality mattress and pillows
• Consider weighted blanket (helps with anxiety)
• Use lavender essential oil for relaxation'''

doc.add_paragraph(sleep_text)

# Add page break
doc.add_page_break()

# Rowing Integration
rowing_heading = doc.add_heading('Part 4: Rowing Routine Integration', level=1)
rowing_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

rowing_text = '''Your daily 30-minute rowing session perfectly aligns with Gary Brecka's exercise philosophy and the 30/30/30 method.

GARY BRECKA'S EXERCISE PRINCIPLES:

1. STEADY-STATE CARDIO FOR FAT BURNING
• Keep heart rate at 60-70% of maximum (approximately 100-120 bpm for age 60)
• This heart rate zone maximizes fat oxidation
• Prevents muscle catabolism (breakdown)
• Sustainable for daily practice

2. CONSISTENCY OVER INTENSITY
• Daily 30-minute sessions are ideal
• Better than intense workouts 2-3x per week
• Builds cardiovascular endurance without excessive cortisol
• Supports mitochondrial health and longevity

3. MORNING TIMING BENEFITS
• Exercise in fasted state (before or right after protein meal)
• Enhanced fat burning and metabolic flexibility
• Improves insulin sensitivity throughout the day
• Sets positive circadian rhythm

ROWING PROTOCOL FOR OPTIMAL RESULTS:

WEEK 1-2: Baseline Establishment
• 30 minutes at comfortable, conversational pace
• Heart rate: 90-100 bpm
• Focus on proper form and breathing
• Track: Distance, calories burned, average pace

WEEK 3-4: Steady-State Optimization  
• 30 minutes at consistent moderate pace
• Heart rate: 100-110 bpm (60% max)
• Maintain steady stroke rate (18-22 strokes per minute)
• Goal: Build aerobic base

WEEK 5-6: Fat-Burning Zone
• 30 minutes at fat-burning heart rate
• Heart rate: 110-120 bpm (65-70% max)
• Stroke rate: 20-24 strokes per minute
• This is your long-term sustainable zone

ADVANCED PROTOCOL (Month 2+):
Option A - Steady State (Most Days)
• 30 minutes at 110-120 bpm
• Consistent pace throughout
• Optimal for fat burning

Option B - Interval Training (1-2x per week)
• 5-minute warm-up (low intensity)
• 20 minutes: 1-minute higher intensity (130 bpm) / 2-minute recovery (100 bpm)
• 5-minute cool-down
• Builds cardiovascular capacity

FORM CHECKLIST:
1. Legs drive the stroke (not arms)
2. Core engaged throughout
3. Relaxed shoulders
4. Rhythmic breathing (exhale on drive, inhale on recovery)
5. Full range of motion

MONITORING METRICS:
• Heart rate (most important - stay in zone)
• Distance per session
• Average pace (minutes per 500m)
• Calories burned
• Stroke rate
• Weekly consistency (aim for 6-7 days)

RECOVERY AND PROGRESSION:
• Listen to your body - rest if overly fatigued
• Gradually increase resistance or pace every 2-3 weeks
• Proper hydration before, during, and after
• Post-rowing stretching (5-10 minutes)
• Consider foam rolling for muscle recovery

INTEGRATION WITH OTHER PROTOCOLS:
• Complete rowing AFTER 30g protein meal
• Maintain fasted or low-insulin state for maximum fat burn
• Hydrate with electrolyte water during workout
• Break fast with remaining breakfast after rowing if needed'''

doc.add_paragraph(rowing_text)

# Add page break
doc.add_page_break()

# Supplement Recommendations
supplement_heading = doc.add_heading('Part 5: Gary Brecka Supplement Recommendations', level=1)
supplement_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

supplement_intro = '''Gary Brecka strongly emphasizes personalized supplementation based on individual needs, ideally determined through genetic testing and blood work. However, he recommends several foundational supplements that benefit most people, particularly those over 50.

IMPORTANT NOTE: Consult with your healthcare provider before starting any new supplement regimen, especially given your age and to rule out any contraindications.

PREFERRED BRANDS (Per Gary Brecka):
• BodyHealth
• Pure Encapsulations
• Thorne
• Cymbiotika
• AOR (Advanced Orthomolecular Research)'''

doc.add_paragraph(supplement_intro)

doc.add_heading('Core Daily Supplements', level=2)

core_supplements = '''
1. METHYLATED MULTIVITAMIN
Why: Many people have MTHFR genetic mutations affecting nutrient absorption
Benefits: Better energy, reduced inflammation, improved mood
Look for: Methylfolate (5-MTHF), methylcobalamin (B12), pyridoxal-5-phosphate (B6)
Recommended: Pure Encapsulations O.N.E. Multivitamin or Thorne Basic Nutrients 2/Day
Dose: As directed (usually 1-2 capsules daily with food)

2. VITAMIN D3 + K2
Why: Critical for bone health, immune function, mood at age 60
Benefits: Calcium absorption, cardiovascular health, reduced inflammation
Recommended: Thorne Vitamin D/K2 Liquid
Dose: 5,000 IU D3 daily (or as determined by blood test)
Note: Take with healthy fat for absorption

3. OMEGA-3 FATTY ACIDS (EPA/DHA)
Why: Reduces inflammation, supports heart and brain health
Benefits: Improved cognitive function, joint health, mood stability
Recommended: Wild-caught fish oil or krill oil
Dose: 2,000-3,000mg combined EPA/DHA daily
Alternative: Eat fatty fish 3-4x per week

4. MAGNESIUM
Why: Involved in 300+ enzymatic reactions, often deficient with age
Benefits: Better sleep, muscle recovery, stress reduction, heart health
Type: Magnesium glycinate (best absorption, sleep support)
Or: Magnesium threonate (brain health/cognitive)
Recommended: Pure Encapsulations Magnesium Glycinate or Cymbiotika Magnesium L-Threonate
Dose: 200-400mg before bed

5. PERFECT AMINO (Essential Amino Acids)
Why: Gary Brecka's favorite supplement for muscle preservation
Benefits: 99% utilization, muscle repair, minimal calories
Recommended: BodyHealth Perfect Amino
Dose: 5-10 tablets daily or 1 scoop powder
Best taken: Morning or post-workout
Perfect for: Age 60+ to prevent muscle loss

6. METHYLFOLATE (L-5-MTHF)
Why: Essential for methylation, reduces homocysteine, supports mood
Benefits: Reduced anxiety, better energy, cardiovascular health
Recommended: Thorne 5-MTHF or Pure Encapsulations
Dose: 400-1,000mcg daily

7. COENZYME Q10 (CoQ10)
Why: Critical for mitochondrial energy production, declines with age
Benefits: Heart health, energy, antioxidant protection
Form: Ubiquinol (more bioavailable than ubiquinone)
Dose: 100-200mg daily with food

SLEEP-SUPPORT SUPPLEMENTS (Take 30-60 min before bed):

8. L-THEANINE
Benefits: Promotes relaxation without drowsiness, improves focus
Recommended: Pure Encapsulations
Dose: 200-400mg before bed

9. MAGNESIUM (as listed above)
Benefits: Muscle relaxation, calms nervous system
Best form for sleep: Glycinate or Threonate

10. MELATONIN (if needed)
Benefits: Regulates sleep-wake cycle
Note: Start low, use only if sleep issues persist
Dose: 0.5-3mg (start with lowest effective dose)
Recommended: Pure Encapsulations

ADDITIONAL SUPPLEMENTS TO CONSIDER:

11. ELECTROLYTES
Why: Crucial for hydration, especially during fasting and exercise
Recommended: LMNT, BodyHealth Complete Multi + Greens, or 10X Health Electrolytes
When: Morning with water, during/after rowing

12. HYDROGEN WATER or HYDROGEN SUPPLEMENTS
Why: Gary Brecka's top recommendation for reducing inflammation and oxidative stress
Benefits: Powerful selective antioxidant, reduces aging effects, improves recovery
Recommended: Echo Go Plus Hydrogen Water Bottle (use code for discount)
Alternative: BodyHealth Hydrogen Boost tablets

13. GUT RESTORE (Probiotics + Immunoglobulins)
Why: Gut health is foundational to overall health
Benefits: Reduces bloating, improves digestion, supports immune function
Recommended: BodyHealth Gut Restore
Dose: As directed

14. NAC (N-Acetyl Cysteine)
Benefits: Powerful antioxidant, supports liver detox, respiratory health
Dose: 600-1,200mg daily

SUPPLEMENT SCHEDULE:

MORNING (with breakfast, 10 AM):
• Methylated Multivitamin
• Vitamin D3 + K2
• Omega-3
• CoQ10
• Electrolytes (in water)

AFTERNOON (with lunch, 2 PM):
• Perfect Amino (if not taken in morning)
• Methylfolate
• Additional Omega-3 if split dose

EVENING (30-60 min before bed, 9:00 PM):
• Magnesium Glycinate
• L-Theanine
• Melatonin (if needed)

POST-ROWING:
• Perfect Amino
• Electrolyte water

IMPORTANT CONSIDERATIONS:
• Start supplements gradually (1-2 at a time) to assess tolerance
• Quality matters - avoid cheap, low-quality brands
• Store properly (some require refrigeration)
• Take with food when indicated for better absorption
• Monitor how you feel - adjust as needed
• Consider genetic testing (10X Health System) for personalized recommendations
• Retest blood work every 3-6 months to track progress

ESTIMATED MONTHLY COST:
Core supplements: $150-250/month
With additional optimizations: $300-400/month
(Investment in longevity and quality of life)'''

doc.add_paragraph(core_supplements)

# Add page break
doc.add_page_break()

# Lifestyle Practices
lifestyle_heading = doc.add_heading('Part 6: Additional Lifestyle Practices', level=1)
lifestyle_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

lifestyle_text = '''Gary Brecka emphasizes several daily practices beyond diet and exercise:

1. COLD EXPOSURE
• 30-60 seconds cold shower finish each morning
• Reduces inflammation and boosts circulation
• Improves mental clarity and resilience
• Activates brown adipose tissue (fat burning)

2. BREATHWORK
• 5-10 minutes daily
• Options: Wim Hof method, box breathing, alternate nostril breathing
• Reduces stress, improves oxygenation
• Balances autonomic nervous system

3. GROUNDING (EARTHING)
• 10-15 minutes daily barefoot contact with earth
• Reduces inflammation and cortisol
• Improves sleep quality
• Best done during morning sun exposure

4. GRATITUDE PRACTICE
• Daily journaling (5 minutes)
• Write 3-5 things you're grateful for
• Improves mental health and emotional resilience
• Reduces stress and anxiety

5. HYDRATION OPTIMIZATION
• Drink half your body weight in ounces daily (95 oz for 190 lbs)
• Add pinch of Celtic sea salt or Himalayan pink salt to water
• First thing in morning and throughout day
• Herbal teas count toward hydration

6. REDUCE TOXIC EXPOSURES
• Filter drinking water
• Choose organic produce when possible (Dirty Dozen)
• Minimize processed foods
• Use natural cleaning and personal care products

7. STRESS MANAGEMENT
• Regular meditation or mindfulness practice
• Limit news and social media consumption
• Prioritize relationships and social connection
• Engage in hobbies and activities you enjoy'''

doc.add_paragraph(lifestyle_text)

# Add page break
doc.add_page_break()

# Implementation Plan
implementation_heading = doc.add_heading('Part 7: 90-Day Implementation Plan', level=1)
implementation_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

implementation_text = '''Making too many changes at once often leads to overwhelm and failure. This phased approach ensures sustainable adoption.

PHASE 1: WEEKS 1-2 (Foundation Building)
Focus: Establish sleep schedule and morning routine

Actions:
☐ Set consistent sleep schedule (10 PM - 6 AM)
☐ Eliminate screens after 8 PM
☐ Begin morning sunlight exposure (10-15 min)
☐ Start 30/30/30 method (30g protein + 30 min rowing)
☐ Track sleep quality and energy levels daily
☐ Begin one core supplement (Vitamin D3 + K2)

Expected Results:
• Better sleep within 7-10 days
• More stable morning energy
• Habit formation of consistent wake time

PHASE 2: WEEKS 3-4 (Nutrition Optimization)
Focus: Implement low-carb nutrition and eliminate processed foods

Actions:
☐ Transition to low-carb, nutrient-dense meals
☐ Eliminate processed foods, refined sugars, and grains
☐ Meal prep 2x per week
☐ Add methylated multivitamin and omega-3
☐ Continue Phase 1 habits
☐ Monitor weight, energy, and mood

Expected Results:
• Initial weight loss (5-7 lbs, mostly water)
• Reduced cravings and stable energy
• Better mental clarity

PHASE 3: WEEKS 5-6 (Time-Restricted Eating)
Focus: Implement 16:8 fasting protocol

Actions:
☐ Begin 10 AM - 6 PM eating window
☐ Continue all previous habits
☐ Add remaining core supplements
☐ Implement evening wind-down protocol
☐ Add breathwork practice (5-10 min daily)

Expected Results:
• Enhanced fat burning
• Improved metabolic flexibility
• Better sleep quality
• Potential adjustment period (hunger, fatigue) - normal and temporary

PHASE 4: WEEKS 7-8 (Advanced Optimization)
Focus: Add cold exposure and optimization practices

Actions:
☐ Begin cold shower finishes (30-60 seconds)
☐ Add grounding practice (10-15 min daily)
☐ Implement gratitude journaling
☐ Fine-tune rowing intensity (heart rate zones)
☐ Continue all previous practices

Expected Results:
• Noticeable improvements in recovery
• Enhanced stress resilience
• Improved mood and mental clarity

PHASE 5: WEEKS 9-12 (Mastery and Refinement)
Focus: Solidify all habits and optimize based on results

Actions:
☐ All protocols running consistently
☐ Consider genetic testing (10X Health) for personalization
☐ Get comprehensive blood work to assess progress
☐ Adjust supplements based on results
☐ Celebrate wins and assess overall progress

Expected Results:
• Significant weight loss (10-15 lbs total)
• Dramatically improved energy and vitality
• Better sleep, mood, and mental clarity
• Reduced inflammation and health markers improved
• Sustainable lifestyle established

TRACKING METRICS:

Weekly:
• Body weight (same time, same conditions)
• Resting heart rate (upon waking)
• Sleep quality (1-10 scale)
• Energy levels (1-10 scale, 3x daily)
• Mood and mental clarity (1-10 scale)
• Rowing metrics (distance, pace, HR)

Monthly:
• Body measurements (waist, chest, hips)
• Progress photos
• Review and adjust protocols
• Assess supplement effectiveness

Quarterly:
• Comprehensive blood work
• Body composition analysis
• Assess overall progress toward goals'''

doc.add_paragraph(implementation_text)

# Add page break
doc.add_page_break()

# Troubleshooting
trouble_heading = doc.add_heading('Part 8: Common Challenges & Solutions', level=1)
trouble_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

trouble_text = '''
CHALLENGE: "I'm hungry during fasting periods"
SOLUTION:
• Ensure adequate protein and fat at meals (increases satiety)
• Drink water with pinch of salt when hungry
• Black coffee or green tea can reduce appetite
• First 2-3 weeks are adjustment period - it gets easier
• Consider starting with 12:12, then 14:10, then 16:8

CHALLENGE: "I feel low energy in the afternoon"
SOLUTION:
• May need to adjust carb intake slightly upward
• Ensure adequate hydration and electrolytes
• Check magnesium levels (common deficiency)
• Power nap (15-20 minutes) can be restorative
• May indicate need for more sleep at night

CHALLENGE: "Not losing weight as expected"
SOLUTION:
• Ensure true low-carb intake (track for 3 days)
• Check for hidden carbs in sauces, dressings
• May need to reduce eating window further (18:6)
• Ensure adequate protein (muscle preservation)
• Consider HIIT interval rowing 1-2x per week
• Check thyroid function (common at age 60)
• Be patient - body composition changes take time

CHALLENGE: "Trouble falling asleep"
SOLUTION:
• Stricter screen curfew (7 PM if needed)
• Ensure complete darkness in bedroom
• Try magnesium glycinate at higher dose (up to 400mg)
• Add L-theanine 30-60 min before bed
• Box breathing for 10-15 minutes in bed
• May need melatonin temporarily (0.5-1mg)
• Avoid caffeine after 12 PM

CHALLENGE: "Muscle soreness from rowing"
SOLUTION:
• Ensure proper form (legs drive, not arms/back)
• Post-rowing stretching routine (10 minutes)
• Adequate protein intake for recovery
• Perfect Amino for muscle repair
• Magnesium for muscle relaxation
• Consider foam rolling
• Take rest day if needed - listen to your body

CHALLENGE: "Social situations and eating window"
SOLUTION:
• Plan ahead - eat larger lunch if dinner out
• Choose protein-focused meals when dining out
• Shift window occasionally if needed (be flexible)
• Don't sacrifice relationships for rigid protocols
• Get back on track next day - consistency matters more than perfection

CHALLENGE: "Supplements are expensive"
SOLUTION:
• Prioritize core 4-5 most important (Multivitamin, D3, Omega-3, Magnesium, Perfect Amino)
• Buy in bulk when possible
• Consider this an investment in longevity vs. future medical costs
• Start with basics, add others as budget allows
• Look for quality brands on sale

CHALLENGE: "Too overwhelming to implement everything"
SOLUTION:
• Follow phased approach - don't do everything at once
• Master one habit before adding another
• Start with sleep optimization (highest ROI)
• Remember: consistency beats perfection
• Even partial implementation yields benefits'''

doc.add_paragraph(trouble_text)

# Add page break
doc.add_page_break()

# Conclusion
conclusion_heading = doc.add_heading('Conclusion & Next Steps', level=1)
conclusion_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

conclusion_text = '''Gary Brecka's health philosophy is built on a simple but powerful premise: your body has an incredible capacity to heal, optimize, and thrive when given the right inputs. At age 60, you're at a critical juncture where the daily choices you make will dramatically impact your health span and longevity for the next 20-30+ years.

This plan provides a comprehensive roadmap, but remember:

• Progress over perfection - small, consistent habits compound dramatically over time
• Your body is unique - adjust protocols based on how YOU feel and respond
• Data is your friend - track metrics and use them to guide decisions
• Community and accountability accelerate results - consider sharing your journey
• Health is wealth - this is the most important investment you'll ever make

THE COMPOUND EFFECT:
By implementing these protocols consistently for 90 days:
• You'll likely lose 10-20 pounds of fat
• Energy levels will increase dramatically
• Sleep quality will improve significantly
• Mental clarity and mood will enhance
• Inflammation and health markers will improve
• You'll feel 5-10 years younger

But more importantly, you'll have established sustainable habits that will serve you for decades to come.

YOUR IMMEDIATE NEXT STEPS:

1. Schedule a comprehensive blood panel with your physician
   • Lipid panel, metabolic panel, inflammatory markers
   • Vitamin D, B12, magnesium levels
   • Thyroid function, testosterone
   • Baseline for tracking progress

2. Order your core supplements
   • Start with top 4-5 priorities
   • Choose quality brands (Gary Brecka approved)

3. Set up your environment for success
   • Clean out pantry of processed foods
   • Stock kitchen with approved foods
   • Prepare bedroom for optimal sleep
   • Set up tracking system (app or journal)

4. Choose your Phase 1 start date
   • Mark it on calendar
   • Commit fully for 2 weeks minimum
   • Share with accountability partner

5. Consider genetic testing
   • 10X Health System (Gary Brecka's company)
   • Provides personalized insights based on YOUR genetics
   • Optimizes supplement and nutrition recommendations

REMEMBER: You're not just adding years to your life - you're adding life to your years. At 60, with consistent implementation of these evidence-based protocols, you can feel and function better than you have in decades.

The time to start is now. Your future self will thank you.

To your health and longevity,

Your Wellness Planning Team

---

RESOURCES:
• Gary Brecka's website: www.theultimatehuman.com
• 10X Health System: www.10xhealthsystem.com
• The Ultimate Human Podcast (available on all platforms)
• Supplement brands: BodyHealth, Pure Encapsulations, Thorne, Cymbiotika

DISCLAIMER: This wellness plan is for educational purposes only and does not constitute medical advice. Always consult with your healthcare provider before making significant changes to your diet, exercise routine, or supplement regimen, especially if you have any pre-existing health conditions or take medications.'''

doc.add_paragraph(conclusion_text)

# Add footer with date
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_para.add_run(f'\nDocument created: {datetime.date.today().strftime("%B %d, %Y")}\nBased on Gary Brecka\'s health philosophy and research')
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(128, 128, 128)

# Save the document
doc.save('Eric_Brown_Wellness_Plan_Gary_Brecka.docx')
print('✅ Wellness plan document created successfully!')
