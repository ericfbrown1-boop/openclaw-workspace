"""Generate Follow-Up Tracker Project Plan Word Document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os


doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)


def add_h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 63, 128)


def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 100, 160)


def tbl(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    doc.add_paragraph()


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.runs[0].font.size = Pt(10)


def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10)


# ── TITLE ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PROJECT FOLLOW-UP TRACKER")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0, 63, 128)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run(
    "Automated Email-Triggered Project Monitoring & Nagging System"
)
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(80, 80, 80)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run(
    f"Concept Design & Implementation Plan | {datetime.now().strftime('%B %d, %Y')}"
)
run3.font.size = Pt(11)
run3.font.color.rgb = RGBColor(120, 120, 120)
run3.italic = True

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run(
    "Prepared by Jarvis AI (Research + Planning Agents) | "
    "For: Eric Brown, CFO & COO, Cohesity"
)
run4.font.size = Pt(9)
run4.font.color.rgb = RGBColor(150, 150, 150)
run4.italic = True

doc.add_page_break()

# ── 1. CONCEPT ──
add_h1("1. The Concept: Email-Triggered Project Autopilot")

body(
    "You send one email from your Cohesity work account. That single email becomes the "
    "trigger for an end-to-end automated follow-up system that tracks the project to "
    "completion. No manual task creation, no separate project management tool, no "
    "remembering to follow up. The system does it all."
)

add_h2("1.1 How It Works (User Experience)")
body("Step 1: Eric composes an email in Outlook from Eric.brown@cohesity.com")
body(
    'Step 2: In the subject line, Eric includes the trigger phrase: '
    '"Follow Up Requested"'
)
body(
    "Step 3: Eric writes the email normally -- assigns the project, describes "
    "deliverables, names owners, sets expectations."
)
body("Step 4: Eric hits Send. That is the last manual step.")
body("Step 5: The system takes over:")
for b in [
    "Detects the trigger email within minutes",
    "Parses the email body to extract: project name, owners, deliverables, "
    "implied timeline, success criteria",
    "Creates a structured project tracker (Google Sheet or internal DB)",
    "Sets up a daily monitoring scan of the email thread",
    "Sends Eric a confirmation: 'Project tracked. First check-in scheduled.'",
    "Every day: scans for replies, measures progress, identifies blockers",
    "If no progress: escalates with a nudge email (to assignees) or alert (to Eric)",
    "Continues until all deliverables are marked complete or Eric cancels",
]:
    bullet(b)

add_h2("1.2 The Trigger Email Format")
body("Subject line pattern:")
p = doc.add_paragraph()
run = p.add_run(
    "  Follow Up Requested: [Project Name] - [Brief Description]"
)
run.font.size = Pt(11)
run.bold = True
run.font.color.rgb = RGBColor(0, 100, 160)

body("Example:")
p = doc.add_paragraph()
run = p.add_run(
    "  Follow Up Requested: Q2 Board Deck - Finance team to deliver "
    "final board materials by April 25"
)
run.font.size = Pt(10)
run.italic = True

body(
    "The body of the email is written naturally -- no special formatting needed. "
    "The AI parses it for context, owners, dates, and deliverables."
)

doc.add_paragraph()

# ── 2. ARCHITECTURE ──
add_h1("2. System Architecture")

add_h2("2.1 Implementation Options")
body(
    "There are three viable architectures. We recommend Option A for fastest "
    "deployment using tools already in your stack."
)

tbl(
    ["Option", "Platform", "Email Access", "Scheduling", "Pros", "Cons"],
    [
        (
            "A: Jarvis + OpenClaw (Recommended)",
            "Jarvis AI on MacBook",
            "Gmail API via gog CLI + Zapier MCP",
            "OpenClaw cron jobs",
            "Already built; full control; no new subscriptions",
            "Requires MacBook running; uses personal Gmail",
        ),
        (
            "B: Claude CoWork + M365",
            "Claude Desktop CoWork",
            "M365 Outlook connector",
            "CoWork scheduled tasks",
            "Native M365 integration; uses work email directly",
            "CoWork is research preview; M365 connector needed",
        ),
        (
            "C: Microsoft Copilot CoWork",
            "M365 Copilot",
            "Native Outlook access",
            "Copilot CoWork automation",
            "Deepest M365 integration; no external tools",
            "Requires Copilot license; less customizable",
        ),
    ],
)

add_h2("2.2 Option A: Jarvis + OpenClaw (Recommended)")
body("This uses your existing infrastructure with zero new tools:")

tbl(
    ["Component", "Tool", "Role"],
    [
        ("Email Detection", "gog gmail search / Zapier MCP", "Scan for 'Follow Up Requested' in sent mail every 15 min"),
        ("Email Parsing", "Claude Opus 4.6 (Jarvis)", "Extract project details, owners, timeline, deliverables from email body"),
        ("Project Tracker", "Google Sheets (via gog/Zapier)", "One sheet per project with columns: Task, Owner, Due Date, Status, Last Activity"),
        ("Thread Monitoring", "gog gmail search (thread ID)", "Daily scan of the original email thread for replies and updates"),
        ("Progress Analysis", "Claude Opus 4.6", "Analyze replies to determine: progress made, blockers, missing responses"),
        ("Nudge Emails", "Zapier MCP gmail_send", "Auto-send follow-up reminders to assignees who haven't responded"),
        ("Eric Alerts", "Telegram (Jarvis)", "Alert Eric when action is needed, deadlines at risk, or project stalls"),
        ("Completion", "Jarvis + Google Sheets", "Mark project complete when all deliverables confirmed; send summary"),
        ("Scheduling", "OpenClaw cron", "15-min detection scan + daily deep analysis at 9 AM PT"),
    ],
)

add_h2("2.3 Option B: Claude CoWork + M365 Connector")
body(
    "Claude CoWork (launched Jan 2026) now supports M365 connectors including "
    "Outlook email reading. A CoWork scheduled task could:"
)
for b in [
    "Run daily at 9 AM: scan Outlook sent folder for 'Follow Up Requested'",
    "Parse email threads using Claude's native language understanding",
    "Create and update a tracking document in a CoWork project folder",
    "Use the Gmail connector (if available) or Slack connector to send nudges",
    "CoWork limitation: runs only when Claude Desktop app is open on your Mac",
    "CoWork advantage: can access local files and create formatted documents",
]:
    bullet(b)

body(
    "Current limitations: CoWork is a research preview (Feb 2026). M365 connector "
    "supports reading Outlook emails but write/send capabilities are evolving. "
    "The Gmail connector is more mature for sending follow-ups."
)

add_h2("2.4 Option C: Microsoft Copilot CoWork")
body(
    "Microsoft's Copilot CoWork (announced Mar 2026) offers the deepest M365 "
    "integration. It can ground work in emails, meetings, messages, files, and data "
    "natively. However, it requires a Copilot license and is less customizable "
    "than the Jarvis approach. Best suited for organizations fully committed to "
    "the M365 Copilot ecosystem."
)

doc.add_paragraph()

# ── 3. DETAILED WORKFLOW ──
add_h1("3. Detailed Workflow (Option A)")

add_h2("3.1 Phase 1: Detection (Every 15 Minutes)")
body("OpenClaw cron job runs every 15 minutes:")
p = doc.add_paragraph()
run = p.add_run(
    '  gog gmail search "from:eric.brown@cohesity.com subject:Follow Up Requested '
    'newer_than:1h" --max 10'
)
run.font.size = Pt(9)
run.font.name = "Courier New"

body(
    "For each new match not already tracked, Jarvis reads the full email and "
    "initiates the parsing pipeline."
)

add_h2("3.2 Phase 2: AI Parsing (Automatic)")
body(
    "Claude Opus 4.6 analyzes the email body and extracts a structured project record:"
)
tbl(
    ["Field", "Extraction Method", "Example"],
    [
        ("Project Name", "From subject line after 'Follow Up Requested:'", "Q2 Board Deck"),
        ("Description", "First paragraph of email body", "Finance team to deliver final board materials"),
        ("Owner(s)", "Named individuals or teams in email body + To/CC", "Sarah Chen, Mike Torres"),
        ("Deliverables", "Action items, bullet points, numbered lists", "1. Revenue slide, 2. Forecast model, 3. Risk summary"),
        ("Due Date", "Explicit date or implied deadline", "April 25, 2026"),
        ("Priority", "Urgency language analysis (ASAP, critical, etc.)", "HIGH"),
        ("Success Criteria", "What 'done' looks like", "Board-ready PDF in shared drive"),
        ("Check-in Frequency", "Default daily; override if specified", "Daily"),
        ("Escalation Path", "Default: nudge at Day 2, Eric alert at Day 4", "Standard"),
    ],
)

add_h2("3.3 Phase 3: Tracker Creation")
body(
    "A new Google Sheet is created (or a new tab added to a master tracker sheet) "
    "with the following structure:"
)
tbl(
    ["Column", "Content"],
    [
        ("A: Task/Deliverable", "Each deliverable as a separate row"),
        ("B: Owner", "Assigned person"),
        ("C: Due Date", "Target completion date"),
        ("D: Status", "NOT STARTED / IN PROGRESS / BLOCKED / COMPLETE"),
        ("E: Last Activity", "Timestamp of last email reply or update"),
        ("F: Days Silent", "Days since last activity (auto-calculated)"),
        ("G: Next Action", "AI-determined next required step"),
        ("H: Nudge Count", "Number of follow-up reminders sent"),
        ("I: Notes", "AI summary of latest thread activity"),
    ],
)

add_h2("3.4 Phase 4: Daily Monitoring (9 AM PT)")
body("Every morning, the monitoring agent:")
for b in [
    "Searches the original email thread for new replies (by thread ID)",
    "Analyzes each reply: Did the sender provide an update? Is there a blocker? "
    "Did they commit to a date?",
    "Updates the Google Sheet tracker with new status information",
    "Calculates 'Days Silent' for each owner -- how long since their last response",
    "Identifies at-risk items: overdue deliverables, owners who haven't responded",
    "Generates a daily digest for Eric (via Telegram) with: "
    "green (on track), yellow (at risk), red (overdue/stalled)",
]:
    bullet(b)

add_h2("3.5 Phase 5: Escalation & Nudging")
body("Automated escalation ladder:")
tbl(
    ["Trigger", "Action", "Tone"],
    [
        ("Day 1: No response", "No action yet (grace period)", "-"),
        ("Day 2: Still no response", "Gentle reminder email to assignee(s)", "Friendly: 'Just checking in on...'"),
        ("Day 3: No progress", "Firmer follow-up email", "Professional: 'Following up on the items below...'"),
        ("Day 4: Still stalled", "Alert Eric via Telegram with specific blockers", "Direct: 'Project X stalled -- [Owner] non-responsive'"),
        ("Day 5+: No movement", "Eric decides: re-send, escalate, or reassign", "Eric's call"),
        ("Due date passed", "Immediate alert to Eric + overdue notice to owners", "Urgent: 'Past due -- immediate attention required'"),
    ],
)

body(
    "Important: Nudge emails are sent FROM Eric's account (maintaining his authority) "
    "but drafted and sent by the system. Each nudge references the original email "
    "thread so recipients see the full context."
)

add_h2("3.6 Phase 6: Completion")
body("A project is marked COMPLETE when:")
for b in [
    "All deliverables in the tracker are marked COMPLETE",
    "Eric manually marks it done (override)",
    "Eric replies 'CLOSE' to the original thread",
    "The specified due date passes and Eric confirms closure",
]:
    bullet(b)

body(
    "Upon completion, the system sends Eric a final summary: project duration, "
    "response times by owner, any delays, and a performance scorecard."
)

doc.add_paragraph()

# ── 4. NUDGE EMAIL TEMPLATES ──
add_h1("4. Sample Nudge Email Templates")

add_h2("4.1 Day 2: Gentle Reminder")
body("Subject: RE: [Original Subject]")
body(
    '"Hi [Name],\n\n'
    "Just checking in on the items below from Eric's email on [date]. "
    "Could you provide a quick status update when you get a chance?\n\n"
    "[List of assigned deliverables]\n\n"
    "Thanks!\n"
    'Eric"'
)

add_h2("4.2 Day 3: Follow-Up")
body("Subject: RE: [Original Subject] - Update Needed")
body(
    '"Hi [Name],\n\n'
    "Following up on the project below. We're approaching the [date] deadline "
    "and I haven't received an update yet. Please share your progress and any "
    "blockers you're facing.\n\n"
    "[List of deliverables with due dates]\n\n"
    "Please reply by end of day.\n\n"
    'Eric"'
)

add_h2("4.3 Overdue Notice")
body("Subject: RE: [Original Subject] - OVERDUE - Action Required")
body(
    '"Hi [Name],\n\n'
    "The following items were due on [date] and are now overdue:\n\n"
    "[List of overdue deliverables]\n\n"
    "Please provide an immediate status update and revised completion date.\n\n"
    'Eric"'
)

doc.add_paragraph()

# ── 5. COHESITY STACK FIT ──
add_h1("5. Integration with Cohesity's Tech Stack")

tbl(
    ["Cohesity Tool", "Integration Opportunity", "Status"],
    [
        ("Outlook / M365", "Trigger email source; thread monitoring", "Ready via Jarvis Zapier MCP or CoWork M365 connector"),
        ("Slack", "Real-time nudge channel; #project-followups", "Ready via Slack webhook or CoWork Slack connector"),
        ("Zoom", "Detect meeting references in emails; auto-schedule check-ins", "Possible via calendar API"),
        ("Salesforce", "Link follow-ups to deal/opportunity tracking", "Possible via Jarvis Salesforce skill"),
        ("Google Workspace (personal)", "Tracker sheets; Gmail fallback", "Ready via gog CLI + Zapier MCP"),
        ("Workday", "Link to employee/team data for owner resolution", "Future enhancement"),
        ("SharePoint / OneDrive", "Monitor deliverable uploads as completion signals", "Possible via M365 connector"),
    ],
)

body(
    "The system is designed to work with Eric's personal Gmail first (fastest path) "
    "with a clear upgrade path to Cohesity M365 when Claude CoWork's enterprise "
    "connectors mature."
)

doc.add_paragraph()

# ── 6. IMPLEMENTATION PLAN ──
add_h1("6. Implementation Plan")

add_h2("6.1 Phase 1: MVP (1-2 Weeks)")
tbl(
    ["Task", "Owner", "Timeline"],
    [
        ("Build email detection cron (gog gmail search)", "Jarvis", "Day 1-2"),
        ("Build AI parsing pipeline (Claude structured outputs)", "Jarvis", "Day 2-3"),
        ("Create Google Sheet tracker template", "Jarvis", "Day 3"),
        ("Build thread monitoring and reply analysis", "Jarvis", "Day 4-5"),
        ("Build nudge email generation and sending", "Jarvis", "Day 5-6"),
        ("Build Telegram alert integration", "Jarvis", "Day 6-7"),
        ("End-to-end testing with sample project email", "Eric + Jarvis", "Day 7-10"),
        ("Deploy as OpenClaw cron jobs", "Jarvis", "Day 10"),
    ],
)

add_h2("6.2 Phase 2: Enhancements (Weeks 3-4)")
for b in [
    "Dashboard view on Mission Control (new tab: Active Follow-Ups)",
    "Slack integration for real-time nudges (in addition to email)",
    "Auto-detection of completion signals (e.g., 'Done' or attachment uploads)",
    "Weekly summary report: all active projects, health status, response rates",
    "Multi-project view: cross-project dependency detection",
]:
    bullet(b)

add_h2("6.3 Phase 3: Enterprise Scale (Month 2+)")
for b in [
    "Claude CoWork integration when M365 connector supports Outlook send",
    "Migrate from personal Gmail to Cohesity M365 for all email operations",
    "Team-wide deployment: allow other execs to use 'Follow Up Requested' trigger",
    "Salesforce integration: link follow-ups to deal stages and pipeline",
    "AI-generated project timeline Gantt charts (auto-updated daily)",
]:
    bullet(b)

doc.add_paragraph()

# ── 7. PRIVACY & GOVERNANCE ──
add_h1("7. Privacy & Governance Considerations")

for b in [
    "Email content parsing: All email content stays within Jarvis's secure environment "
    "(MacBook + encrypted workspace). No data sent to third parties beyond the AI model API.",
    "Nudge emails: Sent as Eric (from Eric's account). Recipients see Eric as the sender, "
    "maintaining chain of command and authority.",
    "Google Sheets tracker: Stored in Eric's personal Google Drive (not shared externally). "
    "Future: migrate to Cohesity SharePoint for enterprise governance.",
    "Escalation controls: Eric can override, pause, or cancel any project tracking at any time "
    "via Telegram command or email reply.",
    "Audit trail: Every action (detection, parsing, nudge, alert) is logged in "
    "memory/followup-tracker/ with timestamps and content hashes.",
    "Opt-out: If a recipient replies 'STOP TRACKING' or similar, the system flags it for "
    "Eric's review rather than auto-removing.",
]:
    bullet(b)

doc.add_paragraph()

# ── 8. COST ──
add_h1("8. Estimated Cost")

tbl(
    ["Component", "Cost", "Notes"],
    [
        ("Jarvis AI (Claude Opus 4.6)", "Included in current plan", "Already running"),
        ("OpenClaw cron jobs", "Free", "Already running"),
        ("Google Sheets API", "Free", "Within quota"),
        ("Zapier MCP (Gmail send)", "Included in current plan", "20 tools available"),
        ("Claude CoWork (future)", "$100-200/mo", "Pro/Max plan; optional upgrade"),
        ("M365 Copilot (future)", "Per-seat license", "Only if using Option C"),
        ("Total (MVP)", "$0 incremental", "Uses existing infrastructure"),
    ],
)

doc.add_paragraph()

# ── 9. RECOMMENDATION ──
add_h1("9. Recommendation")

body(
    "Start with Option A (Jarvis + OpenClaw) for immediate deployment at zero "
    "incremental cost. The MVP can be operational within 10 days using your "
    "existing infrastructure. As Claude CoWork's M365 connectors mature, "
    "migrate the email monitoring to CoWork for deeper Outlook integration."
)

body("The key innovation here is simplicity:")
for b in [
    "ONE email triggers the ENTIRE workflow",
    "No new tools to learn, no project management UI to maintain",
    "The system adapts to natural language -- write emails normally",
    "Escalation is automatic but controllable",
    "Works with your existing Cohesity email + personal Gmail setup",
    "Every project gets the same rigorous follow-up treatment",
]:
    bullet(b)

body(
    "This effectively gives you an AI chief of staff that never forgets to follow up, "
    "never drops a thread, and nags people with exactly the right escalation curve "
    "until the work is done."
)

doc.add_paragraph()

# ── FOOTER ──
p_disc = doc.add_paragraph(
    "This plan was developed by Jarvis AI using Research Agent (market analysis, "
    "tool capabilities) and Planning Agent (architecture design, implementation "
    "timeline). Ready to proceed to MVP build on Eric's approval.\n\n"
    "Sent by Jarvis - AI assistant to Eric Brown"
)
for run in p_disc.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.italic = True

os.makedirs(os.path.expanduser("~/Documents/reports"), exist_ok=True)
outpath = os.path.expanduser(
    "~/Documents/reports/FollowUp_Tracker_Plan_20260405.docx"
)
doc.save(outpath)
print(f"Saved: {outpath}")
