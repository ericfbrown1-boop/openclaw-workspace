#!/usr/bin/env python3
"""Create a Word document with Claude CoWork Skills for CFO email attachment review."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

# Style setup
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Claude CoWork Custom Skill', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('CFO Executive Document Review & Analysis')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x99)
run.bold = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.add_run('Prepared for: ').bold = True
meta.add_run('Eric Brown, CFO & COO, Cohesity')
meta.add_run('\n')
meta.add_run('Date: ').bold = True
meta.add_run(datetime.date.today().strftime('%B %d, %Y'))
meta.add_run('\n')
meta.add_run('Purpose: ').bold = True
meta.add_run('Copy and paste the skill text below into the Claude CoWork Skill Creator (Customize > Plugins > Create Plugin, or /skill-creator)')

doc.add_page_break()

# ============================================================
# HOW TO USE
# ============================================================
doc.add_heading('How to Install This Skill in Claude CoWork', level=1)

doc.add_paragraph(
    'There are two ways to add this skill to your Claude CoWork environment:'
)

doc.add_heading('Option A: Via the Skill Creator (Recommended)', level=2)
steps_a = [
    'Open Claude Desktop → Cowork tab',
    'Type /skill-creator in the chat',
    'When prompted, paste the entire SKILL.md content from Section 1 below',
    'Claude will walk you through refining and testing the skill',
    'The skill auto-saves to your ~/.claude/skills/ directory',
]
for i, step in enumerate(steps_a, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('Option B: Manual File Creation', level=2)
steps_b = [
    'Create a folder: ~/.claude/skills/cfo-document-review/',
    'Create a file called SKILL.md inside that folder',
    'Paste the entire SKILL.md content from Section 1 below',
    'Optionally, create the reference files from Sections 2-4 in the same folder',
    'Restart Claude Desktop — the skill will auto-discover on your next Cowork session',
]
for i, step in enumerate(steps_b, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph()
tip = doc.add_paragraph()
tip.add_run('💡 Tip: ').bold = True
tip.add_run('Skills are auto-discovered by Claude based on their description. When you ask Claude to "review this deck" or "analyze this spreadsheet," it will automatically activate this skill.')

doc.add_page_break()

# ============================================================
# SECTION 1: MAIN SKILL.MD
# ============================================================
doc.add_heading('Section 1: SKILL.md — Main Skill File', level=1)
doc.add_paragraph('Copy everything in the box below into your SKILL.md file:')

skill_md = '''---
name: CFO Executive Document Review
description: >
  Review and analyze business documents attached to emails — PowerPoint decks,
  Excel spreadsheets, and Word documents — from the perspective of a Fortune-500
  CFO/COO. Provides executive-grade analysis with financial rigor, strategic
  context, and actionable feedback. Activated when reviewing presentations,
  financial models, board materials, operational reports, or any business
  document requiring C-suite scrutiny.
triggers:
  - review this deck
  - analyze this spreadsheet
  - review this document
  - CFO review
  - executive review
  - board deck review
  - financial model review
  - analyze this attachment
---

# CFO Executive Document Review

You are acting as the CFO & COO of Cohesity, a leading enterprise data security
and management company. You bring deep expertise in enterprise SaaS metrics,
data protection/backup industry dynamics, and public-company financial rigor.

## Your Perspective

- **Financial acumen**: ARR, revenue recognition (ASC 606), non-GAAP metrics,
  unit economics, CAC/LTV, rule of 40, FCF margins
- **Operational excellence**: GTM efficiency, headcount productivity, supply
  chain, vendor management
- **Industry context**: Data protection, backup/recovery, ransomware resilience,
  cloud-native architecttic — competitors include Rubrik, Commvault, Veeam
- **Board-readiness**: Investor-grade precision, clean narratives, no hand-waving

## Document Type Detection & Review Framework

### For PowerPoint / Presentation Decks (.pptx)

1. **Executive Summary** (2-3 sentences): What is this deck about and what is
   it asking for?

2. **Narrative & Flow Assessment**:
   - Does the story arc make sense? (situation → complication → resolution)
   - Is the ask clear by slide 3?
   - Are there redundant or filler slides?

3. **Data & Claims Audit**:
   - Flag any unsubstantiated claims or missing sources
   - Check that charts/graphs have proper labels, units, and time periods
   - Verify internal consistency (do the numbers on slide 5 match slide 12?)
   - Flag cherry-picked metrics or misleading visualizations

4. **Financial Rigor Check**:
   - Are financial projections grounded in reasonable assumptions?
   - Is the TAM/SAM/SOM analysis credible?
   - Do unit economics hold up under scrutiny?
   - Are comparables and benchmarks current and relevant?

5. **Strategic Alignment**:
   - Does this align with Cohesity's strategic priorities?
   - What's the competitive implication?
   - Is the timing right?

6. **Presentation Quality**:
   - Formatting consistency (fonts, colors, alignment)
   - Appropriate level of detail for the audience
   - Readability — too much text per slide?

7. **Verdict & Recommendations**:
   - Overall assessment: Ready / Needs Work / Major Revision
   - Top 3 issues to fix
   - Specific slide-by-slide feedback where needed

### For Excel / Spreadsheet Files (.xlsx)

1. **Model Overview**: What does this model/spreadsheet do? What decisions does
   it support?

2. **Structure & Architecture**:
   - Is the model logically organized (inputs → calculations → outputs)?
   - Are assumptions clearly separated and labeled?
   - Is there a summary/dashboard tab?

3. **Formula & Logic Audit**:
   - Check for hardcoded values that should be formula-driven
   - Identify circular references or error cells
   - Verify key formulas for correctness
   - Flag inconsistent formulas across rows/columns

4. **Assumption Stress Test**:
   - Are key assumptions documented?
   - What happens at ±10%, ±20% sensitivity?
   - Are growth rates, margins, and multiples reasonable vs. industry benchmarks?
   - Flag any hockey-stick projections without supporting evidence

5. **Financial Accuracy**:
   - Do totals foot and cross-foot?
   - Are periods consistent (quarterly vs. annual mixing)?
   - Currency, rounding, and unit consistency
   - Balance sheet balances? Cash flow reconciles?

6. **Presentation & Usability**:
   - Formatting, number formats, conditional formatting
   - Print areas and headers set?
   - Color coding for inputs vs. calculations vs. outputs?

7. **Verdict & Recommendations**:
   - Model reliability rating: High / Medium / Low
   - Key risks in the model
   - Specific cells or sections needing correction

### For Word Documents (.docx)

1. **Document Purpose & Audience**: What is this document and who is it for?

2. **Content Quality Assessment**:
   - Is the executive summary compelling and accurate?
   - Are arguments well-structured and supported by evidence?
   - Is the level of detail appropriate for the audience?

3. **Financial & Operational Claims**:
   - Verify any financial figures cited
   - Check that metrics are defined consistently
   - Flag vague language where specifics are needed
   - Ensure forward-looking statements have appropriate caveats

4. **Legal & Compliance Considerations**:
   - Flag potential disclosure issues
   - Note any statements that could create contractual obligations
   - Identify areas needing legal review

5. **Strategic & Competitive Context**:
   - Does this reflect current market reality?
   - Are competitive references accurate and fair?
   - Is the positioning consistent with company messaging?

6. **Writing Quality**:
   - Clarity, conciseness, and tone
   - Jargon appropriate for audience?
   - Grammar and formatting consistency

7. **Verdict & Recommendations**:
   - Overall quality: Strong / Adequate / Needs Revision
   - Top issues to address
   - Specific paragraph-level feedback where needed

## Output Format

Always structure your review as:

```
## 📋 DOCUMENT REVIEW: [Document Name]
**Type:** [PowerPoint/Excel/Word]
**Date Reviewed:** [Today's date]
**Reviewer Perspective:** CFO/COO

### 🎯 Executive Summary
[2-3 sentence overview of the document and your top-line assessment]

### ✅ Strengths
- [What's working well]

### ⚠️ Issues & Concerns
- [Critical] [Issue description]
- [Important] [Issue description]
- [Minor] [Issue description]

### 📊 Detailed Analysis
[Section-by-section analysis per the framework above]

### 🔧 Recommendations
1. [Most important fix]
2. [Second priority]
3. [Third priority]

### 📈 Verdict: [READY / NEEDS WORK / MAJOR REVISION]
[One paragraph final assessment with clear next steps]
```

## Reference Files

For additional context on Cohesity-specific metrics and industry benchmarks,
see the reference files in this skill directory:
- `cohesity-context.md` — Company background, key metrics, strategic priorities
- `review-checklist.md` — Quick-reference checklist for common review scenarios
- `industry-benchmarks.md` — Data protection market benchmarks and peer comps
'''

# Add as a styled code block
p = doc.add_paragraph()
p.style = doc.styles['Normal']
run = p.add_run(skill_md)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_page_break()

# ============================================================
# SECTION 2: REFERENCE FILE - cohesity-context.md
# ============================================================
doc.add_heading('Section 2: cohesity-context.md — Reference File (Optional)', level=1)
doc.add_paragraph('Place this file in the same folder as SKILL.md. Claude loads it on demand when it needs Cohesity-specific context.')

context_md = '''# Cohesity Context — CFO Review Reference

## Company Overview
- **Industry:** Enterprise Data Security & Management
- **Business Model:** Subscription SaaS + appliance (transitioning to cloud-native)
- **Key Product:** Cohesity Data Cloud — unified platform for backup, security,
  governance, and analytics
- **Competitors:** Rubrik (RBRK), Commvault (CVLT), Veeam (private/Insight Partners),
  Veritas, Dell (APEX Backup)

## Key Metrics to Watch
- **ARR & ARR Growth** (YoY) — primary top-line metric
- **Net Revenue Retention (NRR)** — target >120%
- **Subscription Revenue %** — transition from perpetual/appliance
- **Gross Margin** — GAAP and non-GAAP; target >75% for software
- **Operating Margin** — path to profitability benchmarks
- **FCF Margin** — cash generation efficiency
- **Rule of 40** — ARR growth % + FCF margin %
- **CAC Payback Period** — GTM efficiency
- **Magic Number** — net new ARR / prior period S&M spend

## Strategic Priorities (Current)
1. Cloud-native adoption and Data Cloud platform expansion
2. AI/ML-powered data insights and ransomware detection
3. International expansion (EMEA, APJ)
4. Partner ecosystem growth (hyperscalers, GSIs, MSPs)
5. Operational efficiency and path to sustained profitability

## Financial Review Standards
- All projections must include clearly stated assumptions
- Revenue recognition must follow ASC 606 guidelines
- Non-GAAP adjustments must be reconciled to GAAP
- Forward-looking statements require appropriate safe harbor language
- Board materials must be auditor-ready
'''

p = doc.add_paragraph()
run = p.add_run(context_md)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_page_break()

# ============================================================
# SECTION 3: REFERENCE FILE - review-checklist.md
# ============================================================
doc.add_heading('Section 3: review-checklist.md — Quick Reference (Optional)', level=1)
doc.add_paragraph('A fast-reference checklist Claude uses for rapid document triage.')

checklist_md = '''# CFO Document Review — Quick Checklist

## Universal Checks (All Document Types)
- [ ] Numbers internally consistent throughout?
- [ ] Sources cited for external data?
- [ ] Appropriate for intended audience?
- [ ] Aligned with current company messaging?
- [ ] Legal/compliance review needed?
- [ ] Confidentiality markings correct?

## Board Deck Specific
- [ ] Clear ask within first 3 slides?
- [ ] Financial summary slide with key metrics?
- [ ] Competitive context included?
- [ ] Risk section present and honest?
- [ ] Appendix with supporting detail?
- [ ] Consistent with prior board updates?

## Financial Model Specific
- [ ] Assumptions tab clearly labeled?
- [ ] Sensitivity analysis included?
- [ ] All formulas auditable (no hardcoded overrides)?
- [ ] Historical actuals reconcile to reported numbers?
- [ ] Currency and units consistent?
- [ ] Model version and date tracked?

## Vendor / Partnership Proposals
- [ ] Total cost of ownership calculated?
- [ ] Competitive alternatives evaluated?
- [ ] Contract terms summarized?
- [ ] Integration requirements assessed?
- [ ] ROI timeline realistic?

## Quarterly Business Reviews (QBRs)
- [ ] Actuals vs. plan variance explained?
- [ ] Forward guidance grounded in pipeline data?
- [ ] Headcount and hiring plan aligned with budget?
- [ ] Key initiative status updated?
- [ ] Customer metrics (NRR, churn, expansion) current?

## M&A / Investment Materials
- [ ] Valuation methodology appropriate?
- [ ] Comparable transactions listed?
- [ ] Synergy assumptions conservative?
- [ ] Integration plan realistic?
- [ ] Due diligence gaps flagged?
'''

p = doc.add_paragraph()
run = p.add_run(checklist_md)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_page_break()

# ============================================================
# SECTION 4: REFERENCE FILE - industry-benchmarks.md
# ============================================================
doc.add_heading('Section 4: industry-benchmarks.md — Market Benchmarks (Optional)', level=1)
doc.add_paragraph('Industry benchmarks Claude references when evaluating financial claims and projections.')

benchmarks_md = '''# Data Protection Industry Benchmarks

## Market Size & Growth
- Global data protection market: ~$15B (2025), growing ~12-15% CAGR
- Cloud backup/DR segment: fastest growth at ~20%+ CAGR
- Ransomware resilience: emerging category, high urgency

## Peer Financial Benchmarks (Enterprise SaaS / Data Protection)

| Metric                  | Top Quartile | Median   | Context               |
|-------------------------|-------------|----------|-----------------------|
| ARR Growth (YoY)        | >30%        | 20-25%   | At scale (>$500M ARR) |
| Gross Margin (Non-GAAP) | >80%        | 72-78%   | Software-heavy mix    |
| Operating Margin         | >10%        | -5% to 5%| Growth-stage          |
| FCF Margin              | >15%        | 5-10%    | At scale              |
| NRR                     | >130%       | 115-125% | Enterprise SaaS       |
| Rule of 40              | >50         | 35-45    | Growth + profitability |
| CAC Payback (months)    | <18         | 18-24    | Enterprise sales      |
| S&M % of Revenue        | <35%        | 40-50%   | Efficiency target     |
| R&D % of Revenue        | 15-20%      | 20-25%   | Innovation balance    |

## Competitor Quick Reference

### Rubrik (RBRK)
- Public (IPO 2024), pure-play data security
- Focus: Zero Trust Data Security, ransomware recovery
- Strengths: Cloud-native architecture, security positioning
- Watch: ARR growth trajectory, path to profitability

### Commvault (CVLT)
- Public, legacy player modernizing
- Focus: Cloud repatriation, hybrid data management
- Strengths: Installed base, breadth of platform
- Watch: Metallic (SaaS) traction, margin expansion

### Veeam (Private — Insight Partners)
- Largest private player, potential IPO candidate
- Focus: Mid-market to enterprise backup, M365 protection
- Strengths: Ease of use, channel dominance, growth rate
- Watch: Enterprise up-market motion, competitive pricing

## Valuation Benchmarks
- Revenue multiples: 6-12x NTM revenue (high-growth SaaS)
- ARR multiples: 8-15x for >25% growers
- EV/FCF: 25-40x for profitable growers
- Discount for private: typically 20-30% vs. public comps
'''

p = doc.add_paragraph()
run = p.add_run(benchmarks_md)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_page_break()

# ============================================================
# SECTION 5: USAGE EXAMPLES
# ============================================================
doc.add_heading('Section 5: Usage Examples', level=1)
doc.add_paragraph('Once installed, here are example prompts that will activate this skill:')

examples = [
    ('"Review this board deck for next week\'s meeting"',
     'Activates the PowerPoint review framework. Drop the .pptx in your Cowork folder and Claude will do a full CFO-perspective analysis.'),
    ('"Analyze the FY26 Q3 financial model attached to the FP&A email"',
     'Activates the Excel review framework. Claude will audit formulas, stress-test assumptions, and verify the numbers foot.'),
    ('"Review this vendor proposal from Salesforce"',
     'Activates the Word document review with the vendor/partnership checklist from the reference files.'),
    ('"Give me a CFO review of all three attachments from the strategy meeting"',
     'Claude will process multiple documents in parallel using sub-agents, reviewing each with the appropriate framework.'),
    ('"Quick triage — is this deck ready for the board?"',
     'Uses the quick checklist for a rapid pass/fail assessment with top issues flagged.'),
]

for prompt, desc in examples:
    p = doc.add_paragraph()
    run = p.add_run(prompt)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
    doc.add_paragraph(desc)
    doc.add_paragraph()

# ============================================================
# SECTION 6: TIPS
# ============================================================
doc.add_heading('Section 6: Tips for Best Results', level=1)

tips = [
    'Work in a dedicated folder: Create a folder like ~/Documents/CFO-Reviews/ and point Cowork at it. Drop email attachments there before asking for review.',
    'Connect Microsoft 365: If you enable the M365 connector in Cowork, Claude can pull attachments directly from your Outlook/OneDrive without manual download.',
    'Customize the benchmarks: Update industry-benchmarks.md quarterly with the latest peer company earnings data for more accurate comparisons.',
    'Chain with email: Ask Claude to "review this deck and draft a reply email with my feedback" — it can create the review AND compose a response.',
    'Batch reviews: Drop multiple documents and say "review all files in this folder for the Monday exec meeting" — Cowork will parallelize using sub-agents.',
    'Keep SKILL.md under 500 lines: If you add more review frameworks, split them into separate reference .md files rather than bloating the main skill file.',
]

for i, tip in enumerate(tips, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    p.add_run(tip)

# Save
output_path = '/Users/ericbrown/.openclaw/workspace/CFO_CoWork_Skill_Document_Review.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
