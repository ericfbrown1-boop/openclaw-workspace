#!/usr/bin/env python3
"""Create Zoom Competitive Intelligence Report."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('Competitive Intelligence Report', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Zoom Communications, Inc. (NASDAQ: ZM)')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x24, 0x40, 0xB3)
run.bold = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Prepared for Eric Brown, CFO & COO, Cohesity\n').bold = True
meta.add_run(f'Date: {datetime.date.today().strftime("%B %d, %Y")}\n')
meta.add_run('Classification: Confidential — Internal Use Only\n')
meta.add_run('Sources: SEC filings, earnings calls, analyst reports, company website, news\n')
meta.add_run('Prepared by: Jarvis (Project Scraper)')

doc.add_page_break()

# ==========================================================
# TABLE OF CONTENTS
# ==========================================================
doc.add_heading('Table of Contents', level=1)
toc = [
    '1. Executive Summary',
    '2. Company Overview',
    '3. Financial Performance (FY2026)',
    '4. Product Portfolio & Strategy',
    '5. AI Strategy & Monetization',
    '6. Competitive Landscape',
    '7. Leadership Team',
    '8. Valuation & Analyst Sentiment',
    '9. SWOT Analysis',
    '10. Strategic Implications for Cohesity',
]
for item in toc:
    doc.add_paragraph(item)
doc.add_page_break()

# ==========================================================
# 1. EXECUTIVE SUMMARY
# ==========================================================
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'Zoom Communications (ZM) has completed its transformation from a pandemic-era video conferencing '
    'provider into an AI-first unified communications platform. The company reported FY2026 revenue of '
    '$4.87B (+4.4% YoY), with Q4 revenue of $1.247B (+5.3% YoY) — its strongest quarterly growth rate '
    'of the fiscal year. While the top-line growth is modest, Zoom\'s financial discipline is exceptional: '
    '$1.9B in free cash flow, $7.8B in cash reserves, and a completed $2.7B share buyback program.'
)
doc.add_paragraph(
    'The most significant strategic development is Zoom\'s AI monetization story, which is materializing '
    'through its Contact Center (CCaaS) business rather than its core meetings platform. Zoom Contact Center '
    'ARR is growing in "high double digits," with 100% of Q4\'s top 10 deals including paid AI components '
    'and 7 of 10 representing competitive displacements of legacy vendors. The Custom AI Companion '
    '($12/user/month vs. Microsoft Copilot at $30) positions Zoom as the value play for enterprise AI adoption.'
)
doc.add_paragraph(
    'FY2027 guidance calls for revenue exceeding $5B (EPS: $5.77-$5.81), implying ~3-4% growth. '
    'The stock trades at ~$78 (down 18% post-earnings), with a consensus analyst target of $95.32 '
    '(~22% upside). The investment thesis hinges on whether AI Companion and Contact Center can '
    'accelerate growth beyond low-single-digits.'
)

p = doc.add_paragraph()
p.add_run('Key Takeaway: ').bold = True
p.add_run('Zoom is no longer just a meetings company. Its AI-first platform strategy, fortress balance '
    'sheet, and aggressive CCaaS expansion make it a relevant player in the broader enterprise software '
    'ecosystem — including potential intersection with Cohesity\'s data security and management platform.')

doc.add_page_break()

# ==========================================================
# 2. COMPANY OVERVIEW
# ==========================================================
doc.add_heading('2. Company Overview', level=1)

overview_data = [
    ('Legal Name', 'Zoom Communications, Inc. (fka Zoom Video Communications)'),
    ('Ticker', 'NASDAQ: ZM'),
    ('Founded', '2011 by Eric Yuan (former Cisco/WebEx VP Engineering)'),
    ('Headquarters', 'San Jose, California'),
    ('Employees', '~8,000 (estimated)'),
    ('Market Cap', '~$24B (as of March 2026)'),
    ('Fiscal Year End', 'January 31'),
    ('CEO', 'Michelle Chang (Eric Yuan transitioned to Chairman/President)'),
    ('Share Structure', 'Dual-class: Class A (1 vote), Class B (10 votes)'),
    ('Revenue (FY2026)', '$4.87B'),
    ('Cash Position', '$7.8B (cash + marketable securities)'),
]

table = doc.add_table(rows=len(overview_data)+1, cols=2)
table.style = 'Light Grid Accent 1'
table.cell(0, 0).text = 'Attribute'
table.cell(0, 1).text = 'Detail'
for i, (attr, val) in enumerate(overview_data, 1):
    table.cell(i, 0).text = attr
    table.cell(i, 1).text = val

doc.add_paragraph()
doc.add_heading('Company History & Evolution', level=2)
doc.add_paragraph(
    'Zoom was founded in 2011 by Eric Yuan, who left Cisco (where he was VP of Engineering for WebEx) '
    'after becoming frustrated with the quality of existing video conferencing solutions. The company '
    'went public in 2019 (IPO at $36/share) and experienced explosive growth during COVID-19, reaching '
    'a peak market cap of ~$160B in October 2020.'
)
doc.add_paragraph(
    'Post-pandemic, Zoom underwent a painful normalization (revenue declined from $4.39B in FY2023 to '
    '$4.53B in FY2024) as hybrid work patterns stabilized. The company pivoted aggressively toward: '
    '(1) AI-powered features, (2) platform expansion beyond meetings, and (3) enterprise upmarket motion. '
    'FY2026 marks the first year of re-accelerating revenue growth after the post-pandemic trough.'
)

doc.add_heading('Business Model: "Land and Expand"', level=2)
doc.add_paragraph(
    'Zoom\'s core business model is a freemium "land and expand" strategy:'
)
land_expand = [
    ('Land', 'Free tier (40-min meetings) creates grassroots adoption within organizations, bypassing IT procurement. Small department contracts for enterprise targets.'),
    ('Upsell', 'Free → Pro → Business → Enterprise tiers as usage scales. Meeting limits, admin controls, and security features drive upgrades.'),
    ('Cross-sell', 'Meetings-only customers expand to Zoom Phone (10M+ seats), Contact Center (1,100+ customers), Workvivo (1,225+ customers), and AI Companion.'),
    ('Monetize AI', 'Base AI Companion is free (retention play). Custom AI Companion ($12/user/month) drives incremental revenue from enterprise customers.'),
]
for phase, desc in land_expand:
    p = doc.add_paragraph()
    run = p.add_run(f'{phase}: ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ==========================================================
# 3. FINANCIAL PERFORMANCE
# ==========================================================
doc.add_heading('3. Financial Performance (FY2026)', level=1)

doc.add_heading('Revenue Summary', level=2)

# Revenue table
rev_table = doc.add_table(rows=6, cols=4)
rev_table.style = 'Light Grid Accent 1'
headers = ['Metric', 'Q4 FY2026', 'FY2026', 'YoY Change']
for i, h in enumerate(headers):
    rev_table.cell(0, i).text = h

rev_data = [
    ('Total Revenue', '$1,247M', '$4,869M', '+5.3% / +4.4%'),
    ('Enterprise Revenue', '~$770M (est.)', '~$2,990M (est.)', '+6.1% (Q3 rate)'),
    ('Online Revenue', '~$477M (est.)', '~$1,879M (est.)', '+2.0% (Q3 rate)'),
    ('Non-GAAP EPS', '$1.44', '$5.95-$5.97', 'Missed by $0.02'),
    ('Free Cash Flow', '~$475M (est.)', '$1,900M', '+30% YoY (Q3 pace)'),
]
for i, (metric, q4, fy, yoy) in enumerate(rev_data, 1):
    rev_table.cell(i, 0).text = metric
    rev_table.cell(i, 1).text = q4
    rev_table.cell(i, 2).text = fy
    rev_table.cell(i, 3).text = yoy

doc.add_paragraph()
doc.add_heading('Key Financial Highlights', level=2)

highlights = [
    'Revenue re-acceleration: Q4 growth of 5.3% was the fastest quarter of FY2026, up from ~3-4% in H1.',
    'Fortress balance sheet: $7.8B in cash & marketable securities — more than the entire annual revenue.',
    'FCF machine: ~$1.9B in free cash flow represents a ~39% FCF margin — exceptional for a $5B revenue company.',
    'Share buybacks: Completed $2.7B buyback program in FY2025-FY2026, reducing dilution.',
    'Online churn at historic low: 2.7% monthly churn in FY2026 (down from 3.8% in 2022), attributed to AI Companion stickiness.',
    'Enterprise growing faster than online: Enterprise revenue +6.1% vs. online +2.0% in Q3, showing upmarket traction.',
    'Rare EPS miss: Q4 non-GAAP EPS of $1.44 missed consensus of $1.46 by 1.37% — first miss in 7 quarters. Small but symbolic.',
    'Strategic investment gains: ~$970M net gain on strategic investments (unrealized) — impressive but non-operational.',
]
for h in highlights:
    doc.add_paragraph(h, style='List Bullet')

doc.add_heading('FY2027 Guidance', level=2)
guidance_table = doc.add_table(rows=4, cols=3)
guidance_table.style = 'Light Grid Accent 1'
guidance_table.cell(0, 0).text = 'Metric'
guidance_table.cell(0, 1).text = 'FY2027 Guidance'
guidance_table.cell(0, 2).text = 'Implied Growth'
g_data = [
    ('Revenue', '>$5.0B', '~3-4% YoY'),
    ('Non-GAAP EPS', '$5.77-$5.81', 'Slight decline from $5.95-$5.97'),
    ('Q1 FY2027 EPS', '$1.40-$1.42', 'vs. consensus $0.86 (different basis)'),
]
for i, (m, g, imp) in enumerate(g_data, 1):
    guidance_table.cell(i, 0).text = m
    guidance_table.cell(i, 1).text = g
    guidance_table.cell(i, 2).text = imp

doc.add_page_break()

# ==========================================================
# 4. PRODUCT PORTFOLIO
# ==========================================================
doc.add_heading('4. Product Portfolio & Strategy', level=1)

doc.add_heading('Platform Architecture: "AI-First Work Platform"', level=2)
doc.add_paragraph(
    'Zoom has evolved from a single-product video conferencing tool into a multi-product platform. '
    'The company now positions itself as an "AI-first work platform for human connection" with two core pillars: '
    'Zoom Workplace (internal collaboration) and Zoom Business Services (customer-facing).'
)

products = [
    ('Zoom Workplace', 'Unified collaboration suite',
     ['Zoom Meetings — Core video conferencing (still the #1 brand in video)',
      'Zoom Phone — Cloud PBX with 10M+ paid seats (launched 2019, fastest-growing segment)',
      'Zoom Chat — Persistent messaging and channels',
      'Zoom Rooms — Conference room hardware/software',
      'Zoom Whiteboard — Digital collaboration canvas',
      'Zoom Clips — Async video messaging',
      'Zoom Docs — AI-powered documents (newer offering)',
      'Workvivo — Employee experience platform (1,225+ customers, ~70-80% YoY growth)']),
    ('Zoom Business Services', 'Customer-facing solutions',
     ['Zoom Contact Center (CCaaS) — AI-first contact center, 1,100+ customers, doubled YoY',
      'Zoom Virtual Agent (ZVA) — AI-powered self-service, included in 4 of top 10 CX deals',
      'Zoom Revenue Accelerator — Conversation intelligence for sales teams',
      'Zoom Events — Virtual and hybrid event platform']),
    ('Zoom AI Companion', 'Cross-platform AI layer',
     ['Free tier: Meeting summaries, chat compose, smart recordings — retention play',
      'Custom AI Companion ($12/user/month): Enterprise workflows, third-party integrations, knowledge retrieval',
      'AI Companion 3.0: Agentic capabilities, cross-app task execution',
      '4x adoption increase YoY; MAU tripled YoY in Q4',
      '100% of top 10 enterprise deals included AI implementation strategies']),
]

for name, subtitle_text, features in products:
    doc.add_heading(name, level=3)
    p = doc.add_paragraph()
    p.add_run(subtitle_text).italic = True
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ==========================================================
# 5. AI STRATEGY
# ==========================================================
doc.add_heading('5. AI Strategy & Monetization', level=1)

doc.add_paragraph(
    'Zoom\'s AI strategy is the most important element of its investment thesis. The company has taken '
    'a differentiated approach compared to Microsoft and Google:'
)

doc.add_heading('The Zoom AI Playbook', level=2)
ai_strategy = [
    ('Free AI = Defensive Moat', 'Unlike Microsoft ($30/user Copilot) or Google ($30/user Gemini), '
     'Zoom offers its base AI Companion for free to all paid users. This is a deliberate retention play '
     'that has driven monthly churn to historic lows (2.7%). The free AI makes Zoom stickier than ever.'),
    ('Custom AI = Offensive Revenue', 'Custom AI Companion at $12/user/month is the monetization engine. '
     'It offers enterprise-grade features: third-party app integrations, custom knowledge bases, workflow '
     'automation across disparate systems. At $12 vs. Microsoft\'s $30, it\'s the value play for '
     'budget-conscious IT departments.'),
    ('Contact Center = AI\'s Killer App', 'The real AI monetization is happening in CCaaS, not meetings. '
     'Zoom Virtual Agent 3.0 automates end-to-end customer resolution with multi-step workflows. '
     'Every single top-10 deal in Q4 included paid AI. ZCX ARR is growing in "high double digits."'),
    ('"System of Action" Vision', 'CEO Michelle Chang and founder Eric Yuan describe Zoom\'s evolution from a '
     '"system of engagement" (meetings/calls) to a "system of action" that automates workflows across '
     'the entire customer journey — from internal collaboration to external customer interactions.'),
]
for title_text, desc in ai_strategy:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('AI Competitive Positioning', level=2)
ai_table = doc.add_table(rows=5, cols=4)
ai_table.style = 'Light Grid Accent 1'
ai_headers = ['Company', 'AI Product', 'Pricing', 'Approach']
for i, h in enumerate(ai_headers):
    ai_table.cell(0, i).text = h
ai_data = [
    ('Zoom', 'AI Companion / Custom', 'Free / $12/user/mo', 'Free base + paid enterprise tier'),
    ('Microsoft', 'Copilot', '$30/user/mo', 'Premium-only, bundled with M365'),
    ('Google', 'Gemini for Workspace', '$30/user/mo', 'Premium-only add-on'),
    ('Cisco', 'Webex AI Assistant', 'Included', 'Bundled with Webex Suite'),
]
for i, (co, prod, price, approach) in enumerate(ai_data, 1):
    ai_table.cell(i, 0).text = co
    ai_table.cell(i, 1).text = prod
    ai_table.cell(i, 2).text = price
    ai_table.cell(i, 3).text = approach

doc.add_page_break()

# ==========================================================
# 6. COMPETITIVE LANDSCAPE
# ==========================================================
doc.add_heading('6. Competitive Landscape', level=1)

doc.add_paragraph(
    'Zoom competes across multiple segments: UCaaS (unified communications), CCaaS (contact center), '
    'and increasingly in the broader enterprise productivity space.'
)

doc.add_heading('UCaaS Market', level=2)
doc.add_paragraph(
    'The global UCaaS market is valued at ~$37B (2026), growing at ~13% CAGR. Key competitors:'
)

competitors = [
    ('Microsoft Teams', '$60B+ Office/M365 revenue', 'Dominant in enterprise due to M365 bundle. '
     '320M+ monthly active users. Zoom is losing some enterprise seats to bundled Teams but winning '
     'on quality and AI value. Two major US financial institutions displaced Teams for Zoom Phone in Q4.'),
    ('Cisco Webex', '~$5B Collaboration segment', 'Legacy enterprise player, strong in hardware/rooms. '
     'Webex AI Assistant included free. Zoom displaced Cisco calling in major bank deal (50K seats).'),
    ('RingCentral (RNG)', '$2.4B ARR', 'Pure-play UCaaS leader. Strong in mid-market. '
     'Partner with Avaya, Mitel. Less AI differentiation than Zoom.'),
    ('8x8', '~$730M revenue', 'Mid-market UCaaS/CCaaS combo. Smaller scale, less AI investment.'),
]
for name, size, analysis in competitors:
    p = doc.add_paragraph()
    run = p.add_run(f'{name} ({size}): ')
    run.bold = True
    p.add_run(analysis)

doc.add_heading('CCaaS Market', level=2)
doc.add_paragraph(
    'Zoom Contact Center is the fastest-growing segment and primary AI monetization vehicle. Competitors include:'
)
ccaas = [
    ('NICE CXone', 'Market leader in cloud CCaaS. Zoom is winning competitive displacements.'),
    ('Genesys', 'Strong in enterprise. Private (Permira). Zoom displaced in 7 of top 10 deals.'),
    ('Five9 (FIVN)', 'Cloud-native CCaaS. $1B revenue. Microsoft attempted acquisition (failed).'),
    ('Amazon Connect', 'AWS-native, pay-per-use. Strong in AWS-centric shops.'),
    ('Talkdesk', 'AI-forward startup. Competing for same "next-gen CCaaS" positioning as Zoom.'),
]
for name, analysis in ccaas:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(analysis)

doc.add_page_break()

# ==========================================================
# 7. LEADERSHIP
# ==========================================================
doc.add_heading('7. Leadership Team', level=1)

leaders = [
    ('Eric Yuan', 'Founder, Chairman & President', 'Founded Zoom in 2011 after 14 years at Cisco/WebEx. '
     'Holds Class B shares with 10x voting power, maintaining significant control. '
     'Transitioned CEO role to Michelle Chang but remains deeply involved in product/AI vision.'),
    ('Michelle Chang', 'CEO (appointed 2025)', 'Former CFO of Zoom. Took over CEO role as Yuan '
     'shifted focus to product and AI strategy. Background in finance brings operational discipline.'),
    ('Kelly Steckelberg', 'Former CFO', 'Led Zoom through IPO and pandemic growth. Transitioned out '
     'as Michelle Chang moved to CEO.'),
]
for name, title_text, bio in leaders:
    p = doc.add_paragraph()
    run = p.add_run(f'{name} — {title_text}')
    run.bold = True
    doc.add_paragraph(bio)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Board Note: ').bold = True
p.add_run('Experienced board with average tenure of 6.9 years. Dual-class share structure gives '
    'Eric Yuan effective control over all major corporate decisions despite transitioning the CEO title.')

doc.add_page_break()

# ==========================================================
# 8. VALUATION
# ==========================================================
doc.add_heading('8. Valuation & Analyst Sentiment', level=1)

doc.add_heading('Current Valuation', level=2)
val_table = doc.add_table(rows=8, cols=2)
val_table.style = 'Light Grid Accent 1'
val_table.cell(0, 0).text = 'Metric'
val_table.cell(0, 1).text = 'Value'
val_data = [
    ('Stock Price', '~$78.05 (March 6, 2026)'),
    ('Market Cap', '~$24B'),
    ('52-Week High', '$96.22 (January 2026)'),
    ('52-Week Low', '~$58'),
    ('EV/Revenue (NTM)', '~4.5x'),
    ('EV/FCF', '~12x'),
    ('Consensus Target', '$95.32 (~22% upside)'),
]
for i, (m, v) in enumerate(val_data, 1):
    val_table.cell(i, 0).text = m
    val_table.cell(i, 1).text = v

doc.add_paragraph()
doc.add_heading('Analyst Ratings', level=2)
ratings = [
    ('Strong Buy', '5 analysts (29%)'),
    ('Buy', '3 analysts (18%)'),
    ('Hold', '8 analysts (47%)'),
    ('Sell', '0 analysts (0%)'),
    ('Strong Sell', '1 analyst (6%)'),
]
for rating, count in ratings:
    p = doc.add_paragraph()
    run = p.add_run(f'{rating}: ')
    run.bold = True
    p.add_run(count)

doc.add_paragraph()
doc.add_heading('Notable Analyst Actions (Post-Earnings)', level=2)
actions = [
    'BTIG: Cut target to $100, maintains Buy — notes AI monetization story intact',
    'KeyBanc: Raised target to $74 from $69, maintains Underweight — skeptical on growth re-acceleration',
    'Consensus: Average target $95.32, range $74-$115',
]
for a in actions:
    doc.add_paragraph(a, style='List Bullet')

doc.add_page_break()

# ==========================================================
# 9. SWOT
# ==========================================================
doc.add_heading('9. SWOT Analysis', level=1)

doc.add_heading('Strengths', level=2)
strengths = [
    'Brand recognition: "Zoom" is a verb — strongest brand in video conferencing globally',
    'Financial fortress: $7.8B cash, $1.9B FCF, no debt — enormous financial flexibility',
    'AI value positioning: Free AI Companion + $12/user Custom tier undercuts Microsoft 2.5x',
    'Platform breadth: Meetings + Phone + Contact Center + Workvivo = full collaboration stack',
    'Contact Center momentum: 1,100+ customers, doubled YoY, winning competitive displacements',
    'Low churn: Historic-low 2.7% monthly churn demonstrates platform stickiness',
    'Zoom Phone scale: 10M+ paid seats — credible voice platform',
]
for s in strengths:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('Weaknesses', level=2)
weaknesses = [
    'Slow top-line growth: 4.4% FY2026 revenue growth is below SaaS peer averages',
    'Post-pandemic overhang: Stock still 85% below all-time high; investor fatigue persists',
    'Dual-class structure: Eric Yuan controls voting — limits shareholder influence',
    'Online segment stagnation: Consumer/SMB revenue growing only 2% — limited upside',
    'EPS miss: Q4 miss, even at 1.37%, breaks the consistency narrative',
    'Limited international diversification: Primarily US/North America revenue base',
]
for w in weaknesses:
    doc.add_paragraph(w, style='List Bullet')

doc.add_heading('Opportunities', level=2)
opportunities = [
    'AI monetization inflection: Custom AI Companion and ZVA could drive meaningful incremental revenue',
    'CCaaS market expansion: Contact center is a $30B+ TAM with legacy vendors ripe for displacement',
    'Workvivo/Meta Workplace migration: Meta sunsetting Workplace → Workvivo is preferred migration path',
    'Platform consolidation: Enterprises want fewer vendors — Zoom can replace Teams + legacy PBX + CCaaS',
    'Share buybacks: With $7.8B cash, aggressive buybacks can drive EPS growth even on modest revenue growth',
    'M&A optionality: Balance sheet supports strategic acquisitions (data security? analytics?)',
]
for o in opportunities:
    doc.add_paragraph(o, style='List Bullet')

doc.add_heading('Threats', level=2)
threats = [
    'Microsoft bundling: Teams included "free" with M365 makes it hard to win net-new enterprise seats',
    'AI commoditization: If meeting summaries become table-stakes across all platforms, Zoom loses differentiation',
    'Economic downturn: IT budget cuts hit discretionary UCaaS spend; enterprises may consolidate to Microsoft',
    'Revenue growth ceiling: If AI doesn\'t accelerate growth beyond 4-5%, multiple expansion is limited',
    'Competition from hyperscalers: Amazon (Connect), Google (Meet/CCAI) have deeper pockets',
    'Founder control risk: Dual-class structure means strategic direction depends heavily on Yuan\'s vision',
]
for t in threats:
    doc.add_paragraph(t, style='List Bullet')

doc.add_page_break()

# ==========================================================
# 10. IMPLICATIONS FOR COHESITY
# ==========================================================
doc.add_heading('10. Strategic Implications for Cohesity', level=1)

doc.add_heading('Why Zoom Matters to Cohesity', level=2)
doc.add_paragraph(
    'While Zoom is not a direct competitor to Cohesity, there are several strategic dimensions worth monitoring:'
)

implications = [
    ('Partnership Opportunity', 'Zoom\'s "Custom AI Companion" integrates with third-party enterprise '
     'applications. As Zoom evolves into a "system of action," there may be opportunities for Cohesity\'s '
     'data security and management platform to integrate with Zoom\'s enterprise workflow layer — particularly '
     'around data governance, compliance recording, and secure data handling for Zoom Contact Center recordings.'),
    ('Competitive Intelligence Template', 'Zoom\'s AI monetization playbook (free base + paid enterprise tier) '
     'is a model worth studying for Cohesity\'s own AI-powered features. The $12 vs. $30 pricing strategy '
     'against Microsoft demonstrates how to win on value in an AI-commoditized market.'),
    ('Enterprise Buyer Dynamics', 'Zoom\'s success displacing Microsoft Teams and Cisco in enterprise accounts '
     'shows that best-of-breed can still win against bundled incumbents when the product is meaningfully better. '
     'This is relevant for Cohesity\'s positioning against bundled backup/recovery from Dell, HPE, and others.'),
    ('Customer Overlap', 'Zoom\'s enterprise customer base (particularly financial services and healthcare) '
     'overlaps significantly with Cohesity\'s target market. Zoom Contact Center customers generating '
     'massive volumes of customer interaction data need robust data management and compliance solutions.'),
    ('M&A Watch', 'With $7.8B in cash and a "system of action" vision, Zoom could acquire in adjacent '
     'spaces (data analytics, security, compliance). Monitor for moves that could bring Zoom closer to '
     'or further from Cohesity\'s competitive orbit.'),
    ('Financial Benchmarking', 'Zoom\'s 39% FCF margin and $1.9B FCF on $4.9B revenue sets a high bar for '
     'financial discipline in enterprise SaaS. Useful as a benchmarking reference for Cohesity\'s own '
     'path to profitability and FCF generation.'),
]
for title_text, desc in implications:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('— End of Report —').italic = True
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Prepared by Jarvis using Project Scraper | Confidential').font.size = Pt(9)

# Save
output_path = '/Users/ericbrown/ProjectScraper/Zoom_Competitive_Intelligence_Report.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
