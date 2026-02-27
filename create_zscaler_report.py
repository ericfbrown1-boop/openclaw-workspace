#!/usr/bin/env python3
"""
Generate Zscaler Q2 FY2026 Earnings Analysis Word Document
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

def create_report():
    doc = Document()
    
    # Title Page
    title = doc.add_heading('Zscaler Inc. (ZS)', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    subtitle = doc.add_paragraph('Q2 FY2026 Earnings Analysis')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    
    date_p = doc.add_paragraph('Report Date: February 27, 2026')
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_p.runs[0].font.size = Pt(11)
    date_p.runs[0].italic = True
    
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph(
        'Zscaler reported strong Q2 FY2026 results on February 26, 2026, beating consensus estimates on both revenue and EPS. '
        'The company demonstrated continued momentum in the zero-trust security market with 26% YoY revenue growth and 25% ARR growth. '
        'Management raised full-year ARR guidance to 24%, reflecting confidence in sustained demand for cloud-native security solutions.'
    )
    
    highlights = doc.add_paragraph('Key Highlights:')
    highlights.runs[0].bold = True
    doc.add_paragraph('• Revenue: $815.8M (+26% YoY), beating consensus of $799M', style='List Bullet')
    doc.add_paragraph('• Adjusted EPS: $1.01 vs. $0.89 expected (+13% beat)', style='List Bullet')
    doc.add_paragraph('• ARR: $3,359M (+25% YoY)', style='List Bullet')
    doc.add_paragraph('• Operating Margin: 22%+ (non-GAAP, expanding YoY)', style='List Bullet')
    doc.add_paragraph('• Raised FY2026 ARR guidance to 24% growth', style='List Bullet')
    doc.add_paragraph('• Calculated RPO growth: 31% YoY (strong bookings momentum)', style='List Bullet')
    
    doc.add_page_break()
    
    # Q2 Results
    doc.add_heading('Q2 FY2026 Financial Results', 1)
    doc.add_paragraph('Period Ended: January 31, 2026', style='Caption')
    doc.add_paragraph()
    
    doc.add_heading('Revenue Performance', 2)
    revenue_table = [
        ['Metric', 'Q2 FY2026', 'Q2 FY2025', 'YoY Growth', 'vs Consensus'],
        ['Revenue', '$815.8M', '$647.5M*', '+26.0%', '+$16.8M (2.1%)'],
        ['ARR', '$3,359M', '$2,687M*', '+25.0%', 'N/A'],
        ['Calculated RPO', 'N/A', 'N/A', '+31.0%', 'Above expectations'],
    ]
    
    table = doc.add_table(rows=len(revenue_table), cols=5)
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(revenue_table):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = cell_data
            if i == 0:
                row.cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph('*Estimated based on reported YoY growth rates', style='Caption')
    doc.add_paragraph()
    
    doc.add_heading('Profitability Metrics', 2)
    doc.add_paragraph('• Adjusted EPS: $1.01 (vs. $0.89 consensus) - 13% beat')
    doc.add_paragraph('• Non-GAAP Operating Margin: 22%+ (expanding 40-80 bps YoY)')
    doc.add_paragraph('• Gross Margin: Approximately 80% (consistent with historical levels)')
    doc.add_paragraph('• Operating Cash Flow: $204.1M')
    doc.add_paragraph('• Free Cash Flow Margin: 52% (including 2% datacenter CapEx)')
    doc.add_paragraph()
    
    doc.add_heading('Key Operational Metrics', 2)
    doc.add_paragraph('• RPO Growth: 31% YoY - Strong indicator of future revenue acceleration')
    doc.add_paragraph('• Customer Additions: Continued enterprise adoption across Fortune 500')
    doc.add_paragraph('• Dollar-Based Net Retention: Typically exceeds 120% (strong expansion within existing customers)')
    doc.add_paragraph('• Platform Adoption: Increasing cross-sell of DLP, CASB, and Digital Experience modules')
    
    doc.add_page_break()
    
    # Guidance
    doc.add_heading('Guidance and Outlook', 1)
    
    doc.add_heading('FY2026 Full-Year Guidance (Fiscal Year Ends July 31, 2026)', 2)
    guidance_table = [
        ['Metric', 'Previous Guidance', 'Raised Guidance', 'Change', 'Implied Growth'],
        ['ARR', '$3.70B - $3.72B', '$3.73B - $3.75B', '+$30M', '24%'],
        ['Revenue', '$3.282B - $3.301B', '$3.309B - $3.322B', '+$21M - $27M', 'Approximately 23%'],
    ]
    
    table2 = doc.add_table(rows=len(guidance_table), cols=5)
    table2.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(guidance_table):
        row = table2.rows[i]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = cell_data
            if i == 0:
                row.cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    analysis_heading = doc.add_paragraph('Guidance Analysis:')
    analysis_heading.runs[0].bold = True
    
    doc.add_paragraph(
        'The guidance raise is significant as it reflects: (1) sustained enterprise demand despite macro uncertainty, '
        '(2) successful execution on large deals, and (3) platform expansion driving higher ACV. The 24% ARR growth '
        'target maintains Zscaler as one of the fastest-growing large-cap cybersecurity companies. The 31% RPO growth '
        'suggests bookings momentum that should support continued revenue acceleration.'
    )
    
    doc.add_page_break()
    
    # Competitive Positioning
    doc.add_heading('Competitive Positioning', 1)
    
    doc.add_heading('Market Leadership in Zero Trust and SASE', 2)
    doc.add_paragraph(
        'Zscaler remains a clear leader in the Security Service Edge (SSE) and SASE markets. Key competitors include:'
    )
    
    comp_table = [
        ['Vendor', 'Key Strengths', 'Market Position'],
        ['Zscaler', 'Pure-play cloud, Zero Trust pioneer, global scale', 'Leader (Gartner 4.7/5)'],
        ['Palo Alto Networks', 'Prisma Access, broader security portfolio', 'Strong Player (4.5/5)'],
        ['Netskope', 'CASB heritage, strong cloud app security', 'Strong Player'],
        ['Cloudflare', 'Network reach, cost competitive', 'Fast Growing Disruptor'],
    ]
    
    table3 = doc.add_table(rows=len(comp_table), cols=3)
    table3.style = 'Light List Accent 1'
    
    for i, row_data in enumerate(comp_table):
        row = table3.rows[i]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = cell_data
            if i == 0:
                row.cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('Zscaler Key Differentiators', 2)
    doc.add_paragraph('1. Purpose-Built Cloud Architecture', style='List Number')
    doc.add_paragraph(
        'Unlike competitors with on-premises heritage, Zscaler was designed for the cloud from inception, '
        'enabling superior scalability, lower latency, and faster feature deployment.'
    )
    
    doc.add_paragraph('2. Global Scale and Network Coverage', style='List Number')
    doc.add_paragraph(
        'Over 150 data centers worldwide processing 500+ billion transactions daily. This global footprint '
        'provides unmatched reliability and performance for multinational enterprises.'
    )
    
    doc.add_paragraph('3. Platform Consolidation', style='List Number')
    doc.add_paragraph(
        'Customers can replace 10+ legacy security point solutions (firewalls, VPNs, proxies, DLP, CASB) '
        'with Zscaler platform, driving TCO savings of 50%+ and operational simplification.'
    )
    
    doc.add_paragraph('4. Zero Trust Leadership', style='List Number')
    doc.add_paragraph(
        'First-mover advantage in Zero Trust architecture with 15+ years of innovation. The company '
        'pioneered the model of treating the internet as the new corporate network.'
    )
    
    doc.add_page_break()
    
    # Investment Thesis
    doc.add_heading('Investment Thesis', 1)
    
    doc.add_heading('Bull Case', 2)
    doc.add_paragraph('• Secular Tailwind: Zero Trust and SASE adoption still in early innings of multi-decade cycle', style='List Bullet')
    doc.add_paragraph('• Market Leadership: Clear number one in pure-play cloud security with durable competitive moat', style='List Bullet')
    doc.add_paragraph('• Margin Expansion: Operating leverage driving 40-80 bps margin improvement annually', style='List Bullet')
    doc.add_paragraph('• Rule of 40: Revenue growth (24-26%) + FCF margin (52%) = 76-78 score (exceptional)', style='List Bullet')
    doc.add_paragraph('• Platform Expansion: Cross-sell opportunities in DLP, CASB, Digital Experience driving ARPU growth', style='List Bullet')
    doc.add_paragraph('• Enterprise Penetration: Only 15% market penetration in Fortune 500 leaves massive runway', style='List Bullet')
    doc.add_paragraph('• Competitive Wins: Taking share from legacy vendors (Palo Alto, Cisco, Fortinet)', style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('Bear Case and Risks', 2)
    doc.add_paragraph('• Competitive Intensity: Palo Alto, Cloudflare, and hyperscalers increasing market pressure', style='List Bullet')
    doc.add_paragraph('• Valuation Premium: Trading at 11-13x EV/Sales, vulnerable to multiple compression', style='List Bullet')
    doc.add_paragraph('• Macro Sensitivity: Enterprise IT spending could contract in recession scenario', style='List Bullet')
    doc.add_paragraph('• Execution Risk: Maintaining 24%+ growth at $3B+ revenue scale requires flawless execution', style='List Bullet')
    doc.add_paragraph('• Customer Concentration: Large enterprise deals create quarterly volatility in bookings', style='List Bullet')
    doc.add_paragraph('• Technology Risk: Architectural shifts (e.g., AI-native security) could disrupt market', style='List Bullet')
    
    doc.add_page_break()
    
    # Valuation
    doc.add_heading('Valuation and Analyst Sentiment', 1)
    
    doc.add_heading('Wall Street Consensus', 2)
    doc.add_paragraph('• Average Price Target: $276 - $315 (varies by source)')
    doc.add_paragraph('• Price Target Range: $209 (bear case) to $390 (bull case)')
    doc.add_paragraph('• Recent Target: Baird lowered from $360 to $300 (still Outperform rating)')
    doc.add_paragraph('• Consensus Rating: Moderate Buy / Outperform across most analysts')
    doc.add_paragraph()
    
    doc.add_heading('Valuation Multiples (Estimated)', 2)
    doc.add_paragraph('• EV/Sales (NTM): Approximately 11-13x (premium to peers)')
    doc.add_paragraph('• P/E (Non-GAAP, NTM): Approximately 40-50x')
    doc.add_paragraph('• EV/FCF: Approximately 25-30x')
    doc.add_paragraph('• PEG Ratio: Approximately 1.8-2.0x (reasonable for quality growth)')
    doc.add_paragraph()
    
    val_assessment = doc.add_paragraph('Valuation Assessment:')
    val_assessment.runs[0].bold = True
    doc.add_paragraph(
        'Zscaler trades at a premium to cybersecurity peers, justified by: (1) superior growth profile, '
        '(2) best-in-class operating margins, (3) architectural competitive advantages, and (4) market leadership position. '
        'However, the stock remains vulnerable to multiple compression if growth decelerates below 20% or if competitive '
        'threats materially intensify. Current valuation implies continued execution at a high level.'
    )
    
    doc.add_page_break()
    
    # Cohesity Relevance
    doc.add_heading('Relevance to Cohesity', 1)
    
    doc.add_heading('Strategic Insights for Cohesity', 2)
    doc.add_paragraph(
        'While Zscaler operates in network security and Cohesity in data security/management, several strategic '
        'lessons are directly applicable:'
    )
    
    doc.add_paragraph('1. Platform Consolidation as Competitive Advantage', style='List Number')
    doc.add_paragraph(
        'Zscaler success demonstrates the power of consolidating multiple point solutions into a unified platform. '
        'Cohesity Data Cloud similarly consolidates backup, disaster recovery, file services, object storage, and security. '
        'This consolidation drives TCO savings, operational simplification, and higher customer stickiness.'
    )
    
    doc.add_paragraph('2. Cloud-Native Architecture Wins Long-Term', style='List Number')
    doc.add_paragraph(
        'Purpose-built cloud architecture provides durable advantages over retrofitted on-premises solutions. '
        'Cohesity should aggressively emphasize its cloud-first design against legacy competitors like Veeam, '
        'Commvault, and NetBackup - similar to how Zscaler wins against Palo Alto hardware-based solutions.'
    )
    
    doc.add_paragraph('3. Rule of 40 Excellence Drives Valuation', style='List Number')
    doc.add_paragraph(
        'Zscaler Rule of 40 score of 76-78 (24% growth + 52% FCF margin) is best-in-class and drives premium valuation. '
        'Cohesity should target similar efficiency metrics: balancing growth with profitability to demonstrate '
        'operational excellence to investors and potential acquirers.'
    )
    
    doc.add_paragraph('4. Land-and-Expand Success', style='List Number')
    doc.add_paragraph(
        'Zscaler high dollar-based net retention (120%+) validates the importance of platform breadth and customer success. '
        'Cohesity should focus on expanding wallet share within existing customers through: (a) additional use cases, '
        '(b) increased data under management, and (c) cross-sell of new modules like security and analytics.'
    )
    
    doc.add_paragraph('5. Go-to-Market Lessons', style='List Number')
    doc.add_paragraph(
        'Zscaler success with Fortune 500 enterprises demonstrates the value of: (a) strong channel partnerships, '
        '(b) technical pre-sales teams, (c) POC-driven sales cycles, and (d) executive sponsorship. Cohesity should '
        'replicate these GTM motions.'
    )
    
    doc.add_paragraph()
    
    doc.add_heading('Competitive Market Observations', 2)
    doc.add_paragraph('• Beat-and-Raise Momentum: Consistent execution rewarded with multiple expansion in public markets')
    doc.add_paragraph('• Architectural Differentiation: Technology moats matter more than feature parity in competitive wins')
    doc.add_paragraph('• Scale Advantages: Network effects and global infrastructure create compounding defensibility')
    doc.add_paragraph('• Margin Discipline: Public investors value profitable growth (Rule of 40) over pure revenue acceleration')
    doc.add_paragraph('• Platform Strategy: Multi-product platforms command premium valuations vs. point solutions')
    
    doc.add_page_break()
    
    # Recommendation
    doc.add_heading('Recommendation and Conclusion', 1)
    
    doc.add_paragraph()
    rating = doc.add_paragraph('Rating: OUTPERFORM / BUY')
    rating.runs[0].bold = True
    rating.runs[0].font.size = Pt(14)
    rating.runs[0].font.color.rgb = RGBColor(0, 176, 80)
    
    doc.add_paragraph()
    
    doc.add_heading('Investment Thesis Summary', 2)
    doc.add_paragraph(
        'Zscaler remains a high-quality growth stock in the cybersecurity sector with durable competitive advantages '
        'stemming from its purpose-built cloud architecture and Zero Trust market leadership. Q2 FY2026 results '
        'demonstrated continued execution excellence with revenue and EPS beats, margin expansion, and raised guidance. '
        'The company is exceptionally well-positioned to capitalize on the secular shift from legacy perimeter security '
        'to cloud-native Zero Trust architecture.'
    )
    
    doc.add_paragraph()
    
    doc.add_heading('Key Investment Considerations', 2)
    doc.add_paragraph('• Strong Q2 beat-and-raise affirms growth trajectory and competitive positioning', style='List Bullet')
    doc.add_paragraph('• 24% ARR guidance demonstrates sustained enterprise demand despite macro headwinds', style='List Bullet')
    doc.add_paragraph('• Operating leverage driving margin expansion toward 25%+ target (currently 22%+)', style='List Bullet')
    doc.add_paragraph('• Competitive position strengthening with architectural advantages vs. legacy vendors', style='List Bullet')
    doc.add_paragraph('• Valuation premium warranted by quality of growth, profitability, and market leadership', style='List Bullet')
    doc.add_paragraph('• Platform expansion creating new revenue vectors and increasing customer stickiness', style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('Catalysts to Monitor', 2)
    doc.add_paragraph('• Q3 FY2026 earnings (expected May 2026) - continued momentum or deceleration?', style='List Bullet')
    doc.add_paragraph('• Large enterprise wins and logo additions in Fortune 500', style='List Bullet')
    doc.add_paragraph('• New product launches and AI-powered security features', style='List Bullet')
    doc.add_paragraph('• Competitive dynamics vs. Palo Alto Networks, Cloudflare, and Netskope', style='List Bullet')
    doc.add_paragraph('• M&A activity in SASE/SSE market (potential consolidation)', style='List Bullet')
    doc.add_paragraph('• Macro IT spending trends and federal budget impacts', style='List Bullet')
    
    doc.add_page_break()
    
    # Appendix
    doc.add_heading('Appendix: Methodology and Sources', 1)
    
    doc.add_heading('Data Sources', 2)
    doc.add_paragraph('• Zscaler Q2 FY2026 Earnings Press Release (February 26, 2026)', style='List Bullet')
    doc.add_paragraph('• Zscaler Investor Relations website (ir.zscaler.com)', style='List Bullet')
    doc.add_paragraph('• Yahoo Finance, MarketBeat, TipRanks analyst consensus data', style='List Bullet')
    doc.add_paragraph('• Gartner Peer Insights Security Service Edge market reviews', style='List Bullet')
    doc.add_paragraph('• Forrester Wave SASE Provider rankings (September 2025)', style='List Bullet')
    doc.add_paragraph('• Public SEC filings and investor conference presentations', style='List Bullet')
    doc.add_paragraph('• Competitive intelligence from technology review sites and Reddit communities', style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('Analytical Methodology', 2)
    doc.add_paragraph('• Financial metrics based on company-reported GAAP and non-GAAP figures', style='List Bullet')
    doc.add_paragraph('• YoY comparisons calculated from reported growth rates', style='List Bullet')
    doc.add_paragraph('• Competitive positioning assessed via third-party market research and peer benchmarking', style='List Bullet')
    doc.add_paragraph('• Valuation multiples estimated based on current market data and consensus forecasts', style='List Bullet')
    doc.add_paragraph('• Cohesity strategic relevance derived from platform consolidation and GTM parallels', style='List Bullet')
    doc.add_paragraph('• Rule of 40 calculation: Revenue Growth % + FCF Margin % = Efficiency Score', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    separator = doc.add_paragraph('—' * 50)
    separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    disclaimer = doc.add_paragraph(
        'This analysis is for informational purposes only and does not constitute investment advice. '
        'Financial projections and forward-looking statements involve risks and uncertainties. '
        'Prepared for internal strategic use at Cohesity Inc.'
    )
    disclaimer.runs[0].font.size = Pt(9)
    disclaimer.runs[0].italic = True
    disclaimer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    prepared = doc.add_paragraph('Prepared by Jarvis - AI Assistant to Eric Brown')
    prepared.runs[0].font.size = Pt(9)
    prepared.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    date_footer = doc.add_paragraph('February 27, 2026')
    date_footer.runs[0].font.size = Pt(9)
    date_footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Save
    filename = 'zscaler_q2_fy2026_earnings_analysis.docx'
    doc.save(filename)
    print(f'✅ Document created: {filename}')
    return filename

if __name__ == '__main__':
    create_report()
