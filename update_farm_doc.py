#!/usr/bin/env python3
"""
Update Farm Trip documentation with farm address
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# File path
doc_path = "/Users/ericbrown/.openclaw/workspace/Farm_Trip_Sept_Oct_2025_Tax_Documentation.docx"

# Open the document
doc = Document(doc_path)

# Farm address to add
farm_address = """3553 Old Mission Road
Traverse City, MI 49686"""

# Find the position to insert (after the first paragraph which should be the title)
# We'll insert after the first paragraph
if len(doc.paragraphs) > 0:
    # Get the first paragraph (title)
    title_para = doc.paragraphs[0]
    
    # Insert new paragraphs after the title
    # We need to work with the underlying XML to insert at a specific position
    # Instead, we'll add to the beginning and then reorder
    
    # Add a blank line after title
    p_blank = doc.paragraphs[0]._element
    p_blank.addnext(doc.add_paragraph()._element)
    
    # Add the "Property Visited" heading
    p_heading = doc.paragraphs[0]._element
    heading_para = doc.add_paragraph()
    heading_para.text = "Property Visited:"
    heading_para.runs[0].bold = True
    heading_para.runs[0].font.size = Pt(12)
    p_heading.addnext(heading_para._element)
    
    # Add the farm address
    address_para = doc.add_paragraph()
    address_para.text = farm_address
    address_para.runs[0].font.size = Pt(11)
    heading_para._element.addnext(address_para._element)
    
    # Add another blank line
    blank_para2 = doc.add_paragraph()
    address_para._element.addnext(blank_para2._element)

# Save the document
doc.save(doc_path)

print("✓ Document updated successfully")
print(f"✓ Farm address added to: {doc_path}")
