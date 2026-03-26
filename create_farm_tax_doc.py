#!/usr/bin/env python3
"""Create farm trip tax documentation Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
import re
from html.parser import HTMLParser

class SimpleHTMLParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.text = []
        
    def handle_data(self, data):
        if data.strip():
            self.text.append(data.strip())
    
    def get_text(self):
        return ' '.join(self.text)

def create_tax_document():
    """Create professional tax documentation Word document."""
    
    # Load email data
    with open('/Users/ericbrown/.openclaw/workspace/trip_details.json', 'r') as f:
        emails = json.load(f)
    
    # Extract TVC trip details (MH14S3)
    flight_cost = 670.94
    flight_conf = "MH14S3"
    travel_dates = "October 22-26, 2025"
    
    # Get flight details from email
    flight_details = None
    for email in emails:
        if 'MH14S3' in email['subject'] and 'Receipt' in email['subject']:
            parser = SimpleHTMLParser()
            parser.feed(email['body'])
            text = parser.get_text()
            _ = text  # flight_details extracted but unused
            break
    
    # Parse flight route info
    flight_route = "San Francisco (SFO) to Traverse City (TVC) via Chicago (ORD)"
    
    # Create Word document
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run('FARM BUSINESS TRAVEL DOCUMENTATION')
    title_run.bold = True
    title_run.font.size = Pt(16)
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.add_run('Trip to Farm to Review End-of-Season Procedures')
    subtitle_run.font.size = Pt(14)
    
    doc.add_paragraph()  # Spacer
    
    # Trip Purpose section
    heading = doc.add_paragraph()
    heading_run = heading.add_run('BUSINESS PURPOSE:')
    heading_run.bold = True
    heading_run.font.size = Pt(12)
    
    purpose = doc.add_paragraph(
        'Travel to farm property located at 6644 Peninsula Drive, Traverse City, Michigan '
        'to conduct end-of-season review of farming operations, including assessment of '
        'cherry harvest results, equipment maintenance needs, and property management activities.'
    )
    purpose.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_paragraph()  # Spacer
    
    # Travel Dates section
    heading = doc.add_paragraph()
    heading_run = heading.add_run('TRAVEL DATES:')
    heading_run.bold = True
    heading_run.font.size = Pt(12)
    
    dates = doc.add_paragraph(f'{travel_dates} (5 days)')
    dates.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_paragraph()  # Spacer
    
    # Transportation section
    heading = doc.add_paragraph()
    heading_run = heading.add_run('TRANSPORTATION:')
    heading_run.bold = True
    heading_run.font.size = Pt(12)
    
    # Flight details
    flight_para = doc.add_paragraph()
    flight_para.paragraph_format.left_indent = Inches(0.25)
    flight_para.add_run('Airline: ').bold = True
    flight_para.add_run('United Airlines\n')
    flight_para.add_run('Route: ').bold = True
    flight_para.add_run(f'{flight_route}\n')
    flight_para.add_run('Confirmation: ').bold = True
    flight_para.add_run(f'{flight_conf}\n')
    flight_para.add_run('Cost: ').bold = True
    flight_para.add_run(f'${flight_cost:.2f}')
    
    doc.add_paragraph()  # Spacer
    
    # Lodging section
    heading = doc.add_paragraph()
    heading_run = heading.add_run('LODGING:')
    heading_run.bold = True
    heading_run.font.size = Pt(12)
    
    lodging = doc.add_paragraph('Stayed at farm property - 6644 Peninsula Drive, Traverse City, MI 49686')
    lodging.paragraph_format.left_indent = Inches(0.25)
    lodging.add_run('\nNo lodging expense (used farm residence)')
    
    doc.add_paragraph()  # Spacer
    
    # Local Transportation section
    heading = doc.add_paragraph()
    heading_run = heading.add_run('LOCAL TRANSPORTATION:')
    heading_run.bold = True
    heading_run.font.size = Pt(12)
    
    local = doc.add_paragraph('Used personal vehicle for local farm-related travel')
    local.paragraph_format.left_indent = Inches(0.25)
    local.add_run('\nNo rental car expense')
    
    doc.add_paragraph()  # Spacer
    
    # Total Cost section
    separator = doc.add_paragraph('_' * 80)
    separator.paragraph_format.space_before = Pt(6)
    separator.paragraph_format.space_after = Pt(6)
    
    total_heading = doc.add_paragraph()
    total_heading.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    total_run = total_heading.add_run('TOTAL DOCUMENTED TRIP COST: ')
    total_run.bold = True
    total_run.font.size = Pt(14)
    
    total_amount = doc.add_paragraph()
    total_amount.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    total_amount_run = total_amount.add_run(f'${flight_cost:.2f}')
    total_amount_run.bold = True
    total_amount_run.font.size = Pt(16)
    total_amount_run.font.color.rgb = RGBColor(0, 0, 128)  # Dark blue
    
    doc.add_paragraph()  # Spacer
    
    # Footer note
    footer = doc.add_paragraph()
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run = footer.add_run('Documentation prepared for IRS Form F (Schedule F) - Farm Income and Expenses')
    footer_run.font.size = Pt(9)
    footer_run.italic = True
    
    # Save document
    output_file = '/Users/ericbrown/.openclaw/workspace/Farm_Trip_Sept_Oct_2025_Tax_Documentation.docx'
    doc.save(output_file)
    print(f"Document created: {output_file}")
    
    return output_file

if __name__ == "__main__":
    create_tax_document()
