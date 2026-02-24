#!/usr/bin/env python3
"""
Create German translations of Cohesity articles with professional styling
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# German translations with professional technical terminology
articles = [
    {
        "filename": "Cohesity_Ransomware_Detection_ML_DE.docx",
        "title_de": "Erkennung von Ransomware-Angriffen mit Cohesity | Modelle für maschinelles Lernen",
        "title_en": "Cohesity Ransomware Attack Detection | Machine-Learning Models",
        "url": "https://www.cohesity.com/blogs/cohesity-ransomware-detection-machine-learning-models/",
        "content": [
            ("heading1", "Erkennung von Ransomware-Angriffen mit Cohesity | Modelle für maschinelles Lernen"),
            ("body", "Warum nutzen Kunden die Cohesity Data Cloud? Sie ist die letzte Verteidigungslinie zum Schutz, zur Erkennung und zur schnellen und sicheren Wiederherstellung wertvoller Daten nach Ransomware-Angriffen. Unsere Technologie minimiert die Angriffsfläche und macht uns besser darin, strengere Backup-SLAs zu erfüllen und alle Arten von Wiederherstellungsvorgängen dramatisch zu beschleunigen. Wir haben auf dieser Grundlage aufgebaut, um Kunden dabei zu helfen, einen Angriff früher zu erkennen und schneller zu reagieren. Wir bieten erweiterte Bedrohungserkennungssignale mit weniger Fehlalarmen sowie schnelle und effiziente Möglichkeiten, Warnmeldungen zu empfangen, zu verwalten und effizient darauf zu reagieren – sowohl um den Angriff zu stoppen als auch um die Bereinigung zu beschleunigen."),
            
            ("heading2", "Die Bösen in Schach halten"),
            ("body", "Ein besonderer Vorteil unseres Backup-Prozesses ist die Sammlung von Informationen über die Daten und deren Veränderungen. Anomale Änderungen in Backup-Daten können präzise Frühwarnungen liefern, ohne viele Fehlalarme zu generieren. Wenn beispielsweise neue Daten plötzlich weniger komprimierbar sind als frühere Arbeiten, ist dies ein potenzielles Kennzeichen dafür, dass sie verschlüsselt wurden. Dies wird manchmal als Entropie-Erkennung bezeichnet, da verschlüsselte Daten sehr zufällig aussehen, weshalb sie nicht komprimiert werden können."),
            
            ("body", "Um laufende Angriffe zu erkennen und Fehlalarme zu vermeiden, speist Cohesity mehrere Metriken in Algorithmen für maschinelles Lernen ein, die in unserer Helios-Steuerungsebene ausgeführt werden, einschließlich, aber nicht beschränkt auf:"),
            
            ("bullet", "Inhaltsinformationen pro Backup: Größe der geschriebenen Daten, Größe der gelesenen Daten, logische Größe"),
            ("bullet", "Entropie-/Komprimierungsverhältnis pro Backup"),
            ("bullet", "Änderungsverfolgungsinformationen pro Backup: Anzahl der hinzugefügten, gelöschten, aktualisierten und unveränderten Dateien"),
            ("bullet", "Aggregierte Statistiken über mehrere Backups: Maximale geschriebene Datenbytes, maximale logische Quellgröße in Bytes, Anzahl erfolgreicher Durchläufe usw."),
            ("bullet", "Trainingssätze, die Änderungsmuster darstellen, die von häufig verwendeter Malware generiert werden"),
            
            ("body", "Die Machine-Learning-Modelle sind kein AI-Washing von Dingen, die genauso gut traditionell durchgeführt werden könnten; es handelt sich um multivariate Modelle. Maschinelles Lernen erfordert eine Baseline, die normales Verhalten definiert, und Cohesity benötigt mindestens 15 gültige historische Datensätze, um eine Erkennung auszulösen. Sobald eine Baseline von Metriken etabliert ist, beginnen unsere Modelle mit einer aggressiven Lernphase und gehen zu einem stationären Zustand über, der kontinuierlich für Präzision und Recall optimiert wird."),
            
            ("body", "Wir haben auch bestimmte heuristische regelbasierte Modelle, die von Menschen geschrieben wurden, um Fehlalarme zu entfernen. Wenn beispielsweise die Daten komprimierbar sind und die Ersatzdateien gleichermaßen komprimierbar sind, ist es unwahrscheinlich, dass ein Krypto-Angriff stattgefunden hat. Durch die Kombination mehrerer Modelle werden sowohl menschliche Intelligenz als auch reines maschinelles Lernen für optimale Ergebnisse kombiniert."),
            
            ("body", "Im Sinne der kontinuierlichen Verbesserung testen wir auch zusätzliche Machine-Learning-Modelle und vergleichen die Ergebnisse. Wenn sich Verbesserungen aus den neuesten Schadensmustern aus dem Feld ergeben, können wir dann die Erkennung aktualisieren, ohne Lücken in der Erkennungsabdeckung zu haben."),
            
            ("heading2", "Erkennung von Ransomware-Angriffen mit Cohesity"),
            ("body", "Die schädlichsten Malware-Angriffe sind Hacker-getrieben und nicht spontane Infektionen durch Malware in freier Wildbahn, die Ihren Laptop oder Server infizieren könnten. Aber Virusschäden sind immer noch möglich. Solche generische Malware versucht nicht, langsam und heimlich zu sein, daher sind ihre Auswirkungen offensichtlich. Wie Sie erwarten würden, haben wir eine 100%ige Erkennung bei Proben der Schadensmuster erreicht, die von Cerber, Cryptxxx, Cryptolocker, Locky und Wannacry erstellt wurden, die in einem Cloud-Konto getrennt von jeder anderen Cohesity-Umgebung ausgeführt wurden."),
            
            ("body", "Da unsere Ransomware-Angriffserkennungsfunktionen ohne zusätzliche Kosten für unsere Kunden in die Plattform integriert und SaaS-basiert sind, haben wir umfangreiche Daten gesammelt, um unsere Modelle zu trainieren. Dies hilft uns, uns an die Auswirkungen realer Angriffe anzupassen, die von Hackern entfesselt wurden, die in Kundensysteme eingedrungen sind und Cohesity-Backups zur Wiederherstellung verwendet haben."),
            
            ("body", "Da unsere Ransomware-Angriffserkennung in unsere Backups integriert ist, können wir eine höhere Effizienz, verbesserte Erkennung, schnellere Reaktionszeiten und ein geringeres Risiko für geschützte Daten bieten. Der zur Erpressung der Daten verwendete Verschlüsselungsprozess hat den Nebeneffekt, dass die Daten im Wesentlichen zufällig aussehen oder, um einen ausgefallenen Begriff zu verwenden, eine hohe Entropie haben. Diese „Entropie" führt dazu, dass die Datenreduzierung ineffektiv wird. Backup-Produkte führen Datenreduzierung durch, sodass sie die erhöhte Entropie bemerken können, ohne erhebliche zusätzliche Arbeit zu leisten. Ein eigenständiges Produkt zu haben, das alle Daten separat liest, um die Entropie zu überprüfen, ist verschwenderisch. Aus diesem Grund prüfen diese Produkte entweder nicht die Entropie, was die Erkennung beeinträchtigt, oder erfordern kostspielige zusätzliche Ressourcen, um die Entropie zu überprüfen, wenn andere Backup-Anbieter separate Produkte zur Angriffsüberwachung anbieten oder vorschlagen."),
            
            ("heading2", "Die Notwendigkeit einer schnellen, sicheren und koordinierten Reaktion"),
            ("body", "Während wir von dem Wert überzeugt sind, den wir durch eine datenzentrische Sicht auf laufende Angriffe bieten, wissen wir, dass in der realen Welt die Korrelation mit anderen Signalen die Agilität der Angriffsreaktion weiter verbessern kann. Um eine schnelle Reaktion auf Vorfälle zu ermöglichen, müssen SecOps-Mitarbeiter die Warnungen erhalten und Vertrauen in sie haben, um sie weiter zu untersuchen. Cohesity kann die Warnungen direkt an Security Orchestration and Automation Response (SOAR)-Plattformen senden. Ein kleines, aber nicht null Level an Fehlalarmen ist einem extremen Vorsichtsmaß vorzuziehen, das manchmal zugunsten der Nichterkennung irren würde. Um eine gute Hygiene bei der Bereinigung vergangener Warnungen aufrechtzuerhalten, bieten wir eine geschlossene Integration mit Palo Alto XSOAR und Cisco SecureX, die eine vollständige Disposition innerhalb der SOAR-Plattform ermöglicht. Durch die Nutzung automatisierter Playbooks müssen sich Benutzer nicht erneut in die Backup-Benutzeroberfläche einloggen, um die Reaktion abzuschließen, wodurch die mittlere Zeit bis zur Erkennung (MTTD) und die mittlere Zeit bis zur Reaktion (MTTR) reduziert werden."),
            
            ("heading2", "Verwaltung einer schnellen Wiederherstellung"),
            ("body", "Kann die Erkennungsausgabe unseren Kunden helfen, schneller wiederherzustellen, wenn der Angriff gestoppt wurde? Absolut! Cohesity generiert eine Liste der Dateien, die angegriffen zu sein scheinen. Und im Gegensatz zu anderen Anbietern, die nur die angegriffenen Dateien zurück auf Server wiederherstellen, die kompromittiert wurden, erzählen Cohesity-Kunden, denen wir bei der Wiederherstellung nach tatsächlichen Angriffen geholfen haben, einstimmig eine andere Geschichte: Ihre Sicherheitsteams und ihre Cyberversicherungsgesellschaften sagen, dass eine Wiederherstellung zurück in die kompromittierten Umgebungen inakzeptabel ist."),
            
            ("body", "Der praxiserprobte Best-Practice-Pfad besteht darin, alle Daten, die die Quarantäne durchlaufen haben, nicht auf ursprünglichen Servern oder VMs wiederherzustellen, sondern stattdessen auf sauber aufgebauten Servern oder VMs, vorzugsweise solchen, die auf die neuesten bekannten Schwachstellen überprüft und als frei von hochriskanten Problemen befunden wurden."),
            
            ("body", "Wenn ein tatsächlicher Angriff stattgefunden hat, haben alle Kunden, denen wir geholfen haben, eine Quarantäneverarbeitung in einer Sandbox für den gesamten Datensatz benötigt, nicht nur für die Dateien, die überschrieben wurden. Für die Untersuchung des gesamten Datensatzes verfügt Cohesity über eine einzigartig skalierbare schnelle Massen-Wiederherstellung und einen einzigartigen sofortigen Network Attached Storage (NAS)-Zugriff auf NAS-Backups. Diese werden durch die einzigartigen Snapshot-Metadaten und die robuste Dateisystemtechnologie von Cohesity ermöglicht, die im Hintergrund bleiben."),
            
            ("heading2", "Wie Cohesity hilft"),
            ("body", "Unsere Grundlage umfasst Unveränderlichkeit, die Ihnen hilft, Ihre Backup-Daten zu schützen. Sie bietet auch eine einzigartige Technologie, die es Organisationen ermöglicht, Bereinigung und Wiederherstellung/Restaurierung zurück zum normalen Betrieb in großem Maßstab zu beschleunigen. Darüber hinaus adressiert das Ransomware-Erkennungsmodell für maschinelles Lernen von Cohesity eine Schlüsselanforderung in der heutigen Cyberwelt und ermöglicht es Kunden, schneller zu reagieren und ihre vollständige Wiederherstellung nach einem Angriff zu beschleunigen – insbesondere wenn es mit führenden Cybersicherheitslösungen von Drittanbietern integriert ist. Zusammen dienen unsere Fähigkeit, Ihre Backup-Daten zu schützen, Bedrohungen zu erkennen und schnell im großen Maßstab wiederherzustellen, dazu, Ihre Fähigkeit zu stärken, die Auswirkungen von Ransomware und anderen Cyber-Bedrohungen zu minimieren."),
        ]
    },
    {
        "filename": "Cohesity_AI_ML_Ransomware_Protection_DE.docx",
        "title_de": "KI/ML-gesteuerte Ransomware-Schutz und -Wiederherstellung",
        "title_en": "AI/ML-Driven Ransomware Protection and Recovery",
        "url": "https://www.cohesity.com/blogs/ai-ml-driven-ransomware-protection-and-recovery/",
        "content": [
            ("heading1", "KI/ML-gesteuerte Ransomware-Schutz und -Wiederherstellung"),
            ("body", "Können wir standhalten und uns erholen? Können wir neuer und fortgeschrittener Ransomware-Malware standhalten, die sich kontinuierlich verändert? Und können wir uns erholen, bevor unser Unternehmen finanzielle Verluste, Kundenabwanderung und Markenschäden erleidet?"),
            
            ("body", "CISOs müssen diese schwierigen Fragen gegenüber Vorständen und Managementteams beantworten, die versichert werden wollen, dass ihre Organisation für die schlimmsten Szenarien vorbereitet und widerstandsfähig ist."),
            
            ("body", "Um Ransomware zu bekämpfen, was dem Überleben einer Organisation gleichkommen kann, benötigen CISOs jeden möglichen Vorteil – die besten Verteidigungen, die beste Wiederherstellung und die besten Prozesse für Incident Response und für Cyber Recovery."),
            
            ("body", "Künstliche Intelligenz (KI) hat sich als Kraftmultiplikator herausgestellt, der CISOs erheblich dabei hilft, Angriffe zu überstehen und sich davon zu erholen. KI kann Organisationen dabei helfen, Angriffen standzuhalten, indem sie große Datensätze nutzt, um Intelligenz abzuleiten und Risikomanagement sowie Sicherheitskontrollen, -praktiken und -prozesse zu automatisieren. Ebenso kann KI Organisationen dabei helfen, sich mit Zuversicht zu erholen, indem sie die Wiederherstellung automatisiert, kritische Daten identifiziert und sicherstellt, dass Daten keine Risiken und Schwachstellen wieder einführen."),
            
            ("heading2", "Wie KI helfen kann, Ransomware-Angriffen standzuhalten"),
            ("body", "KI wird heute verwendet, um hocheffektive Ransomware-Abwehr zu bieten. Dies sind zwar nicht alle, aber Beispiele für Technologien, die KI verwenden, um Organisationen bei der Verteidigung gegen die wachsende Komplexität und Häufigkeit von Ransomware-Angriffen zu helfen:"),
            
            ("bullet", "KI-fähige Multi-Faktor-Authentifizierung (MFA): MFA ist eine äußerst wichtige Kontrolle zur Bekämpfung von Ransomware. Durch die Verwendung von MFA können Ransomware-Banden nicht einfach Passwörter erraten oder Brute-Force-Methoden zum Knacken von Passwörtern verwenden. Und MFA kann mit KI weiter verstärkt werden, um ihren Schutz durch Verhalten (wie Tippgeschwindigkeit), adaptive (erfordert mehrere Authentifizierungen basierend auf Datenrisiko) oder Betrugserkennung (automatisches Blockieren eines Benutzers, wenn sein Zugriff über normale Grenzen hinausgeht) zu stärken."),
            
            ("bullet", "KI-fähige Ransomware-Erkennung: KI kann Netzwerkverkehr oder Dateizugriff analysieren, um Aktivitäten zu identifizieren, die darauf hindeuten könnten, dass ein Ransomware-Angriff unmittelbar bevorsteht oder bereits im Gange ist. Threat-Intelligence-Organisationen identifizieren und detonieren kontinuierlich Malware, um umfassende Indikatoren für Kompromittierungen zu dokumentieren, die eine Frühwarnung vor Ransomware-Aktivitäten bieten."),
            
            ("bullet", "KI-fähige Aktivitäts- und Verhaltensüberwachung: KI kann Zugriffe und Benutzerverhalten betrachten und feststellen, ob die Aktivität verdächtig ist und ein Ransomware-Angriff signalisieren könnte: fehlgeschlagene Anmeldeversuche, übermäßiger Dateizugriff oder andere Aktivitäten, die außerhalb der etablierten Normen durch Anzeichen von Ransomware-Aktivitäten liegen. Die Aktivitätsüberwachung kann Normen sowohl für Benutzer- als auch für Anwendungsverhalten basierend auf der kontinuierlichen Analyse von Aktivitätsprotokollen mit KI festlegen."),
            
            ("heading2", "Wie KI helfen kann, sich von Ransomware-Angriffen zu erholen"),
            ("body", "Ungeachtet umfangreicher Verteidigungen müssen sich Organisationen mit der Realität auseinandersetzen, dass sie einen Ransomware-Angriff erleiden. Angesichts der Komplexität von Unternehmens-Informationssystemen, wachsender Angriffssophistikation und einfacher menschlicher Fehler ist eine Ransomware-Infektion unvermeidlich. Der Schlüssel liegt darin, die Auswirkungen des Angriffs zu begrenzen und alle betroffenen Daten und Anwendungen so schnell wie möglich wiederherzustellen."),
            
            ("body", "In Zukunft kann KI die Sicherheit der Wiederherstellungsplattform verbessern, operative Intelligenz und Automatisierung bereitstellen, um eine schnelle und zuversichtliche Wiederherstellung zu ermöglichen – während sie Entscheidungsunterstützung bietet, um den Wiederherstellungsprozess zu optimieren. Eine schnelle Wiederherstellung ist, wenn Daten und Prozesse in Stunden statt in Tagen wiederhergestellt werden können. Und eine zuversichtliche Wiederherstellung bedeutet, dass die Wiederherstellungsdaten keine Schwachstellen und Bedrohungen wieder einführen, die zu einer Neuinfektion führen könnten. Darüber hinaus wird KI die Verwaltung und das Management der Plattform für optimale Effizienz und Sicherheit rationalisieren und erleichtern."),
            
            ("body", "Die folgenden KI-Fähigkeiten können Organisationen dabei helfen, eine zuversichtliche und schnelle Wiederherstellung zu erreichen und Effizienz und Sicherheit zu verbessern:"),
            
            ("bullet", "KI-Systemverhaltensverfolgung: Nahezu-Echtzeit-Überwachung privilegierter und administrativer Benutzer auf Indikatoren für anomale Aktivitäten."),
            
            ("bullet", "KI-gesteuerte Heilung: Verwendung von KI zur Überwachung der Plattform und Antizipation von Problemen sowie Vorschlag von Abhilfen."),
            
            ("bullet", "KI-fähige optimierte Planung: Basierend auf dem kritischen Bedarf und der Nutzung von Daten, Saisonalität und anderen Variablen kann KI Backup-Pläne anpassen und optimieren, um sicherzustellen, dass RPOs immer erfüllt werden."),
            
            ("bullet", "KI-Stilllegung inaktiver Daten: Als Teil des Backup-Prozesses kann KI Organisationen dabei helfen, festzustellen, welche Daten für die Archivierung inaktiv geworden sind. Dies hilft, die Wiederherstellungszeit zu reduzieren, indem die unnötige Wiederherstellung ungenutzter Daten eliminiert wird, und schafft gleichzeitig Effizienz und Kostenreduzierung bei der Speicherung."),
            
            ("body", "Aber langfristig können Ihre Backup-Daten, die den Großteil Ihrer kritischen Unternehmensdaten darstellen, als sichere Plattform und zeitversetzte Datenquelle für Retrieval Augmented Generation (RAG) dienen, um KI-gesteuerte Suche und Entdeckung von Betriebs-, Transaktions- und anderen Unternehmensdaten zu unterstützen."),
            
            ("body", "Dies ist eine einzigartige Fähigkeit, da kein anderes Unternehmens-Repository die Hyper-Scale-Grundlage, den gesicherten Zugriff, die Analytik und eine zeitversetzte Ansicht von Daten bietet, die von einer modernen Datensicherheits- und Management-Plattform bereitgestellt wird."),
            
            ("heading2", "Ransomware wird weiterhin an Häufigkeit und Raffinesse zunehmen"),
            ("body", "Ransomware wird in naher Zukunft nicht nachlassen. Tatsächlich ist es wahrscheinlich, dass Ransomware-Banden irgendwann in der Zukunft KI verwenden werden, um die Effektivität ihrer Angriffe zu erhöhen. Ab heute gibt es keinen glaubwürdigen Hinweis auf die Verwendung von KI durch diese Bedrohungsakteure, aber sie könnten die folgenden Taktiken mit KI integrieren:"),
            
            ("bullet", "KI zur Automatisierung der Angriffszielerfassung verwenden: Bedrohungsakteure könnten KI verwenden, um die Identifizierung anfälliger Organisationen zu automatisieren. ML könnte verwendet werden, um die Netzwerke und Endpunkte einer Organisation und möglicherweise Mitarbeiter schnell zu durchsuchen, um Schwachstellen und Ziele zu identifizieren."),
            
            ("bullet", "Angriffserstellung: Ransomware-Banden könnten KI verwenden, um große Datenmengen über Angriffe zu analysieren und neue Angriffe zu erstellen, die eine höhere Wahrscheinlichkeit haben, unentdeckt zu bleiben."),
            
            ("heading2", "Standhalten und Wiederherstellung mit einer KI/ML-Grundlage"),
            ("body", "Cybersicherheitslösungen, die KI/ML verwenden, sind in den letzten Jahren entstanden, um bei der Bekämpfung der wachsenden Raffinesse und Zerstörungskraft von Ransomware und anderen bösartigen Aktivitäten zu helfen. Diese Lösungen haben sich als effektiv erwiesen, um Organisationen dabei zu helfen, der Flut von Tausenden von Angriffen pro Jahr standzuhalten."),
            
            ("body", "Datensicherheit und -management, das kritische Unternehmens-Backups und Cyber-Recovery bietet, begann die KI/ML-Reise vor Jahren mit Anomalie-Erkennung, Planung und Optimierung. Da Ransomware-Angriffe jetzt die Fähigkeit von Organisationen bedrohen, ihre Backup-Daten für die Wiederherstellung zu nutzen, wird KI/ML eine immer größere Rolle spielen, um sicherzustellen, dass Organisationen sich mit Zuverlässigkeit und Zuversicht erholen können."),
            
            ("body", "Cohesity verwendet heute bestimmte KI-Einblicke, um Organisationen dabei zu helfen, sich schnell und zuversichtlich zu erholen. Mit ML-Backup werden Snapshots auf ungewöhnliche Änderungen analysiert. Baselines werden für eine Reihe von Variablen festgelegt, einschließlich der Zeitreihen der geschriebenen Daten, Entropie (Zufälligkeit der Daten), der Anzahl der Dateiänderungen und Dateierweiterungsänderungen. Wenn diese Variablen zu weit von den gelernten Baselines abdriften, ist dies ein Hinweis darauf, dass Ransomware oder andere bösartige Aktivitäten auftreten."),
            
            ("body", "Für den Bedrohungsschutz verwendet Cohesity KI/ML-gesteuerte Threat Intelligence, um Backup-Snapshots auf Indikatoren für Kompromittierungen (IOCs) zu analysieren. Diese IOCs bieten eine Frühwarnung, dass Änderungen an den Daten stattfinden, die mit Malware-Aktivitäten zusammenhängen. Und für die Datenklassifizierung verwendet Cohesity KI/ML, um schwer fassbare Muster in sensiblen Daten und/oder Datenelementen zu finden, die sensibel sind und über Dateien fragmentiert sind. Ohne dies können viele sensible Datenelemente übersehen werden."),
        ]
    },
    {
        "filename": "Cohesity_Security_First_Approach_DE.docx",
        "title_de": "Sicherheitsorientierter Ansatz zur Abwehr und schnellen Wiederherstellung nach Ransomware-Angriffen",
        "title_en": "Security-First Approach To Defend And Rapidly Recover From Ransomware Attacks",
        "url": "https://www.cohesity.com/blogs/defend-against-ransomware-attacks-a-security-first-approach/",
        "content": [
            ("heading1", "Sicherheitsorientierter Ansatz zur Abwehr und schnellen Wiederherstellung nach Ransomware-Angriffen"),
            ("body", "Um der sich entwickelnden Cyber-Bedrohungslandschaft zu begegnen, erhöhen Unternehmen weltweit ihre Investitionen in Datensicherheit. Die globalen Ausgaben für Cybersicherheit sind von 3,5 Milliarden US-Dollar im Jahr 2004 auf 124 Milliarden US-Dollar im Jahr 2019 gestiegen. Dieser 35-fache Anstieg wird voraussichtlich bis 2021 1 Billion US-Dollar übersteigen."),
            
            ("body", "Trotz erheblicher Investitionen in Datensicherheit erleben Organisationen aller Größen (von großen multinationalen Unternehmen bis hin zu staatlichen und städtischen Regierungen) einen raschen Anstieg der Häufigkeit und Intensität von Ransomware-Angriffen. Die Auswirkungen dieser Angriffe können lähmend sein und können meist auf eine Kombination aus ungelösten Software-Schwachstellen und internen menschlichen Handlungen/Fehlern sowie raffinierten Taktiken zurückgeführt werden, die zahlreiche Techniken beinhalten, um für einige Zeit unentdeckt zu bleiben und sich in einer Umgebung zu verbreiten, bevor sie sich manifestieren."),
            
            ("body", "Im Jahr 2019 kosteten Cyber-Verletzungen die Weltwirtschaft 2,1 Billionen US-Dollar, und 11,5 Milliarden US-Dollar davon stammten aus Ransomware-Angriffen. Strafverfolgungsbehörden, einschließlich Europol, wiesen darauf hin, dass Ransomware weltweit die größte Bedrohung bleibt. Laut Forrester Research bestätigten jedoch nur 21 Prozent der befragten Organisationen, dass sie Notfallpläne zur Wiederherstellung nach Ransomware-Angriffen haben, und nur 11 Prozent der Befragten sagten, dass sie zuversichtlich seien, ihre Daten innerhalb von drei Tagen nach einem Angriff wiederherzustellen."),
            
            ("heading2", "Warum können sich Organisationen nicht gegen Ransomware verteidigen?"),
            ("body", "Um eine Auszahlung sicherzustellen, greifen Cyberkriminelle nicht nur die Produktionsumgebung an, sondern zielen zunehmend auf Backup-Daten und -Infrastruktur ab – was die „Versicherungspolice" effektiv lahmlegt, auf die Organisationen angewiesen sind, wenn eine Katastrophe eintritt. Die Angreifer nutzen oft Schwachstellen aus, die mit Legacy-Backup-Lösungen verbunden sind, die vor dem Aufkommen der Ransomware-Industrie entwickelt wurden. Bevor die Produktionsumgebung verschlüsselt wird, ist bekannt, dass raffinierte Malware Schattenkopien und Wiederherstellungspunkt-Daten zerstört. Aufgrund ihrer zugrunde liegenden Architektur machen diese Malware die Legacy-Backup-Infrastruktur zu einer leichten Beute anstatt zu einer Verteidigung gegen Ransomware-Angriffe."),
            
            ("body", "Fortgesetzte Schulung von Mitarbeitern zur Cybersicherheit und Investitionen in Sicherheitstools sind wichtig. Organisationen müssen auch eine moderne, robuste Backup-Lösung implementieren, die dabei hilft, Backup-Daten vor Ransomware-Angriffen zu schützen und schnell wiederherzustellen, um Ausfallzeiten zu reduzieren."),
            
            ("body", "Die umfassende Anti-Ransomware-Lösung von Cohesity geht über die Erkennung hinaus. Nach einem typischen Angriffslebenszyklus bietet Cohesity eine End-to-End-Lösung, die Unternehmen dabei hilft:"),
            
            ("bullet", "Ihre Angriffsfläche zu reduzieren"),
            ("bullet", "Backup-Daten mit einzigartiger unveränderlicher Architektur und einfacher richtlinienbasierter Datenverwaltung zu schützen"),
            ("bullet", "Anomalien zu erkennen, die potenzielle Angriffe mit maschinellem Lernen signalisieren"),
            ("bullet", "Tiefe Sichtbarkeit zu gewährleisten, dass die Backups sauber sind und beim Wiederherstellen keine Schwachstellen wieder einführen"),
            ("bullet", "Und am wichtigsten: Schnelle Wiederherstellung zur Reduzierung von Ausfallzeiten"),
            
            ("heading2", "Reduzierung der Angriffsfläche"),
            ("body", "Cohesity-Kunden reduzieren ihren Daten-Fußabdruck, indem sie verschiedene Backup-Komponenten, Disaster Recovery, Dateidienste, Objektspeicher, Dev/Test und Analysen auf einer Web-Scale-Plattform konsolidieren. Kunden reduzieren ihren Daten-Fußabdruck und ihre Angriffsfläche mit der globalen Deduplizierung variabler Länge von Cohesity über Datenquellen und Komprimierung weiter. Dies hilft Unternehmen, ihre Exposition gegenüber Cyberkriminellen zu reduzieren."),
            
            ("heading2", "Verhindern, dass Backup zum Ransomware-Ziel wird"),
            ("body", "Eine moderne Backup-Lösung mit mehrschichtigem Verteidigungsansatz ist erforderlich, um sich gegen raffinierte Ransomware-Angriffe zu verteidigen, die Folgendes umfassen:"),
            
            ("bullet", "Unveränderliches Dateisystem: Im Kern hält das unveränderliche Dateisystem von Cohesity, SpanFS, die Backup-Jobs in zeitbasierten unveränderlichen Snapshots. Der ursprüngliche Backup-Job wird in einem unveränderlichen Zustand gehalten und ist nie zugänglich, was verhindert, dass er von einem externen System gemountet wird. Die einzige Möglichkeit, das Backup im Lese-/Schreibmodus zu mounten, besteht darin, dieses ursprüngliche Backup zu klonen, was automatisch vom System durchgeführt wird. Obwohl Ransomware möglicherweise in der Lage ist, Dateien im gemounteten (Lese-/Schreib-) Backup zu löschen, kann sie den unveränderlichen Snapshot nicht beeinflussen."),
            
            ("bullet", "DataLock: DataLock ist ein WORM für Backup-Snapshots, das eine weitere Schutzschicht gegen Ransomware-Angriffe bietet. Seit Cohesity Pegasus 6.1 verfügbar, ermöglicht diese Fähigkeit Sicherheitsbeauftragten, eine „DataLock"-Richtlinie für ausgewählte Jobs zu erstellen und anzuwenden und eine höhere Ordnung der Unveränderlichkeit für geschützte Daten zu erreichen – etwas, das Sicherheitsbeauftragte und Administratoren nicht ändern/löschen können. Diese Funktion ist in RBAC integriert und eliminiert die Notwendigkeit für Drittanbieter-Tools."),
            
            ("bullet", "Multi-Faktor-Authentifizierung: So sehr wir möchten, dass Passwörter garantierten Schutz bieten, werden Passwörter ständig kompromittiert. Cohesity bietet Multi-Faktor-Authentifizierung, die der beste Weg ist, um Phishing-Systeme und andere Passwort-Hacks zu mindern."),
            
            ("bullet", "Richtlinienbasierte Air Gap: Nichts ist 100% sicher (außer Steuern und Tod); daher fügt die Replikation Ihrer unternehmenskritischen Daten auf einen anderen unveränderlichen Cohesity-Cluster/Standort eine zusätzliche Schutzschicht gegen Ransomware-Angriffe hinzu. Im Gegensatz zu Legacy-Lösungen/Ansätzen, bei denen eine Air-Gap-Lösung kompromittiert werden könnte, weil verschlüsselte/von Ransomware betroffene Daten auf das System im Air-Gap repliziert werden, beeinflusst die Replikation von Daten auf einen anderen Cohesity-Cluster/Standort die Air-Gap-Kopie nicht, weil das unveränderliche Dateisystem auch auf diesem Standort vorhanden ist."),
            
            ("heading2", "Auf maschinellem Lernen basierende Ransomware-Erkennung und umsetzbare Empfehlungen"),
            ("body", "In einer perfekten Welt sollten wir uns keine Sorgen über Ransomware-Angriffe machen müssen, aber leider ist das nicht unsere Welt heute. In einer Situation, in der Ihre primäre Umgebung, Benutzer und Anwendungsinfrastruktur kompromittiert sind, kann Cohesity Helios Ihnen aus dieser Klemme helfen. Mit seiner neuesten Anomalie-Erkennung bietet Helios, unsere SaaS-basierte, maschinengetriebene Lösung, Augen und Sichtbarkeit, wenn Sie nicht in der Lage sind. Mit SmartAssist warnt Helios nicht nur den IT-Administrator, sondern auch das Support-Team von Cohesity, wenn die primäre Dateiänderungsrate außerhalb der Norm liegt. Anomalien werden basierend auf der Übereinstimmung größerer Datenänderungen mit den normalen Mustern erkannt, einschließlich:"),
            
            ("bullet", "Tägliche Änderungsrate bei logischen Daten"),
            ("bullet", "Tägliche Änderungsrate bei gespeicherten Daten (nach Deduplizierung)"),
            ("bullet", "Muster basierend auf historischer Datenaufnahme"),
            ("bullet", "Entropie (Zufälligkeit der Daten)"),
            
            ("body", "Neben der Überwachung der Backup-Datenänderungsrate zur Erkennung eines potenziellen Ransomware-Angriffs helfen die Machine-Learning-Algorithmen von Cohesity auch dabei, eine saubere Kopie der Daten zu finden, die für die Wiederherstellung verwendet werden kann."),
            
            ("heading2", "Tiefe Sichtbarkeit für eine saubere Wiederherstellung"),
            ("body", "Wie gut ist eine Datenwiederherstellung, wenn sie dazu führt, dass Software-Schwachstellen und Cyber-Bedrohungen zurück in die IT-Produktionsumgebung eingeführt werden... dieselben Lücken, die Cyberkriminelle zuvor ausgenutzt haben, um leicht auf Ihre hochgesicherte IT-Umgebung zuzugreifen?"),
            
            ("body", "Cohesity CyberScan gibt Backup-Operatoren tiefe Sichtbarkeit in die Gesundheit und den Wiederherstellungsstatus ihrer Snapshots. Anstatt blind von einem beliebigen Snapshot wiederherzustellen, zeigt CyberScan den Schwachstellenindex jedes Snapshots und umsetzbare Empfehlungen zur Behebung dieser Software-Schwachstellen. Die Lösung wurde entwickelt, um Organisationen dabei zu helfen, nach einem Ransomware-Angriff sauber und vorhersehbar wiederherzustellen, ohne Schwachstellen in die IT-Produktionsumgebung zurück zu kompromittieren oder wieder einzuführen."),
            
            ("heading2", "Schnelle Wiederherstellung zur Reduzierung von Ausfallzeiten"),
            ("body", "Die wichtigste Anforderung nach einem Ransomware-Angriff ist die Fähigkeit, kompromittierte Daten schnell wiederherzustellen. Im Gegensatz zu jeder heute in der Branche verfügbaren Lösung bietet Cohesity die Möglichkeit, Daten über Ihren globalen Fußabdruck hinweg zu lokalisieren, einschließlich in der öffentlichen Cloud. Ihre Apps und Daten werden sofort mit der sofortigen Massen-Wiederherstellung von Cohesity zurückgebracht, indem sie Folgendes bieten:"),
            
            ("bullet", "Unbegrenzte Skalierbarkeit: Eine Web-Scale-Plattform, die IT-Administratoren ermöglicht, ihren Cohesity-Cluster von drei auf unbegrenzte Knoten zu erweitern, mit der Fähigkeit, unbegrenzte Snapshots und Klone ohne Leistungseinbußen zu speichern."),
            
            ("bullet", "Globale umsetzbare Suche: Im Gegensatz zu Legacy-Lösungen, die sich auf Drittanbieter-Suchprodukte verlassen, ermöglicht die einzigartige, Google-ähnliche globale Suchfunktion von Cohesity, Daten und infizierte Dateien schnell zu lokalisieren und entsprechende Korrekturmaßnahmen zu ergreifen. Dies umfasst das Auffinden einer bösartigen Datei über alle Workloads hinweg und das Ergreifen der notwendigen Maßnahmen, um sie einzudämmen."),
            
            ("bullet", "MegaFile: Ein patentierter Ansatz zur intelligenten Verteilung von Dateien über alle Knoten in einem Cluster. Ein Aspekt der Architektur von Cohesity, MegaFile, zerlegt große Dateien in kleinere Teile für parallele Sicherung und Wiederherstellung über Knoten hinweg. Die spezifische Größe dieser Teile ist einzigartig und für maximale Leistung optimiert."),
            
            ("bullet", "Sofortige Massen-Wiederherstellung: Wenn Ransomware zuschlägt, haben Sie es nicht mit einer, zwei oder ein paar VMs/Dateien zu tun, sondern vielmehr mit einem Disaster-Recovery-Szenario, in dem der IT-Administrator Hunderte von VMs wiederherstellen muss. Im Gegensatz zu anderen Backup-Lösungen, ob traditionell oder modern, kann es Tage, wenn nicht Wochen dauern, um wiederherzustellen. Mit der sofortigen Massen-Wiederherstellung von Cohesity können IT-Administratoren Hunderte von VMs sofort, in großem Maßstab und zu jedem Zeitpunkt wiederherstellen."),
            
            ("body", "Ransomware ist so weit verbreitet, dass es praktisch zu einem Haushaltsnamen geworden ist. Dennoch bleibt Ransomware eine gewaltige Bedrohung für Unternehmen, die eine moderne Lösung benötigen, die mehr bietet als nur die Fähigkeit, eine Bedrohung einfach zu erkennen."),
        ]
    }
]

def add_cohesity_styled_paragraph(doc, text, style_type):
    """Add a paragraph with Cohesity professional styling"""
    
    if style_type == "heading1":
        p = doc.add_heading(level=1)
        run = p.add_run(text)
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)  # Cohesity blue
        run.font.name = 'Calibri'
        p.space_after = Pt(12)
        
    elif style_type == "heading2":
        p = doc.add_heading(level=2)
        run = p.add_run(text)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)  # Cohesity blue
        run.font.name = 'Calibri'
        p.space_before = Pt(12)
        p.space_after = Pt(6)
        
    elif style_type == "body":
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        run.font.name = 'Calibri'
        p.space_after = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
    elif style_type == "bullet":
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        run.font.name = 'Calibri'
        p.space_after = Pt(6)
        
    elif style_type == "source":
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(102, 102, 102)  # Gray
        run.font.name = 'Calibri'
        p.space_after = Pt(12)

def create_german_document(article):
    """Create a professionally styled German Word document"""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add title (larger, blue)
    title = doc.add_heading(level=0)
    title_run = title.add_run(article['title_de'])
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 102, 204)  # Cohesity blue
    title_run.font.name = 'Calibri'
    title.space_after = Pt(6)
    
    # Add English title (smaller, gray)
    en_title_p = doc.add_paragraph()
    en_title_run = en_title_p.add_run(f"Original: {article['title_en']}")
    en_title_run.font.size = Pt(12)
    en_title_run.font.italic = True
    en_title_run.font.color.rgb = RGBColor(102, 102, 102)  # Gray
    en_title_run.font.name = 'Calibri'
    en_title_p.space_after = Pt(4)
    
    # Add source URL
    add_cohesity_styled_paragraph(doc, f"Quelle: {article['url']}", "source")
    
    # Add divider line
    doc.add_paragraph("_" * 80).runs[0].font.color.rgb = RGBColor(200, 200, 200)
    
    # Add content with proper styling
    for style_type, text in article['content']:
        add_cohesity_styled_paragraph(doc, text, style_type)
    
    # Add footer space
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_run = footer_p.add_run("Übersetzt von Jarvis - KI-Assistent für Eric Brown")
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_run.font.name = 'Calibri'
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document
    doc.save(article['filename'])
    print(f"✅ Created: {article['filename']}")

def main():
    print("🇩🇪 Creating German translations of Cohesity articles...")
    print()
    
    # Create all three documents
    for article in articles:
        create_german_document(article)
    
    print()
    print("✅ All German documents created successfully!")
    print()
    print("📄 Files created:")
    for article in articles:
        print(f"   - {article['filename']}")

if __name__ == "__main__":
    main()
