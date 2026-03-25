#!/usr/bin/env python3
"""Generate Cohesity / Commvault M&A Analysis Word Document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Color constants
DARK_BLUE = RGBColor(0x0B, 0x1D, 0x3A)
MED_BLUE = RGBColor(0x1A, 0x3C, 0x6E)
ACCENT_BLUE = RGBColor(0x2E, 0x75, 0xB6)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)
HEADER_BG = "0B1D3A"
ALT_ROW_BG = "EBF1F8"

OUTPUT_PATH = "/Users/ericbrown/.openclaw/workspace/reports/Cohesity_Commvault_MA_Analysis.docx"

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = DARK_GRAY


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.name = 'Calibri'
    run.font.size = Pt(size or 9.5)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        '  <w:top w:w="40" w:type="dxa"/>'
        '  <w:bottom w:w="40" w:type="dxa"/>'
        '  <w:left w:w="80" w:type="dxa"/>'
        '  <w:right w:w="80" w:type="dxa"/>'
        '</w:tcMar>'
    )
    tcPr.append(tcMar)


def add_styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW_BG)
            is_first = c_idx == 0
            align = WD_ALIGN_PARAGRAPH.LEFT if is_first else WD_ALIGN_PARAGRAPH.CENTER
            bold = any(kw in str(row_data[0]).lower() for kw in ['total', 'net', 'recommended'])
            set_cell_text(cell, val, bold=bold, color=DARK_GRAY, size=9, align=align)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE if level <= 2 else MED_BLUE
        run.font.name = 'Calibri'
    return h


def add_body(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    run.bold = bold
    return p


def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GRAY
    return p


# ═══════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CONFIDENTIAL")
run.font.size = Pt(12)
run.font.color.rgb = ACCENT_BLUE
run.font.name = 'Calibri'
run.font.letter_spacing = Pt(4)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Merger & Acquisition Analysis")
run.font.size = Pt(28)
run.font.color.rgb = DARK_BLUE
run.font.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Cohesity Acquisition of Commvault Systems")
run.font.size = Pt(18)
run.font.color.rgb = MED_BLUE
run.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\u2501" * 60)
run.font.color.rgb = ACCENT_BLUE
run.font.size = Pt(10)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("March 2026")
run.font.size = Pt(14)
run.font.color.rgb = MED_GRAY
run.font.name = 'Calibri'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prepared for Discussion Purposes Only")
run.font.size = Pt(10)
run.font.color.rgb = MED_GRAY
run.font.italic = True
run.font.name = 'Calibri'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════
add_heading("Table of Contents", level=1)
doc.add_paragraph()

toc_items = [
    ("1.", "Executive Summary"),
    ("2.", "Company Profiles"),
    ("3.", "Purchase Price Analysis"),
    ("4.", "Deal Structure & Financing"),
    ("5.", "Cost Synergies"),
    ("6.", "Comparable Transactions"),
    ("7.", "Strategic Rationale"),
    ("8.", "Key Risks"),
    ("9.", "Disclaimer"),
]

for num, title in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f"{num}  {title}")
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_BLUE
    run.font.name = 'Calibri'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════
add_heading("1. Executive Summary", level=1)

add_body(
    "This analysis evaluates a potential acquisition of Commvault Systems, Inc. (NASDAQ: CVLT) "
    "by Cohesity, Inc. The proposed transaction would create the undisputed #1 data protection "
    "platform globally with pro-forma revenue exceeding $2.8 billion."
)
add_body(
    "Commvault is currently trading at ~$79.41 per share, representing a market capitalization "
    "of approximately $3.5 billion \u2014 significantly below its 52-week high of $200.68. This "
    "dislocation creates a compelling acquisition opportunity at reasonable multiples."
)
add_body(
    "Our recommended offer range of $150\u2013165 per share implies an 89\u2013108% premium to the "
    "current trading price, valuing Commvault at a total enterprise value of $6.5\u20137.2 billion "
    "(5.5\u20136.0x LTM revenue). The deal would be financed through a combination of new debt "
    "($4.5B), Cohesity equity rollover ($2.0B), and sponsor equity ($500M)."
)
add_body(
    "We estimate run-rate cost synergies of $600 million by Year 3, which would reduce "
    "pro-forma leverage from 6.9x to 3.6x Debt/EBITDA \u2014 well within serviceable range "
    "for a business of this scale and recurring revenue profile."
)

add_heading("Transaction Overview", level=2)
add_styled_table(
    ["Metric", "Value"],
    [
        ["Target", "Commvault Systems (NASDAQ: CVLT)"],
        ["Acquirer", "Cohesity, Inc. (Private)"],
        ["Current Share Price", "$79.41 (Mar 20, 2026)"],
        ["Recommended Offer", "$150\u2013165 per share"],
        ["Implied Premium", "89\u2013108%"],
        ["Total Enterprise Value", "$6.5\u20137.2B"],
        ["EV / LTM Revenue", "5.5\u20136.0x"],
        ["EV / ARR", "5.7\u20136.3x"],
        ["Pro-Forma Combined Revenue", "~$2.8B"],
        ["Run-Rate Cost Synergies", "$600M (Year 3)"],
        ["Pro-Forma Leverage (at close)", "6.9x Debt/EBITDA"],
        ["Pro-Forma Leverage (with synergies)", "3.6x Debt/EBITDA"],
    ],
    col_widths=[3.0, 3.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. COMPANY PROFILES
# ═══════════════════════════════════════════════════════════════
add_heading("2. Company Profiles", level=1)

add_heading("2.1 Commvault Systems (Target \u2014 NASDAQ: CVLT)", level=2)
add_body(
    "Commvault is a leading enterprise data protection and management software company, "
    "currently executing a successful cloud/SaaS transition with SaaS ARR growing 44% year-over-year."
)

add_styled_table(
    ["Metric", "Value"],
    [
        ["Stock Price (Mar 20, 2026)", "~$79.41"],
        ["52-Week Range", "$79.51 \u2013 $200.68"],
        ["Shares Outstanding", "44.3M"],
        ["Market Capitalization", "~$3.5B"],
        ["Cash & Equivalents", "$1.03B"],
        ["Total Debt", "$920M"],
        ["Net Cash", "$106M"],
        ["LTM Revenue", "~$1.18B"],
        ["FY2026 Revenue Guidance", "$1,177\u20131,180M"],
        ["Total ARR", "$1,085M (+22% YoY)"],
        ["Subscription ARR", "$941M (+28% YoY)"],
        ["SaaS ARR Growth", "+44% YoY"],
        ["Non-GAAP EBIT Margin", "19\u201320%"],
        ["Non-GAAP Gross Margin", "~81%"],
        ["Free Cash Flow (Guided)", "$215\u2013220M"],
        ["Q3 FY2026 Revenue", "$314M (+19% YoY)"],
    ],
    col_widths=[3.0, 3.5]
)

add_heading("2.2 Cohesity (Acquirer \u2014 Private)", level=2)
add_body(
    "Cohesity is a leading AI-powered data security and management platform. Following "
    "its merger with Veritas Technologies in December 2024, Cohesity is the largest data "
    "protection company globally with approximately 6,000 employees."
)

add_styled_table(
    ["Metric", "Value"],
    [
        ["Valuation (Post-Veritas)", "$7B; IPO target ~$17B"],
        ["Pro-Forma Revenue", "~$1.6B (post-Veritas)"],
        ["ARR", "~$1.5B"],
        ["Revenue Growth (Pre-Veritas)", "~30%"],
        ["Adjusted Profit Margin", "28%"],
        ["Key Investors", "SoftBank, Sequoia Capital"],
        ["Veritas Merger", "Completed December 2024"],
        ["Employees", "~6,000"],
    ],
    col_widths=[3.0, 3.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. PURCHASE PRICE ANALYSIS
# ═══════════════════════════════════════════════════════════════
add_heading("3. Purchase Price Analysis", level=1)

add_heading("3.1 Valuation Methodology", level=2)
add_body(
    "Our valuation is based on EV/Revenue and EV/ARR multiples derived from comparable "
    "transactions in the enterprise software and data protection sectors."
)
add_bullet("SaaS M&A median revenue multiple: 3.8x (2025, per Aventis Advisors)")
add_bullet("Enterprise software PE acquisitions: 12\u201320x EBITDA")
add_bullet("High-growth SaaS (NRR >120%): up to 11.7x revenue")
add_bullet("Control premium for public targets: typically 25\u201340%")

add_heading("3.2 Valuation Scenarios", level=2)

add_styled_table(
    ["", "Low Case (4.5x)", "Base Case (6.0x)", "High Case (8.0x)"],
    [
        ["EV/Revenue Multiple", "4.5x", "6.0x", "8.0x"],
        ["Enterprise Value", "$5.3B", "$7.1B", "$9.4B"],
        ["Less: Debt", "($920M)", "($920M)", "($920M)"],
        ["Plus: Cash", "$1.03B", "$1.03B", "$1.03B"],
        ["Equity Value", "$5.4B", "$7.2B", "$9.5B"],
        ["Price Per Share", "~$122", "~$163", "~$215"],
        ["Premium to Current", "54%", "105%", "171%"],
    ],
    col_widths=[2.0, 1.7, 1.7, 1.7]
)

add_heading("3.3 Recommended Offer", level=2)
add_body("Recommended Offer: $150\u2013165 per share (Base Case)", bold=True)

add_styled_table(
    ["Parameter", "Value"],
    [
        ["Offer Range", "$150\u2013165 per share"],
        ["Implied Premium", "89\u2013108% to current price"],
        ["EV / Revenue", "5.5\u20136.0x"],
        ["EV / ARR", "5.7\u20136.3x"],
        ["Total Equity Value", "$6.6\u20137.3B"],
        ["Total Enterprise Value", "$6.5\u20137.2B"],
    ],
    col_widths=[3.0, 3.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. DEAL STRUCTURE
# ═══════════════════════════════════════════════════════════════
add_heading("4. Deal Structure & Financing", level=1)

add_heading("4.1 Sources of Funds", level=2)
add_body("The transaction will be financed through a mix of new debt, equity rollover, and sponsor equity:")

add_styled_table(
    ["Source", "Amount", "% of Total"],
    [
        ["New Term Loan B", "$3.0B", "43%"],
        ["High Yield Notes", "$1.5B", "21%"],
        ["Cohesity Equity Rollover", "$2.0B", "29%"],
        ["Sponsor Equity (SoftBank + co-investors)", "$500M", "7%"],
        ["Total Sources", "$7.0B", "100%"],
    ],
    col_widths=[3.0, 1.5, 1.5]
)

add_heading("4.2 Uses of Funds", level=2)

add_styled_table(
    ["Use", "Amount"],
    [
        ["Commvault Equity Purchase", "$7.2B"],
        ["Refinance Existing Commvault Debt", "$920M"],
        ["Less: Commvault Cash on Hand", "($1.03B)"],
        ["Transaction Fees", "~$100M"],
        ["Total Uses", "$7.2B"],
    ],
    col_widths=[4.0, 2.0]
)

add_heading("4.3 Pro-Forma Capital Structure", level=2)

add_styled_table(
    ["Metric", "Value"],
    [
        ["Total Debt", "$4.5B ($3.0B TLB + $1.5B HY)"],
        ["Pro-Forma Combined Revenue", "~$2.8B"],
        ["Pro-Forma Combined EBITDA (est.)", "~$650M (23% margin)"],
        ["Leverage Ratio", "6.9x Debt/EBITDA"],
        ["Assessment", "High but manageable with synergies"],
    ],
    col_widths=[3.0, 3.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. COST SYNERGIES
# ═══════════════════════════════════════════════════════════════
add_heading("5. Cost Synergies", level=1)

add_heading("5.1 Year 1\u20133 Synergy Bridge", level=2)
add_body(
    "We project total run-rate synergies of $600 million by Year 3, driven primarily by "
    "headcount rationalization across overlapping functions and infrastructure consolidation."
)

add_styled_table(
    ["Category", "Year 1", "Year 2", "Year 3 (Run-Rate)"],
    [
        ["Headcount Reduction (15\u201320% overlap)", "$150M", "$250M", "$350M"],
        ["Infrastructure / Data Center Consolidation", "$30M", "$60M", "$80M"],
        ["G&A Elimination (Duplicate Corporate)", "$40M", "$70M", "$90M"],
        ["Sales & Marketing Optimization", "$20M", "$50M", "$80M"],
        ["Total Synergies", "$240M", "$430M", "$600M"],
        ["Implementation Costs", "($200M)", "($100M)", "$0"],
        ["Net Synergies", "$40M", "$330M", "$600M"],
    ],
    col_widths=[2.8, 1.2, 1.2, 1.5]
)

add_heading("5.2 Deleveraging Impact", level=2)
add_body("With full run-rate synergies of $600M, the pro-forma financial profile improves significantly:")

add_styled_table(
    ["Metric", "At Close", "With Full Synergies"],
    [
        ["Pro-Forma EBITDA", "~$650M", "~$1,250M"],
        ["Total Debt", "$4.5B", "$4.5B"],
        ["Leverage (Debt/EBITDA)", "6.9x", "3.6x"],
        ["Interest Expense (8% blended)", "$360M", "$360M"],
        ["Interest Coverage", "~1.8x", "~3.5x"],
    ],
    col_widths=[2.5, 1.8, 1.8]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. COMPARABLE TRANSACTIONS
# ═══════════════════════════════════════════════════════════════
add_heading("6. Comparable Transactions", level=1)

add_body(
    "The following recent enterprise software transactions provide context for the proposed "
    "valuation multiples and premiums:"
)

add_styled_table(
    ["Deal (Acquirer / Target)", "Year", "EV/Revenue", "Premium"],
    [
        ["Cohesity / Veritas", "2024", "~4.5x", "N/A (private)"],
        ["Thoma Bravo / Darktrace", "2025", "~7.3x", "44%"],
        ["HPE / Juniper Networks", "2025", "~2.5x", "32%"],
        ["IBM / HashiCorp", "2024", "~9.5x", "64%"],
        ["Sophos / Secureworks", "2024", "~3.5x", "28%"],
    ],
    col_widths=[2.8, 1.0, 1.2, 1.2]
)

add_body(
    "Our recommended 5.5\u20136.0x EV/Revenue multiple positions the Commvault acquisition "
    "at a moderate premium relative to comparable deals, reflecting the company's strong "
    "SaaS transition momentum and recurring revenue base, while accounting for the current "
    "market dislocation in Commvault's share price."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. STRATEGIC RATIONALE
# ═══════════════════════════════════════════════════════════════
add_heading("7. Strategic Rationale", level=1)

rationale = [
    ("Market Leadership: ", "Creates the undisputed #1 data protection platform with $2.8B+ "
     "in combined revenue, far outpacing Veeam, Rubrik, and other competitors."),
    ("SaaS + AI Convergence: ", "Combines Commvault's accelerating SaaS growth (44% YoY SaaS ARR) "
     "with Cohesity's AI-powered data security platform, creating a differentiated offering."),
    ("Cross-Sell Opportunity: ", "Cohesity's 12,000+ customers combined with Commvault's deep "
     "enterprise installed base creates significant upsell and cross-sell potential."),
    ("Synergy-Justified Premium: ", "Run-rate cost synergies of $600M fully justify the acquisition "
     "premium and provide a clear path to deleveraging."),
    ("Competitive Elimination & IPO Positioning: ", "Removes a key competitor while positioning the "
     "combined entity for a premium IPO valuation as the dominant data protection platform."),
]

for i, (prefix, text) in enumerate(rationale, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. ")
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    run.font.bold = True
    run.font.name = 'Calibri'
    run2 = p.add_run(prefix)
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = ACCENT_BLUE
    run2.font.bold = True
    run2.font.name = 'Calibri'
    run3 = p.add_run(text)
    run3.font.size = Pt(10.5)
    run3.font.color.rgb = DARK_GRAY
    run3.font.name = 'Calibri'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. KEY RISKS
# ═══════════════════════════════════════════════════════════════
add_heading("8. Key Risks", level=1)

risks = [
    ("Integration Complexity",
     "Both companies have recently completed major integrations \u2014 Cohesity with Veritas "
     "(Dec 2024) and Commvault's ongoing cloud transformation. Executing a third major "
     "integration while digesting the prior ones creates significant operational risk."),
    ("High Initial Leverage",
     "The 6.9x Debt/EBITDA leverage at close requires flawless synergy execution to "
     "deleverage to sustainable levels. Any shortfall in synergy realization could strain "
     "debt service capacity and limit strategic flexibility."),
    ("Customer Churn Risk",
     "Overlapping customer bases may lead to confusion and churn during the integration "
     "period. Competitors (Veeam, Rubrik, Veritas legacy) will aggressively target "
     "displaced accounts."),
    ("Regulatory Scrutiny",
     "The combined entity would dominate the enterprise data protection market, potentially "
     "attracting antitrust scrutiny from the DOJ, FTC, or international regulators. "
     "Remedies could include divestitures that reduce synergy potential."),
]

for i, (title, desc) in enumerate(risks, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. {title}")
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_BLUE
    run.font.bold = True
    run.font.name = 'Calibri'
    p2 = doc.add_paragraph()
    run2 = p2.add_run(desc)
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = DARK_GRAY
    run2.font.name = 'Calibri'
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. DISCLAIMER
# ═══════════════════════════════════════════════════════════════
add_heading("9. Disclaimer", level=1)

disclaimer_text = (
    "This analysis is for discussion purposes only and does not constitute an offer, "
    "solicitation, or recommendation to buy or sell any securities. The information "
    "contained herein is based on publicly available data and certain assumptions that "
    "may not prove to be correct. No representation or warranty, express or implied, "
    "is made as to the accuracy, completeness, or fairness of the information or "
    "opinions contained in this document.\n\n"
    "This document has been prepared solely for informational purposes and is not "
    "intended to be, and should not be construed as, financial, legal, tax, or "
    "investment advice. Recipients should conduct their own independent analysis and "
    "consult with their own advisors before making any investment decision.\n\n"
    "Past performance is not indicative of future results. Forward-looking statements "
    "and projections are inherently uncertain and actual results may differ materially "
    "from those projected. The authors disclaim any liability for losses incurred in "
    "connection with any action taken based on the information contained herein."
)

p = doc.add_paragraph()
run = p.add_run(disclaimer_text)
run.font.size = Pt(9)
run.font.color.rgb = MED_GRAY
run.font.italic = True
run.font.name = 'Calibri'

doc.save(OUTPUT_PATH)
print(f"Document saved to: {OUTPUT_PATH}")
