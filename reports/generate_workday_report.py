#!/usr/bin/env python3
"""Generate Workday FY26 Earnings Analysis as Word Document."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT = os.path.expanduser("~/Documents/reports/Workday_Q4_FY2026_Earnings_Analysis.docx")
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)

doc = Document()

# ── Styles ──
for level in range(1, 4):
    style = doc.styles[f"Heading {level}"]
    style.font.color.rgb = RGBColor(0, 99, 178)

def add_table(doc, data):
    rows, cols = len(data), len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, val in enumerate(data[0]):
        cell = table.rows[0].cells[j]
        cell.text = str(val)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for i in range(1, rows):
        for j in range(cols):
            cell = table.rows[i].cells[j]
            cell.text = str(data[i][j])
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph("")

def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")

# ═══════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════
doc.add_heading("Workday, Inc. (NASDAQ: WDAY)", level=0)
doc.add_heading("Q4 FY2026 Earnings Analysis — Sell-Side Deep Dive", level=1)
p = doc.add_paragraph("March 27, 2026 | Enterprise AI & HCM Coverage", style="Intense Quote")
doc.add_paragraph("")

# ═══════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "Workday delivered solid Q4 FY2026 results (quarter ended January 31, 2026), "
    "reporting total revenue of $2.532 billion (+14.5% YoY) and subscription revenue of "
    "$2.360 billion (+15.7% YoY). Non-GAAP operating margin expanded 420 basis points YoY "
    "to 30.6% in Q4, and 370 bps for the full year to 29.6%. Free cash flow surged 26.7% "
    "to $2.777 billion (29.1% FCF margin), establishing Workday as one of the most cash-generative "
    "enterprise SaaS companies at scale."
)
doc.add_paragraph(
    "For the full fiscal year 2026, Workday generated $9.552 billion in total revenue (+13.1% YoY) "
    "and $8.833 billion in subscription revenue (+14.5% YoY). The company repurchased $2.9 billion "
    "in shares and ended the year with $5.443 billion in cash and marketable securities."
)
doc.add_paragraph(
    "The most significant strategic development was the March 2026 launch of 'Sana from Workday' — "
    "a conversational AI platform acquired in Q4 FY26 and positioned as 'superintelligence for work.' "
    "With co-founder Aneel Bhusri returning as CEO and an explicit AI pivot underway, Workday is "
    "betting that its trusted position in HR and finance workflows gives it a unique advantage in "
    "enterprise AI agent deployment."
)

# ═══════════════════════════════════════════════
# 2. QUARTERLY RESULTS TABLE
# ═══════════════════════════════════════════════
doc.add_heading("2. Quarterly Results Table", level=1)
doc.add_paragraph("All figures in USD millions except per-share data and percentages.")
add_table(doc, [
    ["Metric", "Q4 FY26", "Q3 FY26", "Q4 FY25", "YoY Change"],
    ["Total Revenue", "$2,532", "$2,498 (est)", "$2,210 (est)", "+14.5%"],
    ["Subscription Revenue", "$2,360", "$2,320 (est)", "$2,040 (est)", "+15.7%"],
    ["GAAP Operating Income", "$174", "—", "$75", "+132%"],
    ["GAAP Op Margin", "6.9%", "—", "3.4%", "+350 bps"],
    ["Non-GAAP Operating Income", "$774", "—", "$584", "+32.5%"],
    ["Non-GAAP Op Margin", "30.6%", "—", "26.4%", "+420 bps"],
    ["GAAP Diluted EPS", "$0.55", "—", "$0.35", "+57%"],
    ["Non-GAAP Diluted EPS", "$2.47", "—", "$1.92", "+29%"],
])

doc.add_heading("Full Year FY2026 Results", level=2)
add_table(doc, [
    ["Metric", "FY2026", "FY2025", "YoY Change"],
    ["Total Revenue", "$9,552", "$8,446 (est)", "+13.1%"],
    ["Subscription Revenue", "$8,833", "$7,713 (est)", "+14.5%"],
    ["GAAP Operating Income", "$721", "$415", "+73.7%"],
    ["GAAP Op Margin", "7.5%", "4.9%", "+260 bps"],
    ["Non-GAAP Operating Income", "$2,824", "$2,186", "+29.2%"],
    ["Non-GAAP Op Margin", "29.6%", "25.9%", "+370 bps"],
    ["GAAP Diluted EPS", "$2.59", "$1.95", "+33%"],
    ["Non-GAAP Diluted EPS", "$9.23", "$7.30", "+26%"],
    ["Operating Cash Flow", "$2,939", "$2,462 (est)", "+19.4%"],
    ["Free Cash Flow", "$2,777", "$2,191 (est)", "+26.7%"],
    ["Cash & Securities", "$5,443", "—", "—"],
    ["Share Repurchases", "$2,900", "—", "~12.8M shares"],
])

# ═══════════════════════════════════════════════
# 3. GUIDANCE ANALYSIS
# ═══════════════════════════════════════════════
doc.add_heading("3. Guidance Analysis", level=1)
doc.add_heading("FY2027 Guidance", level=2)
doc.add_paragraph(
    "Subscription Revenue: ~$9.925 billion to $9.950 billion, representing 12% to 13% growth "
    "year-over-year. At the midpoint of ~$9.94 billion, implied subscription revenue growth is "
    "approximately 12.6% YoY — a modest deceleration from the 14.5% posted in FY2026."
)
doc.add_paragraph(
    "Non-GAAP Operating Margin: Expected to continue expanding from the 29.6% achieved in FY26, "
    "likely targeting the 30%+ range. Management has consistently demonstrated the ability to "
    "expand margins while sustaining double-digit revenue growth."
)
doc.add_paragraph(
    "The guidance implies continued disciplined growth with margin expansion, consistent with "
    "Workday's transition from a hyper-growth cloud company to a mature, large-cap SaaS platform. "
    "The key bull case is whether the Sana AI platform can re-accelerate growth beyond the "
    "guided 12-13% range in FY28 and beyond."
)

# ═══════════════════════════════════════════════
# 4. REVENUE & BACKLOG DEEP DIVE
# ═══════════════════════════════════════════════
doc.add_heading("4. Revenue & Backlog Deep Dive", level=1)
doc.add_paragraph(
    "Subscription Revenue Trajectory: Workday's subscription revenue growth has moderated from "
    "~18% in FY24 to ~14.5% in FY26, reflecting the law of large numbers at $8.8B+ in subscription "
    "revenue. However, absolute dollar growth remains substantial — Workday added over $1.1 billion "
    "in net new subscription revenue in FY26."
)
doc.add_heading("Backlog Metrics", level=2)
add_table(doc, [
    ["Metric", "FY2026", "FY2025", "YoY Change"],
    ["12-Month Sub Backlog", "$8,833", "$7,630 (est)", "+15.8%"],
    ["Total Sub Backlog", "$28,101", "$25,044 (est)", "+12.2%"],
])
doc.add_paragraph(
    "12-month subscription revenue backlog of $8.833 billion (up 15.8% YoY) provides strong "
    "visibility into FY27 subscription revenue guidance of ~$9.94 billion. Total backlog of "
    "$28.1 billion represents ~3.2 years of forward subscription revenue coverage. Both backlog "
    "figures include the impact from acquisitions of Paradox (Q3) and Sana (Q4)."
)
doc.add_heading("Customer Base", level=2)
add_bullet(doc, "11,500+ organizations worldwide")
add_bullet(doc, "65%+ of the Fortune 500 are customers")
add_bullet(doc, "Recent wins: Fairview Health Services (HR, Finance, Supply Chain on one platform)")
add_bullet(doc, "Growing retail/hospitality vertical momentum noted in Q4")

# ═══════════════════════════════════════════════
# 5. PROFITABILITY ANALYSIS
# ═══════════════════════════════════════════════
doc.add_heading("5. Profitability Analysis", level=1)
doc.add_heading("Margin Progression", level=2)
doc.add_paragraph(
    "Non-GAAP Operating Margin expanded from 25.9% in FY25 to 29.6% in FY26 (+370 bps), "
    "demonstrating continued operating leverage. Q4 FY26 reached 30.6%, the first quarter "
    "above 30% — a significant milestone for a company at this scale."
)
doc.add_paragraph(
    "GAAP profitability was impacted by $303 million in restructuring charges in FY26 "
    "(vs $84 million in FY25). Excluding restructuring, GAAP operating margin would have been "
    "~10.7%, indicating the underlying business is firmly profitable on a GAAP basis as well."
)
doc.add_heading("Rule of 40 Assessment", level=2)
doc.add_paragraph(
    "FY2026 Rule of 40: Revenue growth (13.1%) + FCF margin (29.1%) = 42.2. Workday comfortably "
    "exceeds the Rule of 40 threshold, driven by best-in-class free cash flow generation. "
    "This positions WDAY favorably among large-cap SaaS peers."
)

# ═══════════════════════════════════════════════
# 6. CASH FLOW ANALYSIS
# ═══════════════════════════════════════════════
doc.add_heading("6. Cash Flow Analysis", level=1)
doc.add_paragraph(
    "Cash flow generation is a standout strength for Workday. Full-year FY2026 operating cash "
    "flow was $2.939 billion (+19.4% YoY) and free cash flow was $2.777 billion (+26.7% YoY), "
    "representing a 29.1% FCF margin — among the highest in enterprise software."
)
add_table(doc, [
    ["Metric", "FY2026", "FY2025 (est)", "YoY Change"],
    ["Operating Cash Flow", "$2,939M", "$2,462M", "+19.4%"],
    ["Free Cash Flow", "$2,777M", "$2,191M", "+26.7%"],
    ["FCF Margin", "29.1%", "25.9%", "+320 bps"],
    ["Cash & Securities", "$5,443M", "—", "—"],
])
doc.add_paragraph(
    "Capital allocation priorities are clear: Workday repurchased ~12.8 million shares for $2.9 billion "
    "in FY26, partially offsetting SBC dilution while signaling management confidence. The $5.4 billion "
    "cash balance provides ample capacity for continued buybacks and strategic M&A."
)

# ═══════════════════════════════════════════════
# 7. PRODUCT & STRATEGY HIGHLIGHTS
# ═══════════════════════════════════════════════
doc.add_heading("7. Product & Strategy Highlights", level=1)
doc.add_heading("Sana from Workday — The AI Bet", level=2)
doc.add_paragraph(
    "The most important strategic development in Q4/early FY27 is the launch of 'Sana from Workday,' "
    "a conversational AI platform described as 'superintelligence for work.' Workday acquired Sana "
    "(an AI startup founded by Joel Hellermark) in Q4 FY26 and launched the integrated product in "
    "March 2026."
)
doc.add_paragraph("Three-tier product structure:")
add_bullet(doc, "Sana for Workday — New conversational AI interface replacing traditional Workday UI navigation")
add_bullet(doc, "Sana Self-Service Agent — Launches with 300+ skills for HR and finance self-service tasks")
add_bullet(doc, "Sana Enterprise — Connects Workday with third-party apps including Gmail, Salesforce, SharePoint, Box, Confluence, Jira, Slack, Zoom, Notion, and more")
doc.add_paragraph(
    "Monetization: Available to ALL existing Workday customers via Workday Flex Credits at no extra "
    "license cost. This is a land-wide/retain-first strategy, similar to Salesforce's initial "
    "Agentforce positioning — drive adoption now, monetize on consumption later."
)
doc.add_paragraph(
    "Early traction is promising: customer Berner reported 90% adoption within 40 days and retired "
    "400 ChatGPT licenses."
)

doc.add_heading("CEO Transition", level=2)
doc.add_paragraph(
    "Aneel Bhusri, co-founder, returned as CEO with an explicit AI focus: 'We built Workday to bring "
    "innovation back to the worlds of HR and finance, and AI gives us the chance to do it all again.' "
    "Zane Rowe continues as CFO. Gerrit Kazmaier serves as President, Product & Technology."
)

doc.add_heading("Key Product Launches", level=2)
add_bullet(doc, "Sana from Workday — Agentic AI platform (March 2026)")
add_bullet(doc, "Paradox ATS Integration — Conversational applicant tracking (Q3 acquisition)")
add_bullet(doc, "Military Skills Mapper — Veteran career transition tool")
add_bullet(doc, "Insperity HRScale — HR automation partnership")

# ═══════════════════════════════════════════════
# 8. COMPETITIVE LANDSCAPE
# ═══════════════════════════════════════════════
doc.add_heading("8. Competitive Landscape", level=1)
doc.add_paragraph(
    "Workday operates in the large, consolidated enterprise HCM and financial management market. "
    "Key competitors include SAP SuccessFactors, Oracle HCM Cloud, ServiceNow, and Microsoft "
    "Dynamics 365. Workday's competitive advantage rests on three pillars: (1) unified HR + Finance "
    "platform on a single data model, (2) 65%+ Fortune 500 penetration as an installed base moat, "
    "and (3) the Sana AI agent platform as a differentiated next-gen interface."
)
add_table(doc, [
    ["Company", "Focus", "AI Strategy", "Scale"],
    ["Workday (WDAY)", "HCM + Finance", "Sana AI agents, 300+ skills", "$9.6B rev, 11,500+ customers"],
    ["ServiceNow (NOW)", "IT + Workflows", "Now Assist AI copilots", "$11.0B+ rev"],
    ["SAP (SAP)", "ERP + HCM", "Joule AI assistant", "$35B+ rev (cloud growing)"],
    ["Oracle (ORCL)", "ERP + HCM + Cloud", "Oracle AI embedded", "$56B+ rev"],
])
doc.add_paragraph(
    "The Sana acquisition positions Workday uniquely with a conversational AI layer that spans "
    "Workday and third-party apps — a broader scope than most competitors' AI assistants which "
    "are limited to their own ecosystems."
)

# ═══════════════════════════════════════════════
# 9. VALUATION
# ═══════════════════════════════════════════════
doc.add_heading("9. Valuation", level=1)
doc.add_paragraph("Current Trading (as of March 27, 2026):")
add_bullet(doc, "Stock Price: $124.21 (NASDAQ: WDAY)")
add_bullet(doc, "Market Cap: ~$32.5 billion (based on ~262M diluted shares)")
add_bullet(doc, "Enterprise Value: ~$27.1 billion (net cash ~$5.4B)")
add_bullet(doc, "41 sell-side analysts covering the name")

doc.add_heading("Valuation Multiples", level=2)
add_table(doc, [
    ["Multiple", "Value", "Context"],
    ["EV/NTM Revenue", "~2.5x", "Discount to ServiceNow (~10x), Salesforce (~5x)"],
    ["EV/NTM FCF", "~9.8x", "Attractive for 13% grower with 29% FCF margins"],
    ["P/E Non-GAAP (FY26)", "~13.5x", "Below software peer median"],
    ["EV/NTM Sub Revenue", "~2.7x", "Strong backlog coverage"],
])
doc.add_paragraph(
    "At ~2.5x EV/NTM revenue, Workday is trading at a meaningful discount to high-growth SaaS peers, "
    "reflecting the growth deceleration from ~20%+ (FY22-23) to ~13% today. The Sana/AI narrative is "
    "the key re-rating catalyst — if it drives re-acceleration to 15%+ revenue growth, the multiple "
    "should expand materially."
)

# ═══════════════════════════════════════════════
# 10. KEY RISKS & CONCERNS
# ═══════════════════════════════════════════════
doc.add_heading("10. Key Risks & Concerns", level=1)
doc.add_paragraph(
    "Growth Deceleration: Revenue growth has slowed from ~20%+ (FY22-23) to ~13% (FY26). FY27 "
    "guidance implies further modest deceleration (12-13%). The bull case requires Sana/AI "
    "monetization to reverse this trend by FY28."
)
doc.add_paragraph(
    "Sana Monetization Timing: The initial 'free with Flex Credits' packaging is smart for adoption "
    "but delays direct revenue recognition. Watch FY27 for consumption-based pricing signals and "
    "Sana-specific revenue disclosures."
)
doc.add_paragraph(
    "Restructuring Charges: $303 million in FY26 (vs $84M in FY25) — elevated but expected to "
    "normalize in FY27, providing ~$200M+ tailwind to GAAP earnings."
)
doc.add_paragraph(
    "Competitive Intensity: ServiceNow, SAP SuccessFactors, Oracle HCM, and Microsoft Dynamics "
    "are all investing heavily in AI agents. Sana's differentiation is the trusted Workday data "
    "layer — this moat claim needs monitoring."
)
doc.add_paragraph(
    "Macro Exposure: Large enterprise HR/Finance software is resilient but not immune. Any slowdown "
    "in Fortune 500 hiring or IT spend could pressure net new ACV."
)

# ═══════════════════════════════════════════════
# 11. ANALYST RATINGS SUMMARY
# ═══════════════════════════════════════════════
doc.add_heading("11. Analyst Coverage", level=1)
doc.add_paragraph(
    "Workday is covered by 41 sell-side analysts from major firms including Goldman Sachs, "
    "JPMorgan, Morgan Stanley, Bank of America, Barclays, UBS, Wells Fargo, Deutsche Bank, "
    "Evercore, KeyBanc, RBC Capital Markets, TD Cowen, Piper Sandler, and others."
)
doc.add_paragraph(
    "Consensus sentiment is broadly constructive on the AI opportunity but watchful of growth "
    "deceleration. The Sana launch in March 2026 and the CEO transition are the two most recent "
    "catalysts under active debate."
)

# ── Footer ──
doc.add_paragraph("")
doc.add_paragraph("")
p = doc.add_paragraph("Data Sources: investor.workday.com, newsroom.workday.com | Crawled: March 27, 2026")
p.runs[0].font.size = Pt(8)
p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
p = doc.add_paragraph("Tool: Hybrid Crawler v3 (ProjectScraper + Firecrawl) | Analysis: Claude Opus 4.6")
p.runs[0].font.size = Pt(8)
p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
p = doc.add_paragraph("Sent by Jarvis — AI assistant to Eric Brown")
p.runs[0].font.size = Pt(8)
p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

doc.save(OUTPUT)
print(f"Report saved to: {OUTPUT}")
print(f"Size: {os.path.getsize(OUTPUT) / 1024:.1f} KB")
