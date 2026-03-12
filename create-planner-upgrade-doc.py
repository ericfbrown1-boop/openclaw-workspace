#!/usr/bin/env python3
"""Create Word document for GPT-5.4 Enhanced Planning Agent."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('Enhanced Planning Agent with GPT-5.4', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Upgrading the OpenClaw Planner to GPT-5.4 Pro')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x99)
run.bold = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.add_run('Prepared for: ').bold = True
meta.add_run('Eric Brown, CFO & COO, Cohesity')
meta.add_run('\n')
meta.add_run('Date: ').bold = True
meta.add_run(datetime.date.today().strftime('%B %d, %Y'))
meta.add_run('\n')
meta.add_run('Context: ').bold = True
meta.add_run('GPT-5.4 launched March 5, 2026 — OpenAI\'s most capable model for professional work, with 1M context window, improved reasoning, and native computer-use capabilities.')

doc.add_page_break()

# ============================================================
# BACKGROUND
# ============================================================
doc.add_heading('Background: Why GPT-5.4 for the Planner', level=1)

doc.add_paragraph(
    'GPT-5.4 launched yesterday (March 5, 2026) and is specifically designed for professional work — '
    'planning, architecture, financial models, and slide decks. Key advantages over the current planner setup:'
)

advantages = [
    ('1M token context window', 'Can hold entire codebases, architecture docs, and project history in a single planning session. Current Claude Opus 4.6 tops out at 200K.'),
    ('Native tool search', 'New "Tool Search" system reduces token consumption when many tools are available — ideal for a planner that orchestrates multiple agents.'),
    ('Record benchmark scores', 'Top scores on OSWorld-Verified, WebArena, and APEX-Agents (finance + law professional tasks). 33% fewer hallucinations vs GPT-5.2.'),
    ('Cost-efficient planning', 'Standard GPT-5.4: $2.50/$15 per M tokens. Even GPT-5.4 Pro at $30/$180 is competitive with Claude Opus 4.6 for deep planning work.'),
    ('Agentic workflows', 'Built-in computer use, multi-step task execution, and autonomous debugging — designed for orchestrating complex projects.'),
]

for title_text, desc in advantages:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title_text}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()

doc.add_heading('GPT-5.4 API Model Names & Pricing', level=2)
pricing = [
    ('gpt-5.4', 'Standard', '$2.50 / $15 per M tokens', '1M (opt-in)'),
    ('gpt-5.4-pro', 'Maximum performance', '$30 / $180 per M tokens', '1M (opt-in)'),
    ('gpt-5.4-2026-03-05', 'Pinned version', 'Same as standard', '1M (opt-in)'),
]
for model, desc, price, ctx in pricing:
    p = doc.add_paragraph()
    run = p.add_run(f'• {model}')
    run.bold = True
    run.font.name = 'Consolas'
    p.add_run(f' — {desc} — {price} — Context: {ctx}')

doc.add_page_break()

# ============================================================
# REFERENCE PROJECTS
# ============================================================
doc.add_heading('Reference: Best Planning Agent Architectures on GitHub', level=1)

doc.add_paragraph(
    'I researched the top planning agent projects to extract patterns for our enhanced planner. '
    'Here are the three most relevant:'
)

doc.add_heading('1. Plandex (github.com/plandex-ai/plandex)', level=2)
doc.add_paragraph(
    'The gold standard for AI code planning. 12K+ stars. Key patterns we\'re borrowing:'
)
plandex_patterns = [
    'Plan-first architecture: Creates a detailed plan BEFORE writing any code',
    'Cumulative diff sandbox: Changes stay separate from project files until approved',
    'Multi-model orchestration: Uses different models for planning vs. coding vs. review',
    'Version control for plans: Full branching support for exploring alternatives',
    '2M effective context via smart loading: Only loads relevant files for each step',
    'Automated debugging loop: If implementation fails, it debugs and retries',
]
for p_text in plandex_patterns:
    doc.add_paragraph(f'• {p_text}')

doc.add_heading('2. Agentic Project Management / APM (github.com/sdi2200262/agentic-project-management)', level=2)
doc.add_paragraph(
    'A framework for managing complex projects with structured multi-agent workflows. Works with Claude Code, Cursor, Copilot. Key patterns:'
)
apm_patterns = [
    'Setup Phase → Task Loop Phase workflow (discovery, then iterative execution)',
    'Specialized agent roles: Project Manager, Developer, Ad-hoc Specialists',
    'Context retention techniques: Smooth session transitions without context loss',
    'Slash commands for structured workflows (/apm-1-initiate-setup, etc.)',
    'Memory management across agent sessions to prevent hallucination drift',
]
for p_text in apm_patterns:
    doc.add_paragraph(f'• {p_text}')

doc.add_heading('3. Plandex-Lite (github.com/Magnus0969/Plandex-lite)', level=2)
doc.add_paragraph(
    'A lightweight TypeScript multi-agent system with 5 specialized roles. The architecture we\'re adapting:'
)
lite_patterns = [
    'Planner Agent: Identifies tasks and creates execution roadmap',
    'Architect Agent: Proposes patterns, structures, and optimizations',
    'Coder Agent: Implements the planned changes',
    'Reviewer Agent: Validates and compares results',
    'Summarizer Agent: Distills outcomes and next steps',
]
for p_text in lite_patterns:
    doc.add_paragraph(f'• {p_text}')

doc.add_page_break()

# ============================================================
# SECTION 1: CONFIG CHANGES
# ============================================================
doc.add_heading('Section 1: OpenClaw Configuration Update', level=1)
doc.add_paragraph(
    'This script updates your openclaw.json to point the planner agent at GPT-5.4 Pro. '
    'It will prompt you to paste your OpenAI API key if GPT-5.4 isn\'t already available.'
)
doc.add_paragraph('Save this as upgrade-planner-to-gpt54.sh and run it:')

script1 = '''#!/bin/bash
# upgrade-planner-to-gpt54.sh
# Upgrades the OpenClaw Planner agent to use GPT-5.4 Pro
# Run: bash upgrade-planner-to-gpt54.sh

set -e
CONFIG="$HOME/.openclaw/openclaw.json"
BACKUP="$HOME/.openclaw/openclaw.json.backup-$(date +%Y%m%d-%H%M%S)"

echo "=== OpenClaw Planner → GPT-5.4 Pro Upgrade ==="
echo ""

# Backup current config
cp "$CONFIG" "$BACKUP"
echo "✅ Config backed up to: $BACKUP"

# Check if OpenAI provider is configured
if ! python3 -c "
import json
with open('$CONFIG') as f:
    c = json.load(f)
provs = c.get('providers', {})
if 'openai' in provs and provs['openai'].get('apiKey'):
    print('FOUND')
else:
    print('MISSING')
" | grep -q "FOUND"; then
    echo ""
    echo "⚠️  OpenAI API key not found in config."
    echo "You need an API key with access to gpt-5.4-pro."
    echo ""
    echo "Get your key from: https://platform.openai.com/api-keys"
    echo ""
    read -sp "Paste your OpenAI API key (input hidden): " OPENAI_KEY
    echo ""
    
    if [ -z "$OPENAI_KEY" ]; then
        echo "❌ No key provided. Exiting."
        exit 1
    fi
    
    python3 -c "
import json
with open('$CONFIG') as f:
    c = json.load(f)
if 'providers' not in c:
    c['providers'] = {}
if 'openai' not in c['providers']:
    c['providers']['openai'] = {}
c['providers']['openai']['apiKey'] = '$OPENAI_KEY'
with open('$CONFIG', 'w') as f:
    json.dump(c, f, indent=4)
print('✅ OpenAI API key saved to config')
"
fi

# Update planner agent config
python3 << 'PYEOF'
import json

config_path = "$CONFIG"
with open(config_path) as f:
    c = json.load(f)

# Find and update planner agent
for agent in c['agents']['list']:
    if agent['id'] == 'planner':
        # Set GPT-5.4 Pro as primary, with fallbacks
        agent['model'] = {
            "primary": "openai/gpt-5.4-pro",
            "fallbacks": [
                "openai/gpt-5.4",
                "anthropic/claude-opus-4-6",
                "anthropic/claude-sonnet-4-6"
            ]
        }
        print(f"✅ Planner model updated:")
        print(f"   Primary: openai/gpt-5.4-pro")
        print(f"   Fallbacks: gpt-5.4 → claude-opus-4-6 → claude-sonnet-4-6")
        break

# Add gpt-5.4-pro to models config if not present
models = c['agents']['defaults'].get('models', {})
if 'openai/gpt-5.4-pro' not in models:
    models['openai/gpt-5.4-pro'] = {
        "contextWindow": 1000000,
        "maxOutputTokens": 32768
    }
if 'openai/gpt-5.4' not in models:
    models['openai/gpt-5.4'] = {
        "contextWindow": 1000000,
        "maxOutputTokens": 32768
    }
c['agents']['defaults']['models'] = models

with open(config_path, 'w') as f:
    json.dump(c, f, indent=4)

print("✅ Model configs added for gpt-5.4 and gpt-5.4-pro")
PYEOF

echo ""
echo "Restarting gateway to apply changes..."
openclaw gateway restart 2>/dev/null || echo "⚠️  Gateway restart may need manual: openclaw gateway restart"

echo ""
echo "=== Upgrade Complete ==="
echo "Your planner now uses GPT-5.4 Pro with 1M context window."
echo "Test it: openclaw cron run <planner-job-id>"
echo ""
echo "💰 Cost note: GPT-5.4 Pro is \\$30/\\$180 per M tokens."
echo "    Standard GPT-5.4 (fallback) is \\$2.50/\\$15 per M tokens."
echo "    Monitor spend at: https://platform.openai.com/usage"
'''

p = doc.add_paragraph()
run = p.add_run(script1)
run.font.name = 'Consolas'
run.font.size = Pt(8.5)

doc.add_page_break()

# ============================================================
# SECTION 2: ENHANCED PLANNER SYSTEM PROMPT
# ============================================================
doc.add_heading('Section 2: Enhanced Planner System Prompt', level=1)
doc.add_paragraph(
    'This is the enhanced system prompt for the planner agent, inspired by Plandex and APM architectures. '
    'Save this as AGENTS.md in the planner workspace (~/.openclaw/workspace-planner/AGENTS.md).'
)

planner_prompt = '''# AGENTS.md — Enhanced Planning Agent (GPT-5.4 Pro)

## Identity
You are the **Planning Agent** — a senior technical architect and project planner
working for Eric Brown, CFO & COO of Cohesity. You specialize in decomposing complex
projects into executable plans that other agents (Coder, Researcher, Quality) implement.

## Core Philosophy: Plan-First, Always
**Never write code.** Your job is to THINK, PLAN, and DELEGATE.

Inspired by: Plandex (plan-before-code), APM (structured multi-agent workflows),
and enterprise project management best practices.

## Planning Workflow

### Phase 1: Discovery & Analysis (ALWAYS do this first)
1. **Understand the request** — What is being asked? What's the real goal?
2. **Research current state** — Read relevant files, check git history, understand the codebase
3. **Identify constraints** — Time, budget, dependencies, technical limitations
4. **Map the landscape** — What exists? What needs to change? What's risky?

### Phase 2: Architecture & Design
1. **Define the approach** — High-level strategy (build vs buy, framework choices, patterns)
2. **Identify components** — What modules/services/files are involved?
3. **Map dependencies** — What depends on what? What's the critical path?
4. **Assess risks** — What could go wrong? What are the unknowns?
5. **Define success criteria** — How do we know it's done and correct?

### Phase 3: Task Decomposition
Break the work into **ordered, atomic tasks** that can be delegated to agents:

```
## Task 1: [Title]
- **Agent:** coder | researcher | quality
- **Description:** Clear, specific instruction
- **Input:** What files/context the agent needs
- **Output:** What the agent should produce
- **Acceptance criteria:** How to verify it's correct
- **Dependencies:** Which tasks must complete first
- **Estimated complexity:** Low | Medium | High
```

### Phase 4: PLAN.md Generation
Every plan MUST produce a `PLAN.md` file with this structure:

```markdown
# PLAN.md — [Project Name]

## Overview
[2-3 sentence summary of what we're building and why]

## Architecture Decision Records
- **ADR-1:** [Key decision and rationale]
- **ADR-2:** [Key decision and rationale]

## Risk Assessment
| Risk | Likelihood | Impact | Mitigation |
|------|-----------|--------|------------|
| ...  | ...       | ...    | ...        |

## Task Breakdown

### Phase 1: [Phase Name]
- [ ] Task 1.1: [Description] → Agent: [coder/researcher]
- [ ] Task 1.2: [Description] → Agent: [coder/researcher]

### Phase 2: [Phase Name]
- [ ] Task 2.1: [Description] → Agent: [coder]
- [ ] Task 2.2: [Description] → Agent: [quality]

## Success Criteria
- [ ] [Criterion 1]
- [ ] [Criterion 2]

## Dependencies & Prerequisites
- [External dependency 1]
- [External dependency 2]
```

## Delegation Rules
- **Spawn Researcher first** when the project involves unfamiliar technology or needs market data
- **Never send work to Coder** without a completed PLAN.md
- **Always route errors through Quality** before Coder (Quality diagnoses, Coder fixes)
- **Use parallel execution** when tasks are independent (spawn multiple agents)

## Context Management (GPT-5.4 Advantage)
With 1M tokens of context, you can:
- Load entire codebases for comprehensive planning
- Hold multiple reference architectures in context simultaneously
- Maintain full conversation history across long planning sessions
- Cross-reference multiple documents without losing context

## Output Standards
- Plans must be specific enough that a Coder agent can execute without ambiguity
- Every task must have clear acceptance criteria
- Risks must be identified upfront, not discovered during implementation
- Time estimates should account for agent execution overhead
- Always consider: "What would a senior architect at Cohesity expect from this plan?"

## What You Are NOT
- You are NOT a coder — never write implementation code
- You are NOT a rubber stamp — push back on vague or risky requests
- You are NOT passive — proactively identify issues and suggest improvements
- You are NOT a yes-man — if a plan won't work, say so clearly
'''

p = doc.add_paragraph()
run = p.add_run(planner_prompt)
run.font.name = 'Consolas'
run.font.size = Pt(8.5)

doc.add_page_break()

# ============================================================
# SECTION 3: TEST SCRIPT
# ============================================================
doc.add_heading('Section 3: Validation & Test Script', level=1)
doc.add_paragraph(
    'After upgrading, run this script to verify GPT-5.4 is working and test the planner\'s capabilities.'
)

test_script = '''#!/bin/bash
# test-planner-gpt54.sh
# Validates the GPT-5.4 planner upgrade
# Run: bash test-planner-gpt54.sh

set -e
echo "=== GPT-5.4 Planner Validation ==="
echo ""

# 1. Verify config
echo "1. Checking planner config..."
python3 -c "
import json
with open('$HOME/.openclaw/openclaw.json') as f:
    c = json.load(f)
for a in c['agents']['list']:
    if a['id'] == 'planner':
        model = a.get('model', {})
        primary = model.get('primary', 'NOT SET')
        print(f'   Primary model: {primary}')
        print(f'   Fallbacks: {model.get(\"fallbacks\", [])}')
        if 'gpt-5.4' in primary:
            print('   ✅ GPT-5.4 configured')
        else:
            print('   ❌ GPT-5.4 NOT configured')
        break
"

echo ""

# 2. Verify API access
echo "2. Testing OpenAI API access for gpt-5.4..."
python3 -c "
import urllib.request, json, os

config_path = os.path.expanduser('~/.openclaw/openclaw.json')
with open(config_path) as f:
    c = json.load(f)

api_key = c.get('providers', {}).get('openai', {}).get('apiKey', '')
if not api_key:
    print('   ❌ No OpenAI API key found in config')
    exit(1)

# Quick model list check
req = urllib.request.Request(
    'https://api.openai.com/v1/models',
    headers={'Authorization': f'Bearer {api_key}'}
)
try:
    resp = urllib.request.urlopen(req)
    data = json.loads(resp.read())
    models = [m['id'] for m in data['data'] if '5.4' in m['id']]
    if models:
        print(f'   ✅ GPT-5.4 models available: {models}')
    else:
        print('   ⚠️  No gpt-5.4 models found. Check API key tier.')
except Exception as e:
    print(f'   ❌ API error: {e}')
"

echo ""

# 3. Quick planning test
echo "3. Running quick planning test..."
echo "   Spawning planner with a test task..."

# Use OpenClaw to test the planner agent
openclaw rpc agent.turn --agent planner --message \
  "Create a brief PLAN.md for adding a /health endpoint to a Node.js Express API. \
   Keep it under 20 lines. Output only the PLAN.md content." \
  2>/dev/null | head -30

echo ""
echo "=== Validation Complete ==="
echo ""
echo "If all checks passed, your planner is ready."
echo "Try: ask Jarvis to 'plan a new feature' and it will use GPT-5.4 Pro."
'''

p = doc.add_paragraph()
run = p.add_run(test_script)
run.font.name = 'Consolas'
run.font.size = Pt(8.5)

doc.add_page_break()

# ============================================================
# SECTION 4: COST COMPARISON
# ============================================================
doc.add_heading('Section 4: Cost Analysis & Recommendations', level=1)

doc.add_heading('Cost Comparison: Planning Models', level=2)

comparisons = [
    ('Claude Opus 4.6 (current)', '$15 / $75 per M tokens', '200K', 'Excellent reasoning, proven reliability'),
    ('GPT-5.4 Standard', '$2.50 / $15 per M tokens', '1M (opt-in)', 'Best value — 6x cheaper input, 5x cheaper output than Opus'),
    ('GPT-5.4 Pro', '$30 / $180 per M tokens', '1M (opt-in)', 'Maximum performance — 2x input cost of Opus but 5x the context'),
    ('Claude Sonnet 4.6 (fallback)', '$3 / $15 per M tokens', '200K', 'Good fallback, similar cost to GPT-5.4 standard'),
]

for model, price, ctx, notes in comparisons:
    p = doc.add_paragraph()
    run = p.add_run(f'{model}')
    run.bold = True
    p.add_run(f'\n    Price: {price}\n    Context: {ctx}\n    Notes: {notes}')

doc.add_paragraph()
doc.add_heading('Recommendation', level=2)

rec = doc.add_paragraph()
rec.add_run('Start with GPT-5.4 Standard (gpt-5.4) as primary').bold = True
rec.add_run(
    ' — it\'s 6x cheaper than Claude Opus on input tokens, offers 5x the context window, '
    'and benchmarks show it excels at exactly the kind of planning, decomposition, and professional '
    'work we need. Use GPT-5.4 Pro only for complex architecture decisions or when standard '
    'quality isn\'t sufficient.\n\n'
    'Suggested model chain for the planner:\n'
)
chain = [
    'Primary: openai/gpt-5.4 ($2.50/$15 per M) — daily planning tasks',
    'Fallback 1: openai/gpt-5.4-pro ($30/$180 per M) — complex architecture',
    'Fallback 2: anthropic/claude-opus-4-6 ($15/$75 per M) — if OpenAI is down',
]
for item in chain:
    doc.add_paragraph(f'  • {item}')

doc.add_paragraph()
cost_note = doc.add_paragraph()
cost_note.add_run('⚠️ Cost guard: ').bold = True
cost_note.add_run(
    'With the 1M context window, a single planning session could theoretically consume $30+ in input tokens alone '
    'on Pro. The config script sets maxOutputTokens to 32K to keep output costs reasonable. '
    'Monitor usage at platform.openai.com/usage.'
)

doc.add_page_break()

# ============================================================
# SECTION 5: WHAT TO PASTE
# ============================================================
doc.add_heading('Quick Start: What To Do', level=1)

steps = [
    ('Get your OpenAI API key', 'Go to platform.openai.com/api-keys. Make sure your account has access to gpt-5.4 models (requires a funded API account).'),
    ('Run the upgrade script (Section 1)', 'Copy the bash script to your Mac, run it, and paste your API key when prompted. It will update openclaw.json and restart the gateway.'),
    ('Install the planner prompt (Section 2)', 'Copy the AGENTS.md content into ~/.openclaw/workspace-planner/AGENTS.md. This gives the planner its enhanced behavior.'),
    ('Validate (Section 3)', 'Run the test script to confirm GPT-5.4 is accessible and the planner is responding.'),
    ('Test with a real task', 'Ask Jarvis: "Plan a new monitoring dashboard for Cohesity SaaS metrics" — the planner should produce a detailed PLAN.md with task decomposition.'),
]

for i, (title_text, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {i}: {title_text}')
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph(desc)
    doc.add_paragraph()

# Save
output_path = '/Users/ericbrown/.openclaw/workspace/GPT54_Enhanced_Planning_Agent.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
